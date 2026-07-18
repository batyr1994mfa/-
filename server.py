#!/usr/bin/env python3
"""
Локальный веб-интерфейс для Ollama — бэкенд.

Использует тот же файл настроек (~/.config/ollama-terminal/config.json),
что и ollama_terminal.py, поэтому папки/модели, подключённые в терминале,
сразу видны и в браузере, и наоборот.

Запуск:
    pip install -r requirements-web.txt
    python3 server.py
Затем откройте http://localhost:5050 в браузере.
"""

import base64
import getpass
import glob
import html as html_lib
import json
import os
import re
import shutil
import subprocess
import sys
import shlex
import tempfile
import threading
import time
import traceback
import uuid
from pathlib import Path

import requests
from flask import Flask, Response, jsonify, request, send_from_directory

try:
    import psutil
except ImportError:
    psutil = None

# ----------------------------- Настройки -----------------------------------

DEFAULT_OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434")

REASONING_LEVELS = {
    "low":    {"label": "Творческий",  "temperature": 1.0},
    "medium": {"label": "Сбалансированный", "temperature": 0.8},
    "high":   {"label": "Строгий", "temperature": 0.4},
    "max":    {"label": "Точный",    "temperature": 0.0},
}
DEFAULT_REASONING_LEVEL = "low"
DEFAULT_NUM_CTX = 8192
WEB_PORT = int(os.environ.get("WEB_PORT", "5050"))

MAX_FILE_READ_BYTES = 200_000
MAX_IMAGE_READ_BYTES = 10_000_000
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
MAX_UPLOAD_FILE_BYTES = 500_000  # 500 КБ — лимит на текстовый файл при вложении
MAX_UPLOAD_IMAGE_BYTES = 10_000_000  # 10 МБ — лимит на изображение при вложении
MAX_TOOL_ITERATIONS = 12

CONFIG_DIR = Path.home() / ".config" / "ollama-terminal"
CONFIG_FILE = CONFIG_DIR / "config.json"
CHATS_FILE = CONFIG_DIR / "chats.json"
TASKS_FILE = CONFIG_DIR / "tasks.json"
TERMINAL_COMMANDS_FILE = CONFIG_DIR / "terminal_commands.json"
TERMINAL_HISTORY_FILE = CONFIG_DIR / "terminal_history.json"

app = Flask(__name__, static_folder="static", static_url_path="")


# ----------------------------- Состояние (тот же конфиг, что и в терминале) --

class State:
    def __init__(self):
        self.model = "llama3.1"
        self.vision_model = None
        self.allowed_folders: list[Path] = []
        self.auto_approve = False
        self.ollama_host = DEFAULT_OLLAMA_HOST
        self.reasoning_level = DEFAULT_REASONING_LEVEL
        self.active_cancel = {"cancelled": False}
        self.num_ctx = DEFAULT_NUM_CTX
        self.use_cpu = True
        self.use_ram = True
        self.use_vram = True
        # "Режим размышлений" (thinking) — у reasoning-моделей (Qwen3.5, DeepSeek-R1
        # и т.п.) его НЕЛЬЗЯ отключить через PARAMETER в Modelfile (это ограничение
        # самого Ollama: https://github.com/ollama/ollama/issues/14809) — единственный
        # рабочий способ — явно передавать "think" в каждом запросе к /api/chat.
        # По умолчанию выключено: без этого внутренние рассуждения модели (обычно
        # на английском, даже если в SYSTEM указан строго русский язык) могут
        # утекать в видимый ответ.
        self.think_enabled = False
        self.sudo_password = ""
        # "Автономность ИИ" — даёт модели доступ к инструменту run_shell_command,
        # позволяющему выполнять произвольные команды в терминале системы
        # (в т.ч. с sudo, если передан пароль в настройках). Мощная и опасная
        # возможность, поэтому по умолчанию выключена и требует явного
        # включения пользователем через кнопку в боковой панели.
        self.ai_autonomy = False
        # Google Custom Search JSON API — для полноценного поиска через Google
        # (ключ и поисковый движок пользователь создаёт бесплатно на
        # console.cloud.google.com и programmablesearchengine.google.com).
        # Без них google_search тихо использует DuckDuckGo как запасной вариант.
        self.google_api_key = ""
        self.google_cx = ""
        # SMTP для инструмента send_email
        self.smtp_host = ""
        self.smtp_port = 587
        self.smtp_user = ""
        self.smtp_password = ""
        self.smtp_from = ""
        # Telegram-бот для инструмента telegram_send и уведомлений
        self.telegram_bot_token = ""
        self.telegram_chat_id = ""
        # ntfy.sh topic для push-уведомлений (notify)
        self.ntfy_topic = ""
        # Заметки долгосрочной памяти ИИ (remember/recall): {key: text}
        self.memory_notes = {}
        # Папка архива успешных решений (save_solution/recall_solutions) —
        # если пусто, используется путь по умолчанию (см. _solutions_dir()).
        self.solutions_archive_path = ""
        # Приоритет источников знаний, задаваемый пользователем: список из
        # "internet", "archive", "memory" в желаемом порядке (первый — самый
        # приоритетный). Влияет только на порядок, в котором модели
        # рекомендуется их проверять — сами инструменты остаются доступны все.
        self.knowledge_priority = ["internet", "archive", "memory"]
        self.chats: dict = {}          # id -> {id, title, messages, created_at, updated_at}
        self.current_chat_id: str = None
        # Запись истории чатов на диск — довольно тяжёлая операция (полная
        # JSON-сериализация ВСЕХ чатов), а touch()/push_activity() теперь
        # вызываются очень часто (каждые ~15 токенов, на каждый вызов
        # инструмента и т.п.). Если делать реальную запись на каждый такой
        # вызов синхронно, это удерживает GIL и подвисает вообще все запросы
        # к серверу — вплоть до кнопок "Перезапустить"/"Остановить". Поэтому
        # save_chats() теперь только помечает состояние как "грязное", а
        # реальная запись на диск идёт из фонового потока не чаще раза в
        # секунду (см. _chats_saver_loop).
        self._chats_dirty = False
        self._chats_save_lock = threading.Lock()
        self.load()
        self.load_chats()

    def load_chats(self):
        if CHATS_FILE.exists():
            try:
                data = json.loads(CHATS_FILE.read_text(encoding="utf-8"))
                self.chats = data.get("chats", {}) or {}
                self.current_chat_id = data.get("current_chat_id")
            except Exception:
                self.chats = {}
        if not self.chats or self.current_chat_id not in self.chats:
            self.new_chat()

    def save_chats(self):
        # Не пишем на диск прямо сейчас — только помечаем, что есть
        # несохранённые изменения. Реальную запись делает фоновый поток
        # (см. _chats_saver_loop), чтобы частые вызовы touch()/push_activity()
        # во время генерации не тормозили сервер.
        self._chats_dirty = True

    def _save_chats_now(self):
        with self._chats_save_lock:
            try:
                CONFIG_DIR.mkdir(parents=True, exist_ok=True)
                data = {"current_chat_id": self.current_chat_id, "chats": self.chats}
                # Без indent=2: с ростом истории чатов (много сообщений и
                # карточек активности с превью файлов) форматированная
                # сериализация становится заметно медленнее компактной, а
                # человекочитаемость файла тут не критична.
                CHATS_FILE.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            except Exception:
                pass

    def new_chat(self, title="Новый диалог"):
        cid = uuid.uuid4().hex[:12]
        now = time.time()
        self.chats[cid] = {"id": cid, "title": title, "messages": [], "activity": [], "created_at": now, "updated_at": now}
        self.current_chat_id = cid
        self.save_chats()
        return cid

    def current_messages(self) -> list:
        return self.chats[self.current_chat_id]["messages"]

    def messages_for(self, chat_id: str) -> list:
        """Как current_messages(), но для конкретного чата — используется во
        время генерации ответа, чтобы результат не 'уплыл' в другой чат, если
        пользователь переключился на него, пока ИИ ещё печатал."""
        chat = self.chats.get(chat_id)
        if chat is None:
            # чат успели удалить, пока шла генерация — создаём временный, чтобы не упасть
            chat = {"id": chat_id, "title": "Новый диалог", "messages": [],
                    "created_at": time.time(), "updated_at": time.time()}
            self.chats[chat_id] = chat
        return chat["messages"]

    def touch_current(self, maybe_title: str = None):
        chat = self.chats[self.current_chat_id]
        chat["updated_at"] = time.time()
        if maybe_title and chat["title"] == "Новый диалог":
            chat["title"] = maybe_title.strip()[:48] or "Новый диалог"
        self.save_chats()

    def touch(self, chat_id: str, maybe_title: str = None):
        chat = self.chats.get(chat_id)
        if chat is None:
            return
        chat["updated_at"] = time.time()
        if maybe_title and chat["title"] == "Новый диалог":
            chat["title"] = maybe_title.strip()[:48] or "Новый диалог"
        self.save_chats()

    def load(self):
        if not CONFIG_FILE.exists():
            return
        try:
            data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
        except Exception:
            return
        self.model = data.get("model", self.model)
        self.vision_model = data.get("vision_model", self.vision_model)
        folders = data.get("allowed_folders")
        if folders is None:
            single = data.get("allowed_folder")
            folders = [single] if single else []
        loaded = []
        for f in folders:
            p = Path(f)
            if p.exists() and p.is_dir():
                loaded.append(p.resolve())
        if loaded:
            self.allowed_folders = loaded
        self.auto_approve = data.get("auto_approve", self.auto_approve)
        self.ollama_host = data.get("ollama_host", self.ollama_host)
        level = data.get("reasoning_level")
        if level in REASONING_LEVELS:
            self.reasoning_level = level
        try:
            nc = int(data.get("num_ctx", self.num_ctx))
            if nc > 0:
                self.num_ctx = nc
        except (TypeError, ValueError):
            pass
        self.use_cpu = bool(data.get("use_cpu", self.use_cpu))
        self.use_ram = bool(data.get("use_ram", self.use_ram))
        self.use_vram = bool(data.get("use_vram", self.use_vram))
        if not (self.use_cpu or self.use_ram or self.use_vram):
            self.use_cpu = self.use_ram = self.use_vram = True
        self.think_enabled = bool(data.get("think_enabled", self.think_enabled))
        self.sudo_password = data.get("sudo_password", self.sudo_password)
        self.ai_autonomy = bool(data.get("ai_autonomy", self.ai_autonomy))
        self.google_api_key = data.get("google_api_key", self.google_api_key)
        self.google_cx = data.get("google_cx", self.google_cx)
        self.smtp_host = data.get("smtp_host", self.smtp_host)
        self.smtp_port = int(data.get("smtp_port", self.smtp_port) or 587)
        self.smtp_user = data.get("smtp_user", self.smtp_user)
        self.smtp_password = data.get("smtp_password", self.smtp_password)
        self.smtp_from = data.get("smtp_from", self.smtp_from)
        self.telegram_bot_token = data.get("telegram_bot_token", self.telegram_bot_token)
        self.telegram_chat_id = data.get("telegram_chat_id", self.telegram_chat_id)
        self.ntfy_topic = data.get("ntfy_topic", self.ntfy_topic)
        self.memory_notes = data.get("memory_notes", self.memory_notes) or {}
        self.solutions_archive_path = data.get("solutions_archive_path", self.solutions_archive_path)
        kp = data.get("knowledge_priority")
        if isinstance(kp, list) and sorted(kp) == sorted(["internet", "archive", "memory"]):
            self.knowledge_priority = kp

    def compute_num_gpu(self):
        """Возвращает значение параметра num_gpu для Ollama, либо None ('авто').
        CPU и RAM физически неразделимы для Ollama (без GPU всё считается на
        процессоре и лежит в оперативной памяти), поэтому они управляются как
        одна пара; отдельный реальный рычаг у Ollama есть только для GPU/VRAM."""
        if self.use_cpu and self.use_ram and self.use_vram:
            return None  # авто — пусть Ollama сама решает, это стандартное поведение
        if self.use_vram and not self.use_cpu and not self.use_ram:
            return 999  # максимально на видеокарту
        if not self.use_vram:
            return 0  # только CPU/RAM, GPU не используется
        return None  # смешанный случай — доверяем автоматике Ollama

    def reasoning_options(self) -> dict:
        params = REASONING_LEVELS.get(self.reasoning_level, REASONING_LEVELS[DEFAULT_REASONING_LEVEL])
        opts = {"temperature": params["temperature"], "num_ctx": self.num_ctx}
        num_gpu = self.compute_num_gpu()
        if num_gpu is not None:
            opts["num_gpu"] = num_gpu
        return opts

    def save(self):
        try:
            CONFIG_DIR.mkdir(parents=True, exist_ok=True)
            data = {
                "model": self.model,
                "vision_model": self.vision_model,
                "allowed_folders": [str(f) for f in self.allowed_folders],
                "auto_approve": self.auto_approve,
                "ollama_host": self.ollama_host,
                "reasoning_level": self.reasoning_level,
                "num_ctx": self.num_ctx,
                "use_cpu": self.use_cpu,
                "use_ram": self.use_ram,
                "use_vram": self.use_vram,
                "think_enabled": self.think_enabled,
                "sudo_password": self.sudo_password,
                "ai_autonomy": self.ai_autonomy,
                "google_api_key": self.google_api_key,
                "google_cx": self.google_cx,
                "smtp_host": self.smtp_host,
                "smtp_port": self.smtp_port,
                "smtp_user": self.smtp_user,
                "smtp_password": self.smtp_password,
                "smtp_from": self.smtp_from,
                "telegram_bot_token": self.telegram_bot_token,
                "telegram_chat_id": self.telegram_chat_id,
                "ntfy_topic": self.ntfy_topic,
                "memory_notes": self.memory_notes,
                "solutions_archive_path": self.solutions_archive_path,
                "knowledge_priority": self.knowledge_priority,
            }
            CONFIG_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass


state = State()

# id ожидающих подтверждения операций записи/правки файлов -> {"event", "decision"}
pending_confirmations: dict[str, dict] = {}


# ----------------------------- Система заданий --------------------------------

class TasksStore:
    def __init__(self):
        self.tasks: dict = {}  # id -> task
        self.load()

    def load(self):
        if TASKS_FILE.exists():
            try:
                data = json.loads(TASKS_FILE.read_text(encoding="utf-8"))
                self.tasks = data.get("tasks", {}) or {}
            except Exception:
                self.tasks = {}

    def save(self):
        try:
            CONFIG_DIR.mkdir(parents=True, exist_ok=True)
            TASKS_FILE.write_text(
                json.dumps({"tasks": self.tasks}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        except Exception:
            pass

    def create(self, title: str, description: str, files: list = None,
               scheduled_time: str = None, repeat_daily: bool = False,
               command: str = None, sudo_password: str = None,
               run_file: str = None) -> dict:
        tid = uuid.uuid4().hex[:12]
        now = time.time()
        task = {
            "id": tid,
            "title": title,
            "description": description,
            "files": files or [],
            "status": "pending",
            "steps": [],
            "result": None,
            "scheduled_time": scheduled_time,
            "repeat_daily": repeat_daily,
            "command": command,
            "sudo_password": sudo_password,
            "run_file": run_file,
            "last_run": None,
            "created_at": now,
            "updated_at": now,
        }
        self.tasks[tid] = task
        self.save()
        return task

    def get(self, tid: str) -> dict | None:
        return self.tasks.get(tid)

    def list_all(self) -> list:
        items = sorted(self.tasks.values(), key=lambda t: t["created_at"], reverse=True)
        return items

    def update(self, tid: str, **kwargs) -> dict | None:
        task = self.tasks.get(tid)
        if not task:
            return None
        for k, v in kwargs.items():
            task[k] = v
        task["updated_at"] = time.time()
        self.save()
        return task

    def delete(self, tid: str) -> bool:
        if tid in self.tasks:
            del self.tasks[tid]
            self.save()
            return True
        return False

    def add_step(self, tid: str, text: str) -> dict | None:
        task = self.tasks.get(tid)
        if not task:
            return None
        step = {"text": text, "status": "pending"}
        task["steps"].append(step)
        task["updated_at"] = time.time()
        self.save()
        return step

    def update_step(self, tid: str, step_idx: int, status: str) -> bool:
        task = self.tasks.get(tid)
        if not task or step_idx < 0 or step_idx >= len(task["steps"]):
            return False
        task["steps"][step_idx]["status"] = status
        task["updated_at"] = time.time()
        self.save()
        return True


tasks_store = TasksStore()
_last_scheduled_chat_id = None


# ----------------------------- Песочница для файлов --------------------------

class SandboxError(Exception):
    pass


def safe_path(rel_path: str) -> Path:
    if not state.allowed_folders:
        raise SandboxError("Ни одна папка не подключена.")
    p = Path(rel_path).expanduser()
    if p.is_absolute():
        candidates = [p.resolve()]
    else:
        candidates = [(base / p).resolve() for base in state.allowed_folders]
        existing = [c for c in candidates if c.exists()]
        if existing:
            candidates = existing
    for candidate in candidates:
        for base in state.allowed_folders:
            try:
                candidate.relative_to(base.resolve())
                return candidate
            except ValueError:
                continue
    folders_list = ", ".join(str(f) for f in state.allowed_folders)
    raise SandboxError(f"Путь '{rel_path}' вне разрешённых папок ({folders_list}).")


# ----------------------------- Инструменты (те же, что в терминале) ----------

def tool_list_dir(path: str = "") -> str:
    if not path and len(state.allowed_folders) != 1:
        if not state.allowed_folders:
            return "Ни одна папка не подключена."
        return "Подключено несколько папок:\n" + "\n".join(f"[ROOT] {f}" for f in state.allowed_folders)
    target = safe_path(path or str(state.allowed_folders[0]))
    if not target.exists():
        return f"Ошибка: путь не существует: {path or '.'}"
    if target.is_file():
        return f"Это файл, а не папка: {path}"
    entries = []
    for item in sorted(target.iterdir()):
        kind = "DIR " if item.is_dir() else "FILE"
        size = "" if item.is_dir() else f" ({item.stat().st_size} байт)"
        entries.append(f"[{kind}] {item.name}{size}")
    return "\n".join(entries) if entries else "(папка пуста)"


def tool_read_file(path: str) -> str:
    target = safe_path(path)
    if not target.exists() or not target.is_file():
        return f"Ошибка: файл не найден: {path}"
    if target.stat().st_size > MAX_FILE_READ_BYTES:
        return f"Ошибка: файл слишком большой (> {MAX_FILE_READ_BYTES} байт)."
    try:
        return target.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"Ошибка чтения файла: {e}"


def tool_search_files(query: str, path: str = "") -> str:
    roots = [safe_path(path)] if path else list(state.allowed_folders)
    if not roots:
        return "Ни одна папка не подключена."
    matches = []
    for root_dir in roots:
        if not root_dir.exists():
            continue
        for root, _dirs, files in os.walk(root_dir):
            for fname in files:
                fpath = Path(root) / fname
                try:
                    if fpath.stat().st_size > MAX_FILE_READ_BYTES:
                        continue
                    text = fpath.read_text(encoding="utf-8", errors="ignore")
                except Exception:
                    continue
                if query.lower() in text.lower():
                    rel = fpath.relative_to(root_dir)
                    for i, line in enumerate(text.splitlines(), start=1):
                        if query.lower() in line.lower():
                            matches.append(f"[{root_dir.name}] {rel}:{i}: {line.strip()[:200]}")
                if len(matches) > 100:
                    break
    return "\n".join(matches[:100]) if matches else f"Совпадений по '{query}' не найдено."


TOOL_DESCRIPTIONS_PY = {
    "list_dir":          {"icon": "📂", "type": "read",   "verb": "Просмотр папки",       "desc": "ИИ смотрит содержимое директории"},
    "read_file":         {"icon": "◉",  "type": "read",   "verb": "Чтение файла",         "desc": "ИИ читает содержимое файла"},
    "search_files":      {"icon": "⌕",  "type": "search", "verb": "Поиск по файлам",      "desc": "ИИ ищет текст во файлах проекта"},
    "write_file":        {"icon": "＋",  "type": "create", "verb": "Создание файла",       "desc": "ИИ записывает новый файл или перезаписывает существующий"},
    "edit_file":         {"icon": "✎",  "type": "edit",   "verb": "Редактирование",       "desc": "ИИ изменяет часть содержимого файла"},
    "view_image":        {"icon": "👁",  "type": "read",   "verb": "Просмотр изображения", "desc": "ИИ анализирует изображение"},
    "run_shell_command": {"icon": "💻", "type": "exec",   "verb": "Команда в терминале",  "desc": "ИИ выполняет команду в системном терминале"},
    "web_search":        {"icon": "🌐", "type": "search", "verb": "Поиск в интернете",    "desc": "ИИ ищет информацию в интернете"},
    "fetch_url":         {"icon": "🔗", "type": "read",   "verb": "Открытие страницы",    "desc": "ИИ открывает и читает веб-страницу"},
    "wikipedia_search":  {"icon": "📖", "type": "search", "verb": "Поиск в Википедии",    "desc": "ИИ ищет статьи в Википедии для фоновых знаний"},
    "google_search":     {"icon": "🔎", "type": "search", "verb": "Поиск в Google",       "desc": "ИИ ищет информацию через Google"},
    "github_search":     {"icon": "🐙", "type": "search", "verb": "Поиск на GitHub",      "desc": "ИИ ищет репозитории на GitHub"},
    "image_search":      {"icon": "🖼", "type": "search", "verb": "Поиск изображений",    "desc": "ИИ ищет фотографии/картинки"},
    "video_search":      {"icon": "🎬", "type": "search", "verb": "Поиск видео",          "desc": "ИИ ищет видео на YouTube"},
    "download_file":     {"icon": "⬇",  "type": "create", "verb": "Скачивание файла",     "desc": "ИИ скачивает файл по ссылке на диск"},
    "read_document":     {"icon": "📄", "type": "read",   "verb": "Чтение документа",     "desc": "ИИ читает PDF/DOCX/XLSX"},
    "ocr_image":         {"icon": "🔍", "type": "read",   "verb": "Распознавание текста", "desc": "ИИ распознаёт текст на изображении (OCR)"},
    "archive_extract":   {"icon": "📦", "type": "create", "verb": "Распаковка архива",    "desc": "ИИ распаковывает архив"},
    "archive_create":    {"icon": "📦", "type": "create", "verb": "Создание архива",      "desc": "ИИ упаковывает файлы в архив"},
    "diff_files":        {"icon": "⇄",  "type": "read",   "verb": "Сравнение файлов",     "desc": "ИИ сравнивает два файла"},
    "remember":          {"icon": "🧷", "type": "create", "verb": "Запись в память",      "desc": "ИИ сохраняет заметку в долгосрочную память"},
    "recall":            {"icon": "🧷", "type": "read",   "verb": "Чтение памяти",        "desc": "ИИ читает заметки из долгосрочной памяти"},
    "save_solution":     {"icon": "💾", "type": "create", "verb": "Архив решений: запись", "desc": "ИИ сохраняет успешное решение в архив"},
    "recall_solutions":  {"icon": "💾", "type": "read",   "verb": "Архив решений: поиск",  "desc": "ИИ ищет похожие задачи в архиве решений"},
    "process_list":      {"icon": "📊", "type": "read",   "verb": "Список процессов",     "desc": "ИИ смотрит запущенные процессы"},
    "check_process_running": {"icon": "🔍", "type": "read", "verb": "Проверка процесса",  "desc": "ИИ проверяет, запущено ли уже приложение"},
    "git_tool":          {"icon": "🔧", "type": "exec",   "verb": "Git-операция",         "desc": "ИИ выполняет git-команду"},
    "systemd_control":   {"icon": "⚙",  "type": "exec",   "verb": "Управление службой",   "desc": "ИИ управляет системной службой"},
    "notify":            {"icon": "🔔", "type": "create", "verb": "Push-уведомление",     "desc": "ИИ отправляет push-уведомление"},
    "send_email":        {"icon": "✉",  "type": "create", "verb": "Отправка письма",      "desc": "ИИ отправляет email"},
    "telegram_send":     {"icon": "✈",  "type": "create", "verb": "Сообщение в Telegram", "desc": "ИИ отправляет сообщение в Telegram"},
    "python_repl":       {"icon": "🐍", "type": "exec",   "verb": "Python-код",           "desc": "ИИ выполняет Python-код"},
    "lint_format":       {"icon": "🧹", "type": "read",   "verb": "Проверка кода",        "desc": "ИИ проверяет/форматирует код линтером"},
    "get_weather":       {"icon": "🌦", "type": "read",   "verb": "Погода",               "desc": "ИИ смотрит текущую погоду"},
    "get_exchange_rate": {"icon": "💱", "type": "read",   "verb": "Курс валют",           "desc": "ИИ смотрит курс валют"},
    "rss_read":          {"icon": "📰", "type": "read",   "verb": "Чтение RSS",           "desc": "ИИ читает RSS/Atom-ленту"},
    "set_reminder":      {"icon": "⏰", "type": "create", "verb": "Напоминание",          "desc": "ИИ ставит напоминание"},
}


def _act_now() -> str:
    return time.strftime("%H:%M:%S")


def _act_esc(s) -> str:
    return html_lib.escape(str(s if s is not None else ""))


def push_activity(chat_id: str, card_html: str):
    """Добавить готовую HTML-карточку в персистентный лог активности чата —
    в том же формате, в котором фронтенд сохраняет их через saveActivityEntry,
    чтобы карточки одинаково отображались после перезагрузки/переключения
    чата независимо от того, откуда они пришли (интерактивный чат, ручной
    запуск задания или задание по расписанию)."""
    chat = state.chats.get(chat_id)
    if chat is None:
        return
    chat.setdefault("activity", [])
    chat["activity"].append({"html": card_html, "time": time.time() * 1000})
    if len(chat["activity"]) > 200:
        chat["activity"] = chat["activity"][-200:]
    chat["updated_at"] = time.time()
    state.save_chats()


def activity_card_tool_call(name: str, args: dict) -> str:
    path = (args or {}).get("path", "")
    query = (args or {}).get("query", "")
    url = (args or {}).get("url", "")
    misc = (args or {}).get("location") or (args or {}).get("service") or (args or {}).get("key") \
        or (args or {}).get("repo_path") or (args or {}).get("message") or ""
    info = TOOL_DESCRIPTIONS_PY.get(name, {"icon": "⚡", "type": "info", "verb": name, "desc": ""})
    target = path or query or url or misc

    extra_desc = ""
    if name == "write_file" and args:
        content = args.get("content", "")
        lines = content.count("\n") + 1
        extra_desc = f" — {lines} строк, ~{len(content)} символов"
    elif name == "edit_file" and args:
        old_t = (args.get("old_text", "") or "")[:60]
        suffix = "..." if len(args.get("old_text", "")) >= 60 else ""
        extra_desc = f' — замена "{old_t}{suffix}"'
    elif name == "search_files" and query:
        extra_desc = f' — запрос: "{query}"'
    elif name == "run_shell_command" and args:
        cmd = (args.get("command", "") or "")[:80]
        suffix = "..." if len(args.get("command", "") or "") > 80 else ""
        sudo_note = " (sudo)" if args.get("use_sudo") else ""
        extra_desc = f' — $ {cmd}{suffix}{sudo_note}'
    elif name == "web_search" and query:
        extra_desc = f' — запрос: "{query}"'
    elif name == "wikipedia_search" and query:
        extra_desc = f' — запрос: "{query}"'
    elif name == "google_search" and query:
        extra_desc = f' — запрос: "{query}"'
    elif name == "github_search" and query:
        extra_desc = f' — запрос: "{query}"'
    elif name == "download_file" and args:
        extra_desc = f' — с {args.get("url", "")}'
    elif name == "git_tool" and args:
        extra_desc = f' — git {args.get("action", "")}'
    elif name == "systemd_control" and args:
        extra_desc = f' — {args.get("action", "")} {args.get("service", "")}'
    elif name == "python_repl" and args:
        code = (args.get("code", "") or "")[:80]
        extra_desc = f' — {code}{"..." if len(args.get("code","")) > 80 else ""}'
    elif name == "get_exchange_rate" and args:
        extra_desc = f' — {args.get("from_currency","")} → {args.get("to_currency","")}'

    return (
        f'<div class="act-card act-{info["type"]}">'
        f'<div class="act-head">'
        f'<span class="act-icon">{info["icon"]}</span>'
        f'<span class="act-label">{_act_esc(info["verb"])}</span>'
        f'<span class="act-time">{_act_now()}</span>'
        f'</div>'
        f'<div class="act-desc">{_act_esc(info["desc"])}{_act_esc(extra_desc)}</div>'
        + (f'<div class="act-path">→ {_act_esc(target)}</div>' if target else "")
        + '</div>'
    )


def activity_card_file_change(kind: str, path: str, preview: str, auto_approved: bool) -> str:
    if kind == "write":
        is_new = (not preview) or ("успешно создан" in preview)
        card_class, icon = "act-create", "＋"
        verb = "Файл создан" if is_new else "Файл перезаписан"
        desc = "ИИ создал новый файл" if is_new else "ИИ перезаписал содержимое файла"
    elif kind == "edit":
        card_class, icon, verb, desc = "act-edit", "✎", "Файл изменён", "ИИ отредактировал содержимое файла"
    elif kind == "exec":
        card_class, icon, verb, desc = "act-exec", "💻", "Команда выполнена", "ИИ выполнил команду в терминале"
    elif kind == "download":
        card_class, icon, verb, desc = "act-create", "⬇", "Файл скачан", "ИИ скачал файл по ссылке"
    else:
        card_class, icon, verb, desc = "act-info", "⚡", kind, ""

    status_html = ('<span class="act-status ok">✔ подтверждено автоматически</span>' if auto_approved
                   else '<span class="act-status wait">🧠 ожидает вашего подтверждения</span>')
    preview_html = (f'<div class="act-preview">{_act_esc(preview)}</div>' if preview
                     else '<div class="act-preview" style="display:none"></div>')
    actions_html = ""
    if kind in ("write", "edit"):
        actions_html = (
            '<div class="act-mini-actions">'
            f'<button class="act-mini-btn act-open-folder-btn" data-path="{_act_esc(path)}">📂 Открыть папку</button>'
            f'<button class="act-mini-btn act-run-file-btn" data-path="{_act_esc(path)}">▶ Запустить файл</button>'
            '<span class="act-mini-status"></span>'
            '</div>'
        )

    return (
        f'<div class="act-card {card_class}">'
        f'<div class="act-head">'
        f'<span class="act-icon">{icon}</span>'
        f'<span class="act-label">{_act_esc(verb)}</span>'
        f'<span class="act-time">{_act_now()}</span>'
        f'</div>'
        f'<div class="act-desc">{_act_esc(desc)}</div>'
        f'<div class="act-path">→ {_act_esc(path)}</div>'
        f'{status_html}{preview_html}{actions_html}'
        '</div>'
    )


def activity_card_result(text: str, ok: bool) -> str:
    value = text or ""
    is_long = len(value) > 160 or "\n" in value
    cls = "act-ok" if ok else "act-error"
    icon = "✔" if ok else "✖"

    written_path = None
    if ok:
        m = re.search(r"Файл '(.+?)' успешно (?:создан|перезаписан|отредактирован)\.?$", value)
        written_path = m.group(1) if m else None
    actions_html = ""
    if written_path:
        actions_html = (
            '<div class="act-mini-actions">'
            f'<button class="act-mini-btn act-open-folder-btn" data-path="{_act_esc(written_path)}">📂 Открыть папку</button>'
            f'<button class="act-mini-btn act-run-file-btn" data-path="{_act_esc(written_path)}">▶ Запустить файл</button>'
            '<span class="act-mini-status"></span>'
            '</div>'
        )

    if not is_long:
        label = value if ok else f"Ошибка: {value}"
        return (
            f'<div class="act-card {cls}">'
            f'<div class="act-head">'
            f'<span class="act-icon">{icon}</span>'
            f'<span class="act-label">{_act_esc(label)}</span>'
            f'<span class="act-time">{_act_now()}</span>'
            f'</div>{actions_html}</div>'
        )
    label = "Результат получен" if ok else "Ошибка выполнения"
    return (
        f'<div class="act-card {cls}">'
        f'<div class="act-head">'
        f'<span class="act-icon">{icon}</span>'
        f'<span class="act-label">{_act_esc(label)}</span>'
        f'<span class="act-time">{_act_now()}</span>'
        f'</div>'
        f'<div class="act-preview">{_act_esc(value)}</div>'
        f'{actions_html}'
        '</div>'
    )


def make_persisting_emit(chat_id: str):
    """Emit-функция для фоновых процессов (задания по расписанию, ручной запуск
    задания без активного SSE-слушателя) — вместо отправки события в браузер
    сразу строит ту же самую HTML-карточку, что рисует фронтенд для
    интерактивного чата, и сохраняет её в постоянный лог активности чата.
    Так панель "Активность в файлах" одинаково работает независимо от того,
    как был запущен ИИ."""
    def emit(event: dict):
        etype = event.get("type")
        if etype == "file_change":
            card = activity_card_file_change(event.get("kind", ""), event.get("path", ""),
                                              event.get("preview", ""), bool(event.get("auto_approved")))
            push_activity(chat_id, card)
        # confirm_request здесь не обрабатываем — для фоновых задач подтверждение
        # запрашивать некому, см. auto_approve.
    return emit


def request_confirmation(kind: str, path: str, preview: str, emit, force_auto: bool = False) -> bool:
    """Показывает, что именно ИИ собирается сделать с файлом, и либо сразу
    подтверждает (если включено автоподтверждение), либо ждёт решения пользователя.
    force_auto=True используется для фоновых процессов (задание по расписанию
    или ручной запуск задания без открытого чата) — там реального человека,
    который мог бы нажать "разрешить", может не быть, и ожидание решения
    привело бы к зависанию на 10 минут."""
    if state.auto_approve or force_auto:
        emit({"type": "file_change", "kind": kind, "path": path, "preview": preview, "auto_approved": True})
        return True
    cid = str(uuid.uuid4())
    ev = threading.Event()
    pending_confirmations[cid] = {"event": ev, "decision": None}
    emit({"type": "confirm_request", "id": cid, "kind": kind, "path": path, "preview": preview})
    got = ev.wait(timeout=600)  # 10 минут на ответ
    entry = pending_confirmations.pop(cid, {"decision": False})
    if not got:
        return False
    return bool(entry.get("decision"))


def tool_write_file(path: str, content: str, emit, force_auto: bool = False) -> str:
    target = safe_path(path)
    is_new = not target.exists()
    approved = request_confirmation("write", path, content, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ операцию записи."
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return f"Файл '{path}' успешно {'создан' if is_new else 'перезаписан'}."


def tool_edit_file(path: str, old_text: str, new_text: str, emit, force_auto: bool = False) -> str:
    target = safe_path(path)
    if not target.exists() or not target.is_file():
        return f"Ошибка: файл не найден: {path}"
    original = target.read_text(encoding="utf-8", errors="replace")
    count = original.count(old_text)
    if count == 0:
        return "Ошибка: фрагмент 'old_text' не найден дословно в файле."
    if count > 1:
        return f"Ошибка: фрагмент встречается {count} раз(а), уточните, чтобы был уникальным."
    diff_preview = f"--- было ---\n{old_text}\n--- станет ---\n{new_text}"
    approved = request_confirmation("edit", path, diff_preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ операцию редактирования."
    updated = original.replace(old_text, new_text, 1)
    target.write_text(updated, encoding="utf-8")
    return f"Файл '{path}' успешно отредактирован."


MAX_SHELL_OUTPUT_CHARS = 8000
SHELL_COMMAND_TIMEOUT = 120


def tool_run_shell(command: str, emit, use_sudo: bool = False, force_auto: bool = False) -> str:
    """Выполняет команду в системном терминале (bash), в обход песочницы
    allowed_folders — сознательно даёт ИИ полный доступ к системе. Доступен
    модели, только если включена 'Автономность ИИ' (state.ai_autonomy).
    При use_sudo=True выполняет через 'sudo -S' с сохранённым паролем
    (см. настройки/state.sudo_password).

    ВАЖНО про фоновые процессы: раньше вывод захватывался через
    subprocess.run(capture_output=True), который читает stdout/stderr через
    пайп ДО получения EOF. Если команда запускает что-то в фоне через '&' без
    собственного редиректа (например GUI-приложение), тот фоновый процесс
    наследует те же файловые дескрипторы пайпа и держит их открытыми всё
    время своей жизни (часы) — из-за этого чтение зависало и упиралось в
    таймаут, ХОТЯ сама команда уже успешно отработала и приложение
    запустилось. Поэтому вывод теперь пишется в реальный временный файл, а
    ожидание идёт через Popen.wait() — он ждёт только завершения САМОГО
    bash-процесса, а не закрытия дескрипторов его фоновых потомков."""
    if not state.ai_autonomy:
        return ("Инструмент run_shell_command отключён. Пользователь должен сначала включить "
                "«Автономность ИИ» в боковой панели сайта.")
    command = (command or "").strip()
    if not command:
        return "Ошибка: пустая команда."

    preview = f"$ {'sudo ' if use_sudo else ''}{command}"
    approved = request_confirmation("exec", command, preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ выполнение команды."

    tmp_out = tempfile.NamedTemporaryFile(prefix="ai_shell_", suffix=".log", delete=False)
    tmp_out.close()
    tmp_path = tmp_out.name
    wrapped_command = f"{command} > {shlex.quote(tmp_path)} 2>&1"

    proc = None
    timed_out = False
    try:
        if use_sudo:
            proc = subprocess.Popen(
                ["sudo", "-S", "bash", "-c", wrapped_command],
                stdin=subprocess.PIPE, text=True,
            )
            try:
                proc.communicate(
                    input=(state.sudo_password + "\n") if state.sudo_password else "\n",
                    timeout=SHELL_COMMAND_TIMEOUT,
                )
            except subprocess.TimeoutExpired:
                timed_out = True
        else:
            proc = subprocess.Popen(["bash", "-c", wrapped_command])
            try:
                proc.wait(timeout=SHELL_COMMAND_TIMEOUT)
            except subprocess.TimeoutExpired:
                timed_out = True
    except FileNotFoundError:
        Path(tmp_path).unlink(missing_ok=True)
        return "Ошибка: команда 'sudo' или 'bash' не найдена в системе."
    except Exception as e:
        Path(tmp_path).unlink(missing_ok=True)
        return f"Ошибка выполнения: {e}"

    out = ""
    try:
        out = Path(tmp_path).read_text(encoding="utf-8", errors="replace")
    except Exception:
        pass
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    if timed_out:
        # Убиваем только сам bash-обёртку (proc), а не фоновые процессы,
        # которые она успела запустить через '&' — они, скорее всего, уже
        # успешно работают (иначе и не возникло бы зависание на чтении).
        if proc is not None:
            proc.kill()
        out = out.strip()
        note = (
            f"Команда не завершилась за {SHELL_COMMAND_TIMEOUT} сек и была прервана. "
            f"Если это был запуск фонового процесса ('&' в конце) — скорее всего он УЖЕ "
            f"успешно запустился (частая причина: команда сама по себе не отдаёт "
            f"управление, например GUI-приложение без '&' или без редиректа своего "
            f"вывода). Проверь через 'ps aux | grep <имя>' вместо повторного запуска — "
            f"НЕ пытайся запускать другим способом, если процесс уже фактически работает."
        )
        return f"[таймаут]\n{note}" + (f"\n\nВывод до обрыва:\n{out}" if out else "")

    out = out.strip() or "(нет вывода)"
    if len(out) > MAX_SHELL_OUTPUT_CHARS:
        out = out[:MAX_SHELL_OUTPUT_CHARS] + "\n… (вывод обрезан)"
    status = "успешно, код 0" if proc.returncode == 0 else f"код возврата {proc.returncode}"
    return f"[{status}]\n{out}"


# ----------------------------- Веб-поиск и чтение страниц ---------------------

WEB_SEARCH_MAX_RESULTS = 8
WEB_FETCH_MAX_CHARS = 6000
_WEB_UA = ("Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
           "(KHTML, like Gecko) Chrome/124.0 Safari/537.36")


def tool_web_search(query: str, max_results: int = 5) -> str:
    """Ищет в интернете через HTML-версию DuckDuckGo (не требует API-ключа) и
    возвращает список результатов: заголовок, ссылка, краткое описание."""
    query = (query or "").strip()
    if not query:
        return "Ошибка: пустой поисковый запрос."
    try:
        max_results = int(max_results or 5)
    except (TypeError, ValueError):
        max_results = 5
    max_results = max(1, min(max_results, WEB_SEARCH_MAX_RESULTS))

    try:
        resp = requests.post(
            "https://html.duckduckgo.com/html/",
            data={"q": query},
            headers={"User-Agent": _WEB_UA},
            timeout=15,
        )
    except requests.RequestException as e:
        return f"Ошибка поиска: {e}"
    if resp.status_code != 200:
        return f"Ошибка поиска: сервер вернул код {resp.status_code}."

    page = resp.text
    link_pat = re.compile(r'class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', re.S)
    snippet_pat = re.compile(r'class="result__snippet"[^>]*>(.*?)</a>', re.S)
    links = link_pat.findall(page)
    snippets = snippet_pat.findall(page)

    def clean(s):
        return html_lib.unescape(re.sub(r"<[^>]+>", "", s)).strip()

    def real_url(u):
        # DuckDuckGo оборачивает внешние ссылки в /l/?uddg=<encoded_url>
        m = re.search(r"uddg=([^&]+)", u)
        if m:
            import urllib.parse
            return urllib.parse.unquote(m.group(1))
        return u

    results = []
    for i, (raw_url, title) in enumerate(links[:max_results]):
        snippet = clean(snippets[i]) if i < len(snippets) else ""
        results.append(f"{i + 1}. {clean(title)}\n   {real_url(raw_url)}\n   {snippet}")

    if not results:
        return f"По запросу «{query}» ничего не найдено."
    return f"Результаты поиска по запросу «{query}»:\n\n" + "\n\n".join(results)


def tool_fetch_url(url: str) -> str:
    """Открывает страницу по ссылке и возвращает её текстовое содержимое
    (теги убраны), чтобы модель могла прочитать статью/страницу целиком."""
    url = (url or "").strip()
    if not url:
        return "Ошибка: пустой URL."
    if not re.match(r"^https?://", url, re.IGNORECASE):
        url = "https://" + url
    try:
        resp = requests.get(url, headers={"User-Agent": _WEB_UA}, timeout=15)
    except requests.RequestException as e:
        return f"Ошибка загрузки страницы: {e}"
    if resp.status_code != 200:
        return f"Ошибка загрузки: сервер вернул код {resp.status_code}."

    ctype = resp.headers.get("Content-Type", "")
    if "text" not in ctype and "html" not in ctype and "json" not in ctype and ctype:
        return f"Страница не текстовая (Content-Type: {ctype}) — пропущено."

    text = resp.text
    text = re.sub(r"(?is)<(script|style|noscript|svg).*?</\1>", "", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = html_lib.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    text = text.strip()
    if not text:
        return f"[{url}]\n(страница пуста или содержимое не удалось извлечь)"
    if len(text) > WEB_FETCH_MAX_CHARS:
        text = text[:WEB_FETCH_MAX_CHARS] + "\n… (обрезано)"
    return f"[{url}]\n{text}"


def tool_wikipedia_search(query: str, lang: str = "ru", max_results: int = 3) -> str:
    """Ищет статьи в Википедии через официальное MediaWiki API (без ключа) и
    возвращает краткое содержание (extract) для каждой найденной статьи —
    удобный источник дополнительных фоновых знаний."""
    query = (query or "").strip()
    if not query:
        return "Ошибка: пустой запрос."
    lang = re.sub(r"[^a-z-]", "", (lang or "ru").lower()) or "ru"
    try:
        max_results = max(1, min(int(max_results or 3), 5))
    except (TypeError, ValueError):
        max_results = 3

    try:
        resp = requests.get(
            f"https://{lang}.wikipedia.org/w/api.php",
            params={
                "action": "query", "format": "json", "generator": "search",
                "gsrsearch": query, "gsrlimit": max_results,
                "prop": "extracts|info", "exintro": 1, "explaintext": 1,
                "exchars": 700, "inprop": "url",
            },
            headers={"User-Agent": _WEB_UA}, timeout=15,
        )
    except requests.RequestException as e:
        return f"Ошибка запроса к Википедии: {e}"
    if resp.status_code != 200:
        return f"Ошибка Википедии: сервер вернул код {resp.status_code}."

    try:
        data = resp.json()
    except ValueError:
        return "Ошибка: Википедия вернула некорректный ответ."

    pages = (data.get("query") or {}).get("pages") or {}
    if not pages:
        # Пробуем английскую Википедию, если в указанном языковом разделе ничего нет
        if lang != "en":
            return tool_wikipedia_search(query, lang="en", max_results=max_results)
        return f"В Википедии ничего не найдено по запросу «{query}»."

    # У MediaWiki порядок словаря не гарантирует порядок релевантности —
    # сортируем по 'index', который API проставляет по релевантности поиска.
    ordered = sorted(pages.values(), key=lambda p: p.get("index", 999))
    parts = []
    for i, page in enumerate(ordered[:max_results]):
        title = page.get("title", "")
        extract = (page.get("extract") or "").strip()
        url = page.get("fullurl", "")
        parts.append(f"{i + 1}. {title}\n   {url}\n   {extract}")
    return f"Википедия ({lang}) по запросу «{query}»:\n\n" + "\n\n".join(parts)


def _google_html_scrape(query: str, max_results: int):
    """Парсит обычную HTML-страницу выдачи google.com/search — без ключей и
    OAuth, но и без официальных гарантий: Google активно борется со
    скрейпингом (капчи, изменение разметки, временные блокировки по IP),
    поэтому это best-effort попытка, а не надёжный API. При неудаче
    вызывающий код откатывается на DuckDuckGo."""
    resp = requests.get(
        "https://www.google.com/search",
        params={"q": query, "num": max_results, "hl": "ru"},
        headers={
            "User-Agent": _WEB_UA,
            "Accept-Language": "ru-RU,ru;q=0.9,en;q=0.8",
        },
        timeout=15,
    )
    if resp.status_code != 200:
        return None
    page = resp.text
    if "unusual traffic" in page.lower() or "recaptcha" in page.lower() or "/sorry/" in page:
        return None  # Google показал капчу/блокировку — сдаёмся, пусть решает fallback

    # Разметка Google часто меняется и не документирована официально — этот
    # паттерн ловит типичный вид блока результата: ссылка с <h3>-заголовком
    # внутри, затем где-то рядом текстовый сниппет.
    block_pat = re.compile(r'<a href="(https?://(?!www\.google\.)[^"]+)"[^>]*>\s*<h3[^>]*>(.*?)</h3>', re.S)
    matches = block_pat.findall(page)
    if not matches:
        return None

    def clean(s):
        return html_lib.unescape(re.sub(r"<[^>]+>", "", s)).strip()

    out = []
    seen_urls = set()
    for url, title in matches:
        if url in seen_urls:
            continue
        seen_urls.add(url)
        out.append((clean(title), url, ""))
        if len(out) >= max_results:
            break
    return out or None


def tool_google_search(query: str, max_results: int = 5) -> str:
    """Поиск в Google. По умолчанию (без ключей) парсит обычную HTML-выдачу
    google.com/search напрямую — быстрый вариант без регистрации, но
    неофициальный и не гарантированно стабильный (Google может показать
    капчу). Если задан google_api_key/google_cx в настройках — используется
    официальный Custom Search JSON API вместо парсинга (надёжнее, но с
    лимитом 100 бесплатных запросов/день). Если оба способа не сработали —
    тихий откат на DuckDuckGo, чтобы инструмент всё равно вернул результат."""
    query = (query or "").strip()
    if not query:
        return "Ошибка: пустой поисковый запрос."
    try:
        max_results = max(1, min(int(max_results or 5), 10))
    except (TypeError, ValueError):
        max_results = 5
    api_err = ""

    if state.google_api_key and state.google_cx:
        try:
            resp = requests.get(
                "https://www.googleapis.com/customsearch/v1",
                params={"key": state.google_api_key, "cx": state.google_cx,
                        "q": query, "num": max_results},
                timeout=15,
            )
            data = resp.json()
        except requests.RequestException as e:
            return f"Ошибка запроса к Google API: {e}"
        except ValueError:
            data = None
        if data is not None and resp.status_code == 200:
            items = data.get("items") or []
            if not items:
                return f"Google ничего не нашёл по запросу «{query}»."
            parts = []
            for i, item in enumerate(items[:max_results]):
                title = item.get("title", "")
                link = item.get("link", "")
                snippet = (item.get("snippet") or "").replace("\n", " ").strip()
                parts.append(f"{i + 1}. {title}\n   {link}\n   {snippet}")
            return f"Google (официальный API) по запросу «{query}»:\n\n" + "\n\n".join(parts)
        # Ключ настроен, но запрос не удался — не молчим, а сообщаем и всё
        # равно пробуем HTML-парсинг ниже, чтобы результат хоть как-то пришёл.
        api_err = (data.get("error") or {}).get("message", "") if isinstance(data, dict) else ""

    try:
        scraped = _google_html_scrape(query, max_results)
    except requests.RequestException:
        scraped = None
    if scraped:
        parts = [f"{i + 1}. {title}\n   {url}" for i, (title, url, snippet) in enumerate(scraped)]
        return f"Google (прямой поиск) по запросу «{query}»:\n\n" + "\n\n".join(parts)

    # И официальный API (если был настроен), и прямой парсинг не сработали —
    # используем DuckDuckGo, чтобы инструмент всё равно вернул хоть что-то.
    fallback = tool_web_search(query, max_results)
    note = f" ({api_err})" if api_err else ""
    return (f"[Google недоступен (капча/блокировка, или ошибка API{note}, или не настроен) — "
            f"использован DuckDuckGo как запасной поисковик]\n\n" + fallback)


def tool_github_search(query: str, max_results: int = 5) -> str:
    """Поиск репозиториев на GitHub через публичный REST API (без токена,
    с ограничением по частоте запросов). Полезно, чтобы найти проект/код для
    просмотра (fetch_url) или скачивания (download_file)."""
    query = (query or "").strip()
    if not query:
        return "Ошибка: пустой поисковый запрос."
    try:
        max_results = max(1, min(int(max_results or 5), 10))
    except (TypeError, ValueError):
        max_results = 5

    try:
        resp = requests.get(
            "https://api.github.com/search/repositories",
            params={"q": query, "sort": "stars", "order": "desc", "per_page": max_results},
            headers={"User-Agent": _WEB_UA, "Accept": "application/vnd.github+json"},
            timeout=15,
        )
    except requests.RequestException as e:
        return f"Ошибка запроса к GitHub: {e}"
    try:
        data = resp.json()
    except ValueError:
        return f"Ошибка GitHub: сервер вернул код {resp.status_code}, некорректный ответ."
    if resp.status_code != 200:
        msg = data.get("message", "") if isinstance(data, dict) else ""
        return f"Ошибка GitHub API ({resp.status_code}): {msg or 'превышен лимит запросов, попробуйте позже'}."

    items = data.get("items") or []
    if not items:
        return f"На GitHub ничего не найдено по запросу «{query}»."
    parts = []
    for i, item in enumerate(items[:max_results]):
        name = item.get("full_name", "")
        url = item.get("html_url", "")
        stars = item.get("stargazers_count", 0)
        lang = item.get("language") or "?"
        desc = (item.get("description") or "").strip()
        clone_url = item.get("clone_url", "")
        parts.append(f"{i + 1}. {name} (★{stars}, {lang})\n   {url}\n   clone: {clone_url}\n   {desc}")
    return f"GitHub по запросу «{query}»:\n\n" + "\n\n".join(parts)


_TRANSLIT_RU_EN_HINTS = {
    # Небольшой словарь самых частых тем — если запрос на русском и Openverse
    # ничего не находит, пробуем эти же слова по-английски (метаданные на
    # большинстве стоков и в Wikimedia Commons — в основном на английском).
    "планета": "planet", "земля": "earth", "космос": "space", "звезда": "star",
    "звёзды": "stars", "детская комната": "kids room", "детская": "kids room",
    "кухня": "kitchen", "интерьер": "interior design", "дом": "house",
    "природа": "nature", "горы": "mountains", "море": "sea", "лес": "forest",
    "город": "city", "цветы": "flowers", "животные": "animals", "кот": "cat",
    "собака": "dog", "машина": "car", "еда": "food", "закат": "sunset",
}


def _guess_english_query(query: str) -> str:
    q = query.lower()
    for ru, en in _TRANSLIT_RU_EN_HINTS.items():
        if ru in q:
            q = q.replace(ru, en)
    return q


def _openverse_search(query: str, max_results: int):
    resp = requests.get(
        "https://api.openverse.org/v1/images/",
        params={"q": query, "page_size": max_results},
        headers={"User-Agent": _WEB_UA}, timeout=15,
    )
    if resp.status_code != 200:
        return []
    try:
        data = resp.json()
    except ValueError:
        return []
    out = []
    for r in (data.get("results") or [])[:max_results]:
        # ВАЖНО: используем "thumbnail" (проксируется через собственный CDN
        # Openverse), а не "url" (прямая ссылка на оригинал у источника,
        # например staticflickr.com). Многие хостинги блокируют хотлинкинг
        # и вместо картинки отдают HTTP 200 с "заглушкой-предупреждением"
        # (характерные диагональные полосы) — <img> её показывает как будто
        # это реальное фото, так как ошибки загрузки формально нет.
        url = r.get("thumbnail") or r.get("url") or ""
        title = (r.get("title") or query).strip().replace("[", "").replace("]", "")
        if url:
            out.append((title, url))
    return out


def _wikimedia_commons_search(query: str, max_results: int):
    resp = requests.get(
        "https://commons.wikimedia.org/w/api.php",
        params={
            "action": "query", "format": "json", "generator": "search",
            "gsrsearch": f"filetype:bitmap {query}", "gsrnamespace": 6, "gsrlimit": max_results,
            "prop": "imageinfo", "iiprop": "url", "iiurlwidth": 800,
        },
        headers={"User-Agent": _WEB_UA}, timeout=15,
    )
    if resp.status_code != 200:
        return []
    try:
        data = resp.json()
    except ValueError:
        return []
    pages = (data.get("query") or {}).get("pages") or {}
    out = []
    for page in pages.values():
        infos = page.get("imageinfo") or []
        if not infos:
            continue
        info = infos[0]
        url = info.get("thumburl") or info.get("url") or ""
        title = (page.get("title") or query).replace("File:", "").rsplit(".", 1)[0]
        if url:
            out.append((title, url))
        if len(out) >= max_results:
            break
    return out


def tool_image_search(query: str, max_results: int = 6) -> str:
    """Ищет изображения через Openverse и (если пусто) Wikimedia Commons —
    без ключа. Если русский запрос не даёт результатов, автоматически
    пробует упрощённый английский вариант. Возвращает markdown ![]() —
    модель должна скопировать эти строки в ответ пользователю без изменений."""
    query = (query or "").strip()
    if not query:
        return "Ошибка: пустой запрос."
    try:
        max_results = max(1, min(int(max_results or 6), 12))
    except (TypeError, ValueError):
        max_results = 6

    tried_queries = [query]
    en_guess = _guess_english_query(query)
    if en_guess != query.lower():
        tried_queries.append(en_guess)

    found = []
    try:
        for q in tried_queries:
            found = _openverse_search(q, max_results)
            if found:
                break
        if not found:
            for q in tried_queries:
                found = _wikimedia_commons_search(q, max_results)
                if found:
                    break
    except requests.RequestException as e:
        return f"Ошибка поиска изображений: {e}"

    if not found:
        return (
            f"Изображений по запросу «{query}» не найдено ни в Openverse, ни в Wikimedia "
            f"Commons. НЕ используй google_search как замену — он даёт ссылки на страницы "
            f"сайтов, а не на сами картинки, и в чате не отобразится. Вместо этого повтори "
            f"image_search с более простым запросом из 1-2 английских слов (например "
            f"'planet earth' вместо длинной фразы на русском)."
        )

    lines = [
        "Найденные изображения — ОБЯЗАТЕЛЬНО вставь эти markdown-строки "
        "![...](...) в свой ответ пользователю дословно, не пересказывая словами, "
        "чтобы картинки реально показались в чате:"
    ]
    for title, url in found[:max_results]:
        lines.append(f"![{title}]({url})")
    return "\n".join(lines)


_YT_ID_RE = re.compile(
    r"(?:youtube\.com/(?:watch\?v=|shorts/|embed/)|youtu\.be/)([a-zA-Z0-9_-]{6,15})"
)


def tool_video_search(query: str, max_results: int = 4) -> str:
    """Ищет видео на YouTube (через обычный веб-поиск, ограниченный сайтом
    youtube.com). Возвращает ссылки — модель должна вставить их в ответ как
    есть (можно просто голым текстом), тогда в чате появится встроенный
    видеоплеер."""
    query = (query or "").strip()
    if not query:
        return "Ошибка: пустой запрос."
    try:
        max_results = max(1, min(int(max_results or 4), 6))
    except (TypeError, ValueError):
        max_results = 4

    try:
        resp = requests.post(
            "https://html.duckduckgo.com/html/",
            data={"q": f"{query} site:youtube.com"},
            headers={"User-Agent": _WEB_UA},
            timeout=15,
        )
    except requests.RequestException as e:
        return f"Ошибка поиска видео: {e}"
    if resp.status_code != 200:
        return f"Ошибка поиска видео: сервер вернул код {resp.status_code}."

    page = resp.text
    link_pat = re.compile(r'class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>', re.S)

    def clean(s):
        return html_lib.unescape(re.sub(r"<[^>]+>", "", s)).strip()

    def real_url(u):
        m = re.search(r"uddg=([^&]+)", u)
        if m:
            import urllib.parse
            return urllib.parse.unquote(m.group(1))
        return u

    found = []
    seen_ids = set()
    for raw_url, title in link_pat.findall(page):
        url = real_url(raw_url)
        m = _YT_ID_RE.search(url)
        if not m:
            continue
        vid = m.group(1)
        if vid in seen_ids:
            continue
        seen_ids.add(vid)
        found.append((clean(title), f"https://www.youtube.com/watch?v={vid}"))
        if len(found) >= max_results:
            break

    if not found:
        return f"Видео на YouTube по запросу «{query}» не найдено."
    lines = [
        "Найденные видео на YouTube — вставь эти ссылки в ответ пользователю "
        "как обычный текст (без изменений), встроенный плеер появится сам:"
    ]
    for title, url in found:
        lines.append(f"{title}\n{url}")
    return "\n\n".join(lines)


MAX_DOWNLOAD_BYTES = 200 * 1024 * 1024  # 200 МБ


def tool_download_file(url: str, path: str, emit, force_auto: bool = False) -> str:
    """Скачивает файл по ссылке (например, найденной через google_search) и
    сохраняет его в одну из подключённых папок. Требует подтверждения
    пользователя, как write_file — реально пишет на диск."""
    url = (url or "").strip()
    if not url:
        return "Ошибка: пустой URL."
    if not re.match(r"^https?://", url, re.IGNORECASE):
        url = "https://" + url
    if not path:
        return "Ошибка: не указан путь сохранения (path)."

    try:
        target = safe_path(path)
    except SandboxError as e:
        return f"Ошибка доступа: {e}"

    preview = f"Скачать: {url}\n→ Сохранить как: {path}"
    approved = request_confirmation("download", path, preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ скачивание файла."

    try:
        with requests.get(url, headers={"User-Agent": _WEB_UA}, stream=True, timeout=30) as resp:
            if resp.status_code != 200:
                return f"Ошибка скачивания: сервер вернул код {resp.status_code}."
            total = 0
            target.parent.mkdir(parents=True, exist_ok=True)
            tmp_path = target.with_suffix(target.suffix + ".part")
            with open(tmp_path, "wb") as f:
                for chunk in resp.iter_content(chunk_size=1024 * 256):
                    if not chunk:
                        continue
                    total += len(chunk)
                    if total > MAX_DOWNLOAD_BYTES:
                        f.close()
                        tmp_path.unlink(missing_ok=True)
                        return f"Ошибка: файл больше лимита {MAX_DOWNLOAD_BYTES // 1_000_000} МБ, скачивание прервано."
                    f.write(chunk)
            tmp_path.replace(target)
    except requests.RequestException as e:
        return f"Ошибка скачивания: {e}"
    except OSError as e:
        return f"Ошибка записи файла: {e}"

    return f"Файл скачан: '{path}' ({total} байт) ← {url}"


# ============================================================================
# ДОПОЛНИТЕЛЬНЫЕ ИНСТРУМЕНТЫ ИИ (документы, OCR, архивы, память, система,
# git, уведомления, email/telegram, python-repl, линтеры, погода, курсы, RSS)
# ============================================================================

def tool_read_document(path: str) -> str:
    """Читает текст из PDF, DOCX или XLSX (обычный read_file понимает только
    текстовые файлы). Определяет формат по расширению."""
    target = safe_path(path)
    if not target.exists():
        return f"Ошибка: файл '{path}' не найден."
    ext = target.suffix.lower()

    if ext == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            return "Ошибка: не установлена библиотека pdfplumber. Установите: pip install pdfplumber"
        try:
            parts = []
            with pdfplumber.open(target) as pdf:
                for i, page in enumerate(pdf.pages[:200]):
                    text = page.extract_text() or ""
                    if text.strip():
                        parts.append(f"--- Страница {i + 1} ---\n{text.strip()}")
            content = "\n\n".join(parts) or "(текст не найден — возможно, это скан; попробуйте ocr_image)"
        except Exception as e:
            return f"Ошибка чтения PDF: {e}"

    elif ext == ".docx":
        try:
            import docx
        except ImportError:
            return "Ошибка: не установлена библиотека python-docx. Установите: pip install python-docx"
        try:
            doc = docx.Document(str(target))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_text = []
            for t in doc.tables:
                for row in t.rows:
                    tables_text.append(" | ".join(c.text.strip() for c in row.cells))
            content = "\n".join(paras)
            if tables_text:
                content += "\n\n[Таблицы]\n" + "\n".join(tables_text)
        except Exception as e:
            return f"Ошибка чтения DOCX: {e}"

    elif ext in (".xlsx", ".xlsm"):
        try:
            import openpyxl
        except ImportError:
            return "Ошибка: не установлена библиотека openpyxl. Установите: pip install openpyxl"
        try:
            wb = openpyxl.load_workbook(str(target), data_only=True, read_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"--- Лист: {ws.title} ---")
                for row in ws.iter_rows(max_row=500, values_only=True):
                    if any(c is not None for c in row):
                        parts.append(" | ".join("" if c is None else str(c) for c in row))
            content = "\n".join(parts)
        except Exception as e:
            return f"Ошибка чтения XLSX: {e}"

    else:
        return (f"Формат '{ext}' не поддерживается read_document — используйте read_file "
                f"для обычных текстовых файлов.")

    if len(content) > 20000:
        content = content[:20000] + "\n… (обрезано)"
    return content or "(пусто)"


def tool_ocr_image(path: str, lang: str = "rus+eng") -> str:
    """Распознаёт текст на изображении (скриншот, скан, фото документа) через
    Tesseract OCR. Требует системный пакет tesseract-ocr и pip-пакет pytesseract."""
    target = safe_path(path)
    if not target.exists():
        return f"Ошибка: файл '{path}' не найден."
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ("Ошибка: не установлены pytesseract/Pillow. Установите: "
                "pip install pytesseract Pillow")
    try:
        img = Image.open(target)
        text = pytesseract.image_to_string(img, lang=lang or "rus+eng")
    except pytesseract.TesseractNotFoundError:
        return ("Ошибка: системная программа tesseract-ocr не найдена. Установите: "
                "sudo apt install tesseract-ocr tesseract-ocr-rus")
    except Exception as e:
        return f"Ошибка OCR: {e}"
    text = text.strip()
    if len(text) > 10000:
        text = text[:10000] + "\n… (обрезано)"
    return text or "(текст на изображении не распознан)"


def tool_archive_extract(path: str, dest: str, emit, force_auto: bool = False) -> str:
    """Распаковывает архив (zip/tar/tar.gz/tar.bz2/tar.xz) в указанную папку."""
    src = safe_path(path)
    if not src.exists():
        return f"Ошибка: архив '{path}' не найден."
    dest_dir = safe_path(dest)
    preview = f"Распаковать: {path}\n→ в папку: {dest}"
    approved = request_confirmation("write", dest, preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ распаковку архива."
    try:
        dest_dir.mkdir(parents=True, exist_ok=True)
        shutil.unpack_archive(str(src), str(dest_dir))
    except Exception as e:
        return f"Ошибка распаковки: {e}"
    return f"Архив '{path}' распакован в '{dest}'."


def tool_archive_create(source: str, dest: str, emit, force_auto: bool = False) -> str:
    """Упаковывает папку/файл в архив. Формат определяется по расширению dest
    (.zip, .tar.gz, .tar.bz2, .tar.xz)."""
    src = safe_path(source)
    if not src.exists():
        return f"Ошибка: '{source}' не найден."
    dest_path = safe_path(dest)
    name = dest_path.name.lower()
    if name.endswith(".zip"):
        fmt, base = "zip", str(dest_path)[:-4]
    elif name.endswith(".tar.gz") or name.endswith(".tgz"):
        fmt, base = "gztar", re.sub(r"\.(tar\.gz|tgz)$", "", str(dest_path))
    elif name.endswith(".tar.bz2"):
        fmt, base = "bztar", str(dest_path)[:-8]
    elif name.endswith(".tar.xz"):
        fmt, base = "xztar", str(dest_path)[:-7]
    else:
        return "Ошибка: путь назначения должен заканчиваться на .zip, .tar.gz, .tar.bz2 или .tar.xz"

    preview = f"Запаковать: {source}\n→ в архив: {dest}"
    approved = request_confirmation("write", dest, preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ создание архива."
    try:
        Path(base).parent.mkdir(parents=True, exist_ok=True)
        root_dir = src.parent if src.is_file() else src
        base_dir = src.name if src.is_file() else None
        result_path = shutil.make_archive(base, fmt, root_dir=str(root_dir), base_dir=base_dir)
    except Exception as e:
        return f"Ошибка создания архива: {e}"
    return f"Архив создан: {result_path}"


def tool_diff_files(path_a: str, path_b: str) -> str:
    """Построчное сравнение (unified diff) двух текстовых файлов."""
    import difflib
    a = safe_path(path_a)
    b = safe_path(path_b)
    if not a.exists():
        return f"Ошибка: файл '{path_a}' не найден."
    if not b.exists():
        return f"Ошибка: файл '{path_b}' не найден."
    try:
        lines_a = a.read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
        lines_b = b.read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
    except Exception as e:
        return f"Ошибка чтения файлов: {e}"
    diff = list(difflib.unified_diff(lines_a, lines_b, fromfile=path_a, tofile=path_b, lineterm=""))
    if not diff:
        return "Файлы идентичны."
    out = "\n".join(diff)
    if len(out) > 12000:
        out = out[:12000] + "\n… (обрезано)"
    return out


def tool_remember(key: str, text: str) -> str:
    """Сохраняет заметку в долгосрочную память ИИ (переживает разные чаты)."""
    key = (key or "").strip()
    if not key:
        return "Ошибка: не указан ключ заметки."
    state.memory_notes[key] = (text or "").strip()
    state.save()
    return f"Запомнено под ключом «{key}»."


def tool_recall(key: str = "") -> str:
    """Достаёт заметку из долгосрочной памяти по ключу, либо список всех
    заметок, если ключ не указан."""
    key = (key or "").strip()
    if key:
        if key in state.memory_notes:
            return f"{key}: {state.memory_notes[key]}"
        return f"Заметки с ключом «{key}» нет."
    if not state.memory_notes:
        return "Память пуста — заметок ещё нет."
    lines = [f"- {k}: {v[:200]}{'...' if len(v) > 200 else ''}" for k, v in state.memory_notes.items()]
    return "Все заметки:\n" + "\n".join(lines)


# ----------------------------------------------------------------------------
# Архив успешных решений (save_solution / recall_solutions).
#
# Задача: ИИ должен применять уже проверенные рабочие решения вместо того,
# чтобы каждый раз изобретать заново — и не "забывать" про архив через
# несколько сообщений. Полагаться на то, что модель САМА не забудет вызвать
# run_shell_command → ls/cat каждый раз, ненадёжно (особенно для маленьких
# локальных моделей — внимание к системным инструкциям падает по мере роста
# диалога). Поэтому здесь два независимых механизма:
#   1) save_solution / recall_solutions — инструменты, которые модель может
#      вызывать сама.
#   2) solutions_auto_digest() — вызывается СЕРВЕРОМ автоматически на КАЖДОМ
#      сообщении (см. build_system_message) и подмешивает наиболее похожие
#      прошлые решения прямо в system prompt, без участия модели вообще.
# ----------------------------------------------------------------------------

def _solutions_dir() -> Path:
    raw = (state.solutions_archive_path or "").strip()
    d = Path(raw).expanduser() if raw else (Path.home() / "AI_memory" / "solutions")
    d.mkdir(parents=True, exist_ok=True)
    return d


def _load_solution_files() -> list:
    d = _solutions_dir()
    items = []
    for f in sorted(d.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)[:300]:
        try:
            data = json.loads(f.read_text(encoding="utf-8", errors="replace"))
        except Exception:
            continue
        if isinstance(data, dict) and data.get("task") and data.get("solution"):
            items.append({"file": f.name, "task": str(data["task"]), "solution": str(data["solution"])})
    return items


_STOPWORDS_RU_EN = {
    "как", "для", "это", "что", "или", "при", "если", "нужно", "надо", "чтобы",
    "the", "and", "for", "with", "how", "what", "this", "that", "you", "your",
}


def _score_overlap(query: str, task: str) -> int:
    q_words = {w for w in re.findall(r"[a-zA-Zа-яёА-ЯЁ0-9_./-]{3,}", query.lower()) if w not in _STOPWORDS_RU_EN}
    t_words = {w for w in re.findall(r"[a-zA-Zа-яёА-ЯЁ0-9_./-]{3,}", task.lower()) if w not in _STOPWORDS_RU_EN}
    return len(q_words & t_words)


def tool_save_solution(task: str, solution: str) -> str:
    """Сохраняет успешно решённую техническую задачу в архив (JSON-файл с
    меткой времени) — чтобы в будущем не решать её заново, а взять готовый
    рабочий вариант. Вызывай сразу после того, как задача пользователя решена
    и результат подтверждён рабочим (например, команда/скрипт отработали без
    ошибок)."""
    task = (task or "").strip()
    solution = (solution or "").strip()
    if not task or not solution:
        return "Ошибка: нужно указать и task, и solution."
    d = _solutions_dir()
    fname = f"task_{time.strftime('%Y%m%d_%H%M%S')}.json"
    payload = {"task": task, "solution": solution, "created_at": time.strftime("%Y-%m-%d %H:%M:%S")}
    try:
        (d / fname).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        return f"Ошибка сохранения в архив: {e}"
    return f"Решение сохранено в архив: {d / fname}"


def tool_recall_solutions(query: str, max_results: int = 5) -> str:
    """Ищет в архиве успешных решений записи, похожие на запрос (по общим
    словам в описании задачи), и возвращает их готовые решения."""
    query = (query or "").strip()
    if not query:
        return "Ошибка: пустой запрос."
    try:
        max_results = max(1, min(int(max_results or 5), 10))
    except (TypeError, ValueError):
        max_results = 5
    items = _load_solution_files()
    if not items:
        return f"Архив решений пуст ({_solutions_dir()}) — ещё нечего вспоминать."
    scored = [(it, _score_overlap(query, it["task"])) for it in items]
    scored = [x for x in scored if x[1] > 0]
    scored.sort(key=lambda x: x[1], reverse=True)
    top = scored[:max_results]
    if not top:
        return f"В архиве ({len(items)} записей) не нашлось похожих задач по запросу «{query}»."
    lines = [f"Похожие задачи из архива по запросу «{query}»:"]
    for it, score in top:
        lines.append(f"- Задача: {it['task']}\n  Решение: {it['solution']}")
    return "\n".join(lines)


def solutions_auto_digest(user_text: str, max_items: int = 2) -> str:
    """Вызывается сервером автоматически (не моделью!) при сборке system
    prompt на КАЖДОМ сообщении. Если текущее сообщение пользователя похоже на
    задачу из архива — подмешивает готовое решение прямо в system prompt, без
    необходимости, чтобы модель сама вспомнила про run_shell_command/ls/cat."""
    user_text = (user_text or "").strip()
    if not user_text or len(user_text) < 6:
        return ""
    try:
        items = _load_solution_files()
    except Exception:
        return ""
    if not items:
        return ""
    # Если пользователь поставил архив на 1-е место приоритета — снижаем
    # порог совпадения (достаточно 1 общего слова вместо 2) и показываем
    # больше кандидатов, т.к. в этом случае пользователь явно хочет, чтобы
    # архив проверялся агрессивнее, а не только при точном совпадении фраз.
    archive_is_top = bool(state.knowledge_priority) and state.knowledge_priority[0] == "archive"
    min_score = 1 if archive_is_top else 2
    if archive_is_top:
        max_items = max(max_items, 4)
    scored = [(it, _score_overlap(user_text, it["task"])) for it in items]
    scored = [x for x in scored if x[1] >= min_score]
    if not scored:
        if archive_is_top:
            # Ни одного совпадения по ключевым словам, но архив стоит первым
            # в приоритете — не молчим полностью: показываем список ЗАГОЛОВКОВ
            # (без решений) как явный намёк модели, что архив непустой и
            # ОБЯЗАТЕЛЬНО стоит проверить его вручную через recall_solutions,
            # прежде чем изобретать решение с нуля.
            titles = [it["task"] for it in items[:8]]
            return (
                "Архив успешных решений НЕ пуст (в нём есть записи, но ни одна не совпала "
                "по ключевым словам с текущим запросом автоматически) — заголовки сохранённых "
                "задач: " + "; ".join(titles) + ". Раз архив в приоритете (1-е место), "
                "ОБЯЗАТЕЛЬНО вызови recall_solutions с сутью текущего запроса вручную, прежде "
                "чем предлагать решение самостоятельно — возможно, там есть применимое решение, "
                "которое просто не поймал автоматический подбор по словам."
            )
        return ""
    scored.sort(key=lambda x: x[1], reverse=True)
    top = scored[:max_items]
    lines = [
        "В архиве успешных решений (АРХИВ_РЕШЕНИЙ) найдены задачи, похожие на текущий запрос "
        "пользователя. Это уже проверенные рабочие решения — ИСПОЛЬЗУЙ их как основу вместо "
        "того, чтобы придумывать решение с нуля, если они применимы к текущей задаче:"
    ]
    for it, _score in top:
        lines.append(f"— Прошлая задача: {it['task']}\n  Рабочее решение: {it['solution']}")
    return "\n".join(lines)


def tool_process_list(filter_text: str = "") -> str:
    """Список запущенных процессов (имя, PID, CPU%, RAM%), опционально
    отфильтрованный по подстроке в имени."""
    try:
        import psutil
    except ImportError:
        return "Ошибка: не установлен psutil."
    filter_text = (filter_text or "").strip().lower()
    rows = []
    for p in psutil.process_iter(["pid", "name", "cpu_percent", "memory_percent"]):
        try:
            info = p.info
            name = info.get("name") or ""
            if filter_text and filter_text not in name.lower():
                continue
            rows.append((info.get("pid"), name, info.get("cpu_percent") or 0, info.get("memory_percent") or 0))
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    rows.sort(key=lambda r: r[2], reverse=True)
    rows = rows[:40]
    if not rows:
        return "Процессы не найдены."
    lines = [f"PID {pid:<7} {name[:30]:<30} CPU {cpu:5.1f}%  RAM {mem:5.1f}%" for pid, name, cpu, mem in rows]
    return "\n".join(lines)


def tool_check_process_running(hints: str) -> str:
    """Проверяет, УЖЕ ли запущено приложение — по нескольким ключевым словам
    сразу, а не по одному точному имени. Простой grep/pgrep по названию
    программы часто не срабатывает для Electron/AppImage-приложений: их
    РЕАЛЬНЫЙ процесс может называться 'electron', временным путём распаковки
    AppImage и т.п., никак не совпадая с именем программы. Поэтому здесь
    проверяются сразу: полная командная строка процессов (ps, не только имя)
    И заголовки открытых GUI-окон (wmctrl, если установлен) — заголовок окна
    почти всегда содержит настоящее имя программы, даже если её процесс
    называется иначе. Используй этот инструмент ПЕРЕД повторным запуском
    любого GUI-приложения, чтобы не плодить дубликаты процессов."""
    hints_list = [h.strip() for h in re.split(r"[,;]+", hints or "") if h.strip()]
    if not hints_list:
        return ("Ошибка: укажите хотя бы одно ключевое слово через запятую — и название "
                "программы, и общие варианты, например: 'lmstudio, LM Studio, electron'.")

    found_procs = []
    try:
        proc = subprocess.run(["ps", "-eo", "pid,comm,args", "--no-headers"],
                               capture_output=True, text=True, timeout=10)
        for line in proc.stdout.splitlines():
            low = line.lower()
            if any(h.lower() in low for h in hints_list):
                found_procs.append(line.strip()[:200])
    except Exception:
        pass

    found_windows = []
    try:
        wproc = subprocess.run(["wmctrl", "-l"], capture_output=True, text=True, timeout=5)
        if wproc.returncode == 0:
            for line in wproc.stdout.splitlines():
                low = line.lower()
                if any(h.lower() in low for h in hints_list):
                    found_windows.append(line.strip())
    except FileNotFoundError:
        pass
    except Exception:
        pass

    if not found_procs and not found_windows:
        return (
            f"Совпадений НЕ найдено ни среди процессов, ни среди открытых окон по словам "
            f"{hints_list} (проверка через 'ps' и 'wmctrl -l'). Приложение, скорее всего, "
            f"ещё не запущено — можно запускать."
        )

    lines = [
        "НАЙДЕНЫ совпадения — приложение, скорее всего, УЖЕ запущено. "
        "НЕ запускай его ещё раз, это создаст дублирующийся процесс:"
    ]
    if found_procs:
        lines.append("Процессы (PID, имя, командная строка):\n" + "\n".join(found_procs[:10]))
    if found_windows:
        lines.append("Открытые окна:\n" + "\n".join(found_windows[:10]))
    if not found_windows:
        lines.append(
            "(wmctrl не нашёл окно с таким названием — либо wmctrl не установлен "
            "[sudo apt install wmctrl], либо приложение ещё загружается)"
        )
    return "\n\n".join(lines)


GIT_READ_ACTIONS = {"status", "diff", "log"}
GIT_WRITE_ACTIONS = {"add", "commit", "push", "pull"}


def tool_git(action: str, repo_path: str, args: str, emit, force_auto: bool = False) -> str:
    """Обёртка над git: status/diff/log — без подтверждения (только чтение);
    add/commit/push/pull — требуют подтверждения пользователя."""
    action = (action or "").strip().lower()
    if action not in GIT_READ_ACTIONS | GIT_WRITE_ACTIONS:
        return f"Ошибка: неизвестное git-действие '{action}'. Доступны: status, diff, log, add, commit, push, pull."
    repo_dir = safe_path(repo_path or ".")
    if not repo_dir.exists():
        return f"Ошибка: путь репозитория '{repo_path}' не найден."

    extra = shlex.split(args) if args else []
    cmd = ["git", "-C", str(repo_dir), action] + extra

    if action in GIT_WRITE_ACTIONS:
        preview = f"$ git {action} {' '.join(extra)}\n(в репозитории: {repo_path})"
        approved = request_confirmation("exec", repo_path, preview, emit, force_auto=force_auto)
        if not approved:
            return "Пользователь ОТКЛОНИЛ git-операцию."

    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
    except FileNotFoundError:
        return "Ошибка: git не установлен в системе."
    except subprocess.TimeoutExpired:
        return "Ошибка: git-команда выполнялась слишком долго."
    out = (proc.stdout or "") + (("\n" + proc.stderr) if proc.stderr else "")
    out = out.strip() or "(нет вывода)"
    if len(out) > 8000:
        out = out[:8000] + "\n… (обрезано)"
    status = "успешно" if proc.returncode == 0 else f"код возврата {proc.returncode}"
    return f"[{status}]\n{out}"


SYSTEMD_ACTIONS = {"status", "start", "stop", "restart", "enable", "disable"}


def tool_systemd(action: str, service: str, emit, use_sudo: bool = True, force_auto: bool = False) -> str:
    """Управление системными службами (systemctl). Требует включённой
    'Автономности ИИ' (затрагивает всю систему) и подтверждения пользователя."""
    if not state.ai_autonomy:
        return ("Инструмент systemd_control отключён. Включите «Автономность ИИ» "
                "в боковой панели сайта.")
    action = (action or "").strip().lower()
    if action not in SYSTEMD_ACTIONS:
        return f"Ошибка: неизвестное действие '{action}'. Доступны: {', '.join(sorted(SYSTEMD_ACTIONS))}."
    service = (service or "").strip()
    if not service:
        return "Ошибка: не указано имя службы."

    preview = f"systemctl {action} {service}" + (" (sudo)" if use_sudo else "")
    approved = request_confirmation("exec", service, preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ операцию со службой."

    base_cmd = ["systemctl", action, service]
    try:
        if use_sudo:
            proc = subprocess.run(
                ["sudo", "-S"] + base_cmd,
                input=(state.sudo_password + "\n") if state.sudo_password else "\n",
                capture_output=True, text=True, timeout=30,
            )
        else:
            proc = subprocess.run(base_cmd, capture_output=True, text=True, timeout=30)
    except FileNotFoundError:
        return "Ошибка: systemctl не найден (не systemd-система?)."
    except subprocess.TimeoutExpired:
        return "Ошибка: команда выполнялась слишком долго."
    out = (proc.stdout or "") + (("\n" + proc.stderr) if proc.stderr else "")
    out = out.strip() or "(нет вывода)"
    status = "успешно" if proc.returncode == 0 else f"код возврата {proc.returncode}"
    return f"[{status}]\n{out}"


def tool_notify(message: str, title: str = "") -> str:
    """Отправляет push-уведомление через ntfy.sh (нужен ntfy_topic в настройках).
    Установите приложение ntfy или откройте https://ntfy.sh/<topic> в браузере
    на телефоне, чтобы получать уведомления."""
    if not state.ntfy_topic:
        return "Ошибка: не задан ntfy_topic в настройках — некуда отправлять уведомление."
    message = (message or "").strip()
    if not message:
        return "Ошибка: пустое сообщение."
    try:
        resp = requests.post(
            f"https://ntfy.sh/{state.ntfy_topic}",
            data=message.encode("utf-8"),
            headers={"Title": (title or "Локальный ИИ").encode("utf-8")},
            timeout=10,
        )
    except requests.RequestException as e:
        return f"Ошибка отправки уведомления: {e}"
    if resp.status_code != 200:
        return f"Ошибка отправки уведомления: сервер вернул код {resp.status_code}."
    return "Уведомление отправлено."


def tool_send_email(to: str, subject: str, body: str, emit, force_auto: bool = False) -> str:
    """Отправляет email через SMTP, настроенный в боковой панели."""
    if not (state.smtp_host and state.smtp_user and state.smtp_password and state.smtp_from):
        return "Ошибка: SMTP не настроен (host/user/password/from) — заполните настройки почты в боковой панели."
    to = (to or "").strip()
    if not to:
        return "Ошибка: не указан получатель."

    preview = f"Кому: {to}\nТема: {subject}\n\n{(body or '')[:500]}"
    approved = request_confirmation("write", f"email → {to}", preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ отправку письма."

    import smtplib
    from email.mime.text import MIMEText
    msg = MIMEText(body or "", "plain", "utf-8")
    msg["Subject"] = subject or "(без темы)"
    msg["From"] = state.smtp_from
    msg["To"] = to
    try:
        with smtplib.SMTP(state.smtp_host, state.smtp_port, timeout=20) as server:
            server.starttls()
            server.login(state.smtp_user, state.smtp_password)
            server.sendmail(state.smtp_from, [to], msg.as_string())
    except Exception as e:
        return f"Ошибка отправки письма: {e}"
    return f"Письмо отправлено на {to}."


def tool_telegram_send(message: str, emit, force_auto: bool = False) -> str:
    """Отправляет сообщение через Telegram-бота, настроенного в боковой панели
    (нужны telegram_bot_token и telegram_chat_id)."""
    if not (state.telegram_bot_token and state.telegram_chat_id):
        return "Ошибка: Telegram не настроен — заполните токен бота и chat_id в боковой панели."
    message = (message or "").strip()
    if not message:
        return "Ошибка: пустое сообщение."

    approved = request_confirmation("write", "Telegram", f"Отправить в Telegram:\n\n{message[:500]}",
                                     emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ отправку сообщения в Telegram."

    try:
        resp = requests.post(
            f"https://api.telegram.org/bot{state.telegram_bot_token}/sendMessage",
            json={"chat_id": state.telegram_chat_id, "text": message},
            timeout=15,
        )
        data = resp.json()
    except Exception as e:
        return f"Ошибка отправки в Telegram: {e}"
    if not data.get("ok"):
        return f"Ошибка Telegram API: {data.get('description', 'неизвестная ошибка')}"
    return "Сообщение отправлено в Telegram."


_repl_namespaces: dict = {}


def tool_python_repl(code: str, chat_id: str, emit, force_auto: bool = False) -> str:
    """Выполняет Python-код в постоянном интерпретаторе, сохраняющем
    переменные между вызовами в рамках одного чата. Как и run_shell_command,
    доступен только при включённой 'Автономности ИИ' и требует подтверждения —
    это полноценное выполнение кода на компьютере пользователя."""
    if not state.ai_autonomy:
        return ("Инструмент python_repl отключён. Включите «Автономность ИИ» "
                "в боковой панели сайта.")
    code = code or ""
    if not code.strip():
        return "Ошибка: пустой код."

    preview = f"```python\n{code}\n```"
    approved = request_confirmation("exec", "python_repl", preview, emit, force_auto=force_auto)
    if not approved:
        return "Пользователь ОТКЛОНИЛ выполнение кода."

    ns = _repl_namespaces.setdefault(chat_id or "default", {"__name__": "__main__"})
    import io
    import contextlib
    buf = io.StringIO()
    try:
        with contextlib.redirect_stdout(buf), contextlib.redirect_stderr(buf):
            try:
                compiled = compile(code, "<python_repl>", "eval")
                result = eval(compiled, ns)
                if result is not None:
                    print(repr(result))
            except SyntaxError:
                exec(compile(code, "<python_repl>", "exec"), ns)
    except Exception:
        buf.write("\n" + traceback.format_exc())
    out = buf.getvalue().strip() or "(без вывода)"
    if len(out) > 8000:
        out = out[:8000] + "\n… (обрезано)"
    return out


def tool_lint_format(path: str, tool: str, fix: bool, emit, force_auto: bool = False) -> str:
    """Проверяет (или, если fix=True, исправляет) стиль кода файла с помощью
    black/ruff/eslint — если они установлены в системе."""
    target = safe_path(path)
    if not target.exists():
        return f"Ошибка: файл '{path}' не найден."
    tool = (tool or "ruff").strip().lower()
    cmd_map = {
        "black": (["black", "--check", "--diff", str(target)], ["black", str(target)]),
        "ruff": (["ruff", "check", str(target)], ["ruff", "check", "--fix", str(target)]),
        "eslint": (["npx", "--yes", "eslint", str(target)], ["npx", "--yes", "eslint", "--fix", str(target)]),
    }
    if tool not in cmd_map:
        return f"Ошибка: неизвестный линтер '{tool}'. Доступны: black, ruff, eslint."
    check_cmd, fix_cmd = cmd_map[tool]

    if fix:
        approved = request_confirmation("write", path, f"Отформатировать/исправить файл ({tool} --fix)",
                                         emit, force_auto=force_auto)
        if not approved:
            return "Пользователь ОТКЛОНИЛ форматирование файла."
        cmd = fix_cmd
    else:
        cmd = check_cmd

    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
    except FileNotFoundError:
        return f"Ошибка: инструмент '{tool}' не установлен (pip install {tool} или npm i -g eslint)."
    except subprocess.TimeoutExpired:
        return "Ошибка: превышено время ожидания."
    out = (proc.stdout or "") + (("\n" + proc.stderr) if proc.stderr else "")
    return out.strip() or "Проблем не найдено."


def tool_get_weather(location: str) -> str:
    """Текущая погода для места (через бесплатный open-meteo API, без ключа)."""
    location = (location or "").strip()
    if not location:
        return "Ошибка: не указано место."
    try:
        geo = requests.get(
            "https://geocoding-api.open-meteo.com/v1/search",
            params={"name": location, "count": 1, "language": "ru"}, timeout=15,
        ).json()
    except requests.RequestException as e:
        return f"Ошибка геокодирования: {e}"
    results = geo.get("results") or []
    if not results:
        return f"Не удалось найти место «{location}»."
    r = results[0]
    lat, lon = r["latitude"], r["longitude"]
    name = r.get("name", location)
    country = r.get("country", "")
    try:
        wx = requests.get(
            "https://api.open-meteo.com/v1/forecast",
            params={"latitude": lat, "longitude": lon,
                    "current": "temperature_2m,relative_humidity_2m,wind_speed_10m,weather_code",
                    "timezone": "auto"}, timeout=15,
        ).json()
    except requests.RequestException as e:
        return f"Ошибка получения погоды: {e}"
    cur = wx.get("current") or {}
    if not cur:
        return f"Не удалось получить погоду для {name}."
    return (f"Погода в {name}, {country}: {cur.get('temperature_2m')}°C, "
            f"влажность {cur.get('relative_humidity_2m')}%, "
            f"ветер {cur.get('wind_speed_10m')} км/ч "
            f"(код погодного явления WMO: {cur.get('weather_code')}).")


def tool_get_exchange_rate(amount: float, from_currency: str, to_currency: str) -> str:
    """Курс валют через frankfurter.app (данные ЕЦБ, без ключа). Может не
    поддерживать некоторые валюты, недоступные в ЕЦБ."""
    from_currency = (from_currency or "USD").strip().upper()
    to_currency = (to_currency or "EUR").strip().upper()
    try:
        amount = float(amount or 1)
    except (TypeError, ValueError):
        amount = 1.0
    try:
        resp = requests.get(
            "https://api.frankfurter.app/latest",
            params={"amount": amount, "from": from_currency, "to": to_currency}, timeout=15,
        )
        data = resp.json()
    except Exception as e:
        return f"Ошибка получения курса: {e}"
    if resp.status_code != 200 or "rates" not in data:
        return (f"Не удалось получить курс {from_currency}→{to_currency} "
                f"(валюта может быть не поддержана источником).")
    rate = data["rates"].get(to_currency)
    if rate is None:
        return f"Валюта {to_currency} не найдена в ответе."
    return f"{amount} {from_currency} = {rate} {to_currency} (на дату {data.get('date')})"


def tool_rss_read(url: str, max_items: int = 10) -> str:
    """Читает RSS/Atom-ленту и возвращает последние записи (заголовок, ссылка, дата)."""
    import xml.etree.ElementTree as ET
    url = (url or "").strip()
    if not url:
        return "Ошибка: не указана ссылка на ленту."
    try:
        max_items = max(1, min(int(max_items or 10), 30))
    except (TypeError, ValueError):
        max_items = 10
    try:
        resp = requests.get(url, headers={"User-Agent": _WEB_UA}, timeout=15)
    except requests.RequestException as e:
        return f"Ошибка загрузки ленты: {e}"
    if resp.status_code != 200:
        return f"Ошибка загрузки ленты: сервер вернул код {resp.status_code}."
    try:
        root = ET.fromstring(resp.content)
    except ET.ParseError as e:
        return f"Ошибка разбора ленты: {e}"

    ns_atom = "{http://www.w3.org/2005/Atom}"
    items = []
    for item in root.iter("item"):  # RSS 2.0
        title = (item.findtext("title") or "").strip()
        link = (item.findtext("link") or "").strip()
        date = (item.findtext("pubDate") or "").strip()
        items.append((title, link, date))
    if not items:
        for entry in root.iter(f"{ns_atom}entry"):  # Atom
            title = (entry.findtext(f"{ns_atom}title") or "").strip()
            link_el = entry.find(f"{ns_atom}link")
            link = link_el.get("href") if link_el is not None else ""
            date = (entry.findtext(f"{ns_atom}updated") or "").strip()
            items.append((title, link, date))

    if not items:
        return "Записей в ленте не найдено (или неизвестный формат)."
    parts = [f"{i + 1}. {t}\n   {l}\n   {d}" for i, (t, l, d) in enumerate(items[:max_items])]
    return f"Лента {url}:\n\n" + "\n\n".join(parts)


def tool_set_reminder(minutes: float, message: str, chat_id: str) -> str:
    """Ставит однократное напоминание: через указанное число минут в панель
    активности текущего чата придёт уведомление с сообщением (и push через
    ntfy, если ntfy_topic настроен)."""
    try:
        minutes = float(minutes)
    except (TypeError, ValueError):
        return "Ошибка: некорректное количество минут."
    if minutes <= 0 or minutes > 1440:
        return "Ошибка: укажите от 0 до 1440 минут (24 часа)."
    message = (message or "Напоминание").strip()
    target_chat = chat_id or "default"

    def fire():
        card = (f'<div class="act-card act-info">'
                f'<div class="act-head"><span class="act-icon">⏰</span>'
                f'<span class="act-label">Напоминание</span></div>'
                f'<div class="act-desc">{html_lib.escape(message)}</div></div>')
        try:
            push_activity(target_chat, card)
        except Exception:
            pass
        if state.ntfy_topic:
            try:
                tool_notify(message, "Напоминание")
            except Exception:
                pass

    timer = threading.Timer(minutes * 60, fire)
    timer.daemon = True
    timer.start()
    return f"Напоминание поставлено через {minutes} мин.: «{message}»."



def friendly_ollama_error(e: Exception) -> str:
    """Превращает сырые ошибки Ollama (особенно длинный JSON про переполнение
    контекста) в понятное сообщение на русском, вместо простыни технического
    текста прямо в чате."""
    text = str(e)
    if "exceed_context_size_error" in text or "context size" in text.lower():
        m = re.search(r'"n_prompt_tokens":(\d+).*?"n_ctx":(\d+)', text)
        if m:
            prompt_tok, ctx = m.group(1), m.group(2)
            return (
                f"Промпт не влез в контекст модели: нужно {prompt_tok} токенов, а лимит "
                f"контекста сейчас {ctx}. Увеличьте размер контекста (поле рядом с выбором "
                f"модели в шапке сайта) хотя бы до {int(prompt_tok) + 2000}, либо начните "
                f"новый чат — этот стал слишком длинным для текущего лимита."
            )
        return (
            "Промпт не влез в лимит контекста модели. Увеличьте размер контекста (поле "
            "рядом с выбором модели) или начните новый чат."
        )
    return text[:600]


def tool_view_image(path: str):
    target = safe_path(path)
    if not target.exists() or not target.is_file():
        return None, f"Ошибка: файл не найден: {path}"
    if target.suffix.lower() not in IMAGE_EXTENSIONS:
        return None, f"'{path}' не похож на изображение."
    if target.stat().st_size > MAX_IMAGE_READ_BYTES:
        return None, f"Изображение слишком большое (> {MAX_IMAGE_READ_BYTES // 1_000_000} МБ)."
    try:
        data = base64.b64encode(target.read_bytes()).decode("ascii")
    except Exception as e:
        return None, f"Ошибка чтения изображения: {e}"
    return data, f"Изображение '{path}' отправлено модели."


TOOLS_SCHEMA = [
    {"type": "function", "function": {
        "name": "list_dir",
        "description": "Список файлов/папок. Пустой путь = список подключённых корневых папок, если их несколько.",
        "parameters": {"type": "object", "properties": {"path": {"type": "string"}}},
    }},
    {"type": "function", "function": {
        "name": "read_file",
        "description": "Прочитать текстовый файл.",
        "parameters": {"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]},
    }},
    {"type": "function", "function": {
        "name": "search_files",
        "description": "Найти текст во всех подключённых папках (рекурсивно) или в указанной подпапке.",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string"}, "path": {"type": "string"}}, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "write_file",
        "description": "Создать/перезаписать файл. Требует подтверждения пользователя.",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string"}, "content": {"type": "string"}}, "required": ["path", "content"]},
    }},
    {"type": "function", "function": {
        "name": "edit_file",
        "description": "Заменить уникальный фрагмент текста в файле. Требует подтверждения.",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string"}, "old_text": {"type": "string"}, "new_text": {"type": "string"}},
            "required": ["path", "old_text", "new_text"]},
    }},
    {"type": "function", "function": {
        "name": "view_image",
        "description": "Посмотреть изображение из разрешённой папки (нужна vision-модель).",
        "parameters": {"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]},
    }},
    {"type": "function", "function": {
        "name": "web_search",
        "description": (
            "Найти актуальную информацию в интернете (новости, факты, документацию и т.п.). "
            "Возвращает список результатов: заголовок, ссылка, краткое описание. Используй, "
            "когда нужны данные, которых нет в твоих знаниях, или свежая информация."
        ),
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Поисковый запрос"},
            "max_results": {"type": "integer", "description": "Сколько результатов вернуть (по умолчанию 5, максимум 8)"},
        }, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "fetch_url",
        "description": (
            "Открыть веб-страницу по ссылке и получить её текстовое содержимое "
            "(например, чтобы прочитать статью или страницу, найденную через web_search)."
        ),
        "parameters": {"type": "object", "properties": {
            "url": {"type": "string", "description": "Полный URL страницы"},
        }, "required": ["url"]},
    }},
    {"type": "function", "function": {
        "name": "wikipedia_search",
        "description": (
            "Поиск и краткое содержание статей Википедии — источник дополнительных "
            "фоновых знаний: определения, история, факты о людях/местах/явлениях. "
            "Используй как приоритетный источник справочной информации."
        ),
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Что искать в Википедии"},
            "lang": {"type": "string", "description": "Языковой раздел, например 'ru' или 'en' (по умолчанию ru)"},
            "max_results": {"type": "integer", "description": "Сколько статей вернуть (по умолчанию 3)"},
        }, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "google_search",
        "description": (
            "Поиск через Google — используй для просмотра актуальной информации в интернете "
            "и для того, чтобы найти прямые ссылки на файлы/страницы для последующего "
            "скачивания через download_file или чтения через fetch_url."
        ),
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Поисковый запрос"},
            "max_results": {"type": "integer", "description": "Сколько результатов вернуть (по умолчанию 5, максимум 10)"},
        }, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "github_search",
        "description": (
            "Поиск репозиториев на GitHub — используй, чтобы найти проект, инструмент "
            "или исходный код по теме. Возвращает название, ссылку, ссылку для клонирования "
            "и описание. Для скачивания найденного файла используй download_file."
        ),
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Поисковый запрос (можно использовать синтаксис GitHub, например 'language:python stars:>100')"},
            "max_results": {"type": "integer", "description": "Сколько результатов вернуть (по умолчанию 5, максимум 10)"},
        }, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "image_search",
        "description": (
            "Найти фотографии/картинки по теме (свободно лицензированные, через Openverse "
            "и Wikimedia Commons). Возвращает результат сразу в виде markdown ![]() — вставь "
            "эти строки в свой финальный ответ пользователю дословно, без изменений, чтобы "
            "картинки реально показались в чате. Используй, когда пользователь просит "
            "'покажи фото', 'как выглядит', 'пришли картинки' и т.п. Если инструмент ответил "
            "'не найдено' — вызови его ЕЩЁ РАЗ с более простым запросом (1-2 слова, лучше на "
            "английском), а не переключайся на google_search/web_search: они дают ссылки на "
            "страницы сайтов, а не на сами файлы картинок, и в чате не отобразятся."
        ),
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Что искать, например 'детская комната скандинавский стиль'"},
            "max_results": {"type": "integer", "description": "Сколько картинок (по умолчанию 6, максимум 12)"},
        }, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "video_search",
        "description": (
            "Найти видео на YouTube по теме. Вставь возвращённые ссылки в свой ответ "
            "пользователю как обычный текст (без изменений) — чат сам покажет встроенный "
            "видеоплеер: один — крупный, несколько — сеткой поменьше. Используй, когда "
            "пользователь просит видео, обзор, туториал и т.п."
        ),
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string"},
            "max_results": {"type": "integer", "description": "Сколько видео (по умолчанию 4, максимум 6)"},
        }, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "download_file",
        "description": (
            "Скачать файл по прямой ссылке (например, найденной через google_search) и "
            "сохранить его в одну из подключённых папок. Требует подтверждения пользователя."
        ),
        "parameters": {"type": "object", "properties": {
            "url": {"type": "string", "description": "Прямая ссылка на файл"},
            "path": {"type": "string", "description": "Путь для сохранения внутри подключённой папки"},
        }, "required": ["url", "path"]},
    }},
    {"type": "function", "function": {
        "name": "read_document",
        "description": "Прочитать текст из PDF, DOCX или XLSX-файла (обычный read_file подходит только для текстовых файлов).",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "Путь к файлу внутри подключённой папки"},
        }, "required": ["path"]},
    }},
    {"type": "function", "function": {
        "name": "ocr_image",
        "description": "Распознать текст на изображении (скриншот, скан, фото документа) через OCR.",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "Путь к изображению"},
            "lang": {"type": "string", "description": "Языки OCR через '+', например 'rus+eng' (по умолчанию rus+eng)"},
        }, "required": ["path"]},
    }},
    {"type": "function", "function": {
        "name": "archive_extract",
        "description": "Распаковать архив (zip/tar/tar.gz/tar.bz2/tar.xz) в указанную папку. Требует подтверждения.",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "Путь к архиву"},
            "dest": {"type": "string", "description": "Папка назначения"},
        }, "required": ["path", "dest"]},
    }},
    {"type": "function", "function": {
        "name": "archive_create",
        "description": "Упаковать файл или папку в архив (формат по расширению dest: .zip/.tar.gz/.tar.bz2/.tar.xz). Требует подтверждения.",
        "parameters": {"type": "object", "properties": {
            "source": {"type": "string", "description": "Файл или папка для упаковки"},
            "dest": {"type": "string", "description": "Путь к создаваемому архиву"},
        }, "required": ["source", "dest"]},
    }},
    {"type": "function", "function": {
        "name": "diff_files",
        "description": "Показать построчные различия (unified diff) между двумя текстовыми файлами.",
        "parameters": {"type": "object", "properties": {
            "path_a": {"type": "string"}, "path_b": {"type": "string"},
        }, "required": ["path_a", "path_b"]},
    }},
    {"type": "function", "function": {
        "name": "remember",
        "description": "Сохранить факт/заметку в долгосрочную память, которая переживает разные чаты (например, предпочтения пользователя, важные детали проекта).",
        "parameters": {"type": "object", "properties": {
            "key": {"type": "string", "description": "Короткий ключ заметки"},
            "text": {"type": "string", "description": "Текст заметки"},
        }, "required": ["key", "text"]},
    }},
    {"type": "function", "function": {
        "name": "recall",
        "description": "Достать заметку из долгосрочной памяти по ключу, либо список всех заметок, если ключ не указан.",
        "parameters": {"type": "object", "properties": {
            "key": {"type": "string", "description": "Ключ заметки (необязательно)"},
        }, "required": []},
    }},
    {"type": "function", "function": {
        "name": "save_solution",
        "description": (
            "Сохранить успешно решённую техническую задачу (запуск софта, скрипт, "
            "настройка системы) в архив готовых решений — чтобы в будущем не решать "
            "такую же задачу с нуля, а взять проверенный рабочий вариант. "
            "ВАЖНО: вызывай ТОЛЬКО после РЕАЛЬНОГО подтверждения результата, а не просто "
            "потому что команда вернула код 0. Код 0 означает лишь 'команда не упала с "
            "ошибкой' — это не то же самое, что 'результат действительно появился и "
            "работает' (например, .desktop-файл может успешно скопироваться командой cp, "
            "но не отображаться на рабочем столе или не запускаться, пока не станет "
            "'доверенным' в файловом менеджере). Подтверждением считается: пользователь "
            "прямо сказал, что всё сработало, ИЛИ ты независимо перепроверил результат "
            "(например, повторно прочитал/показал файл, проверил процесс через ps и т.п.), "
            "а не просто увидел код возврата 0 у команды, которая его создала."
        ),
        "parameters": {"type": "object", "properties": {
            "task": {"type": "string", "description": "Краткое описание задачи, например 'запуск LM Studio без sandbox на Linux'"},
            "solution": {"type": "string", "description": "Рабочее решение целиком: команда/флаги/шаги"},
        }, "required": ["task", "solution"]},
    }},
    {"type": "function", "function": {
        "name": "recall_solutions",
        "description": (
            "Вручную поискать в архиве готовых решений задачи, похожие на запрос. "
            "Обычно этого не требуется — сервер и так автоматически подмешивает похожие "
            "прошлые решения в начало разговора при каждом сообщении; используй этот "
            "инструмент, если хочешь поискать целенаправленно по другому запросу."
        ),
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string"},
            "max_results": {"type": "integer", "description": "Сколько записей вернуть (по умолчанию 5)"},
        }, "required": ["query"]},
    }},
    {"type": "function", "function": {
        "name": "process_list",
        "description": "Список запущенных процессов системы (имя, PID, CPU%, RAM%), опционально отфильтрованный по имени.",
        "parameters": {"type": "object", "properties": {
            "filter_text": {"type": "string", "description": "Подстрока для фильтра по имени процесса (необязательно)"},
        }, "required": []},
    }},
    {"type": "function", "function": {
        "name": "check_process_running",
        "description": (
            "Проверить, УЖЕ ли запущено приложение — по нескольким ключевым словам сразу "
            "(и по процессам, и по заголовкам открытых окон). Обычный pgrep по одному "
            "точному имени часто НЕ находит Electron/AppImage-приложения — их реальный "
            "процесс называется иначе, чем сама программа. ОБЯЗАТЕЛЬНО вызывай этот "
            "инструмент перед запуском/повторным запуском любого GUI-приложения, передав "
            "несколько вариантов слов (название программы, вендора, 'electron' и т.п.) — "
            "иначе рискуешь запустить несколько дублирующихся копий одного приложения."
        ),
        "parameters": {"type": "object", "properties": {
            "hints": {"type": "string", "description": "Ключевые слова через запятую, например: 'lmstudio, LM Studio, electron'"},
        }, "required": ["hints"]},
    }},
    {"type": "function", "function": {
        "name": "git_tool",
        "description": "Обёртка над git. action=status/diff/log — без подтверждения; add/commit/push/pull — требуют подтверждения пользователя.",
        "parameters": {"type": "object", "properties": {
            "action": {"type": "string", "description": "status | diff | log | add | commit | push | pull"},
            "repo_path": {"type": "string", "description": "Путь к репозиторию (внутри подключённой папки)"},
            "args": {"type": "string", "description": "Доп. аргументы командной строки, например \"-m 'fix bug'\" для commit"},
        }, "required": ["action", "repo_path"]},
    }},
    {"type": "function", "function": {
        "name": "notify",
        "description": "Отправить push-уведомление пользователю на телефон/десктоп через ntfy.sh (нужен настроенный ntfy_topic).",
        "parameters": {"type": "object", "properties": {
            "message": {"type": "string"}, "title": {"type": "string"},
        }, "required": ["message"]},
    }},
    {"type": "function", "function": {
        "name": "send_email",
        "description": "Отправить email через SMTP, настроенный пользователем в боковой панели. Требует подтверждения.",
        "parameters": {"type": "object", "properties": {
            "to": {"type": "string"}, "subject": {"type": "string"}, "body": {"type": "string"},
        }, "required": ["to", "subject", "body"]},
    }},
    {"type": "function", "function": {
        "name": "telegram_send",
        "description": "Отправить сообщение через Telegram-бота, настроенного пользователем в боковой панели. Требует подтверждения.",
        "parameters": {"type": "object", "properties": {
            "message": {"type": "string"},
        }, "required": ["message"]},
    }},
    {"type": "function", "function": {
        "name": "lint_format",
        "description": "Проверить (fix=false) или исправить (fix=true) стиль кода файла с помощью black/ruff/eslint, если они установлены. fix=true требует подтверждения.",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string"},
            "tool": {"type": "string", "description": "black | ruff | eslint"},
            "fix": {"type": "boolean", "description": "Исправить найденные проблемы (по умолчанию false — только проверка)"},
        }, "required": ["path", "tool"]},
    }},
    {"type": "function", "function": {
        "name": "get_weather",
        "description": "Текущая погода для указанного места (город/адрес).",
        "parameters": {"type": "object", "properties": {
            "location": {"type": "string"},
        }, "required": ["location"]},
    }},
    {"type": "function", "function": {
        "name": "get_exchange_rate",
        "description": "Курс обмена между двумя валютами (данные ЕЦБ через frankfurter.app).",
        "parameters": {"type": "object", "properties": {
            "amount": {"type": "number", "description": "Сумма (по умолчанию 1)"},
            "from_currency": {"type": "string", "description": "Код валюты-источника, например USD"},
            "to_currency": {"type": "string", "description": "Код целевой валюты, например EUR"},
        }, "required": ["from_currency", "to_currency"]},
    }},
    {"type": "function", "function": {
        "name": "rss_read",
        "description": "Прочитать RSS/Atom-ленту и вернуть последние записи (заголовок, ссылка, дата).",
        "parameters": {"type": "object", "properties": {
            "url": {"type": "string"},
            "max_items": {"type": "integer", "description": "Сколько записей вернуть (по умолчанию 10)"},
        }, "required": ["url"]},
    }},
    {"type": "function", "function": {
        "name": "set_reminder",
        "description": "Поставить однократное напоминание — через указанное число минут в панели активности появится уведомление (плюс push через ntfy, если настроен).",
        "parameters": {"type": "object", "properties": {
            "minutes": {"type": "number", "description": "Через сколько минут напомнить (1–1440)"},
            "message": {"type": "string", "description": "Текст напоминания"},
        }, "required": ["minutes", "message"]},
    }},
]

TOOL_SCHEMA_SHELL = {"type": "function", "function": {
    "name": "run_shell_command",
    "description": (
        "Выполнить произвольную команду в системном терминале (bash) компьютера "
        "пользователя. В отличие от остальных инструментов, НЕ ограничена "
        "подключёнными папками — полный доступ к файловой системе и системе в целом. "
        "Передай use_sudo=true, если команде нужны права администратора (root); "
        "пароль sudo уже сохранён на сервере, вводить его не нужно. Каждый вызов "
        "требует подтверждения пользователя (если не включено автоподтверждение). "
        "Используй с осторожностью — команда выполняется по-настоящему и может "
        "необратимо изменить систему.\n\n"
        "ВАЖНО про запуск GUI-приложений и долгих процессов: ВСЕГДА запускай их в фоне "
        "с редиректом вывода, например: "
        "nohup /путь/к/приложению > /dev/null 2>&1 & disown; echo запущено. "
        "Если вывод команды не редиректнуть, чтение результата может упереться в таймаут "
        "(120 сек), даже когда приложение УЖЕ успешно запустилось — в этом случае "
        "инструмент вернёт пометку '[таймаут]' с пояснением. Увидев '[таймаут]', "
        "проверь реальный статус через 'ps aux | grep <имя>' и, если процесс уже "
        "работает, считай задачу решённой — НЕ пытайся запускать снова другим способом, "
        "это плодит дублирующиеся процессы."
    ),
    "parameters": {"type": "object", "properties": {
        "command": {"type": "string", "description": "Команда bash для выполнения"},
        "use_sudo": {"type": "boolean", "description": "Выполнить с правами администратора через sudo"},
    }, "required": ["command"]},
}}

TOOL_SCHEMA_SYSTEMD = {"type": "function", "function": {
    "name": "systemd_control",
    "description": "Управление системной службой через systemctl (status/start/stop/restart/enable/disable). Затрагивает всю систему.",
    "parameters": {"type": "object", "properties": {
        "action": {"type": "string", "description": "status | start | stop | restart | enable | disable"},
        "service": {"type": "string", "description": "Имя службы, например nginx"},
        "use_sudo": {"type": "boolean", "description": "Выполнить через sudo (по умолчанию true)"},
    }, "required": ["action", "service"]},
}}

TOOL_SCHEMA_PYREPL = {"type": "function", "function": {
    "name": "python_repl",
    "description": "Выполнить Python-код в постоянном интерпретаторе (переменные сохраняются между вызовами в рамках чата) — реальное выполнение кода на компьютере пользователя.",
    "parameters": {"type": "object", "properties": {
        "code": {"type": "string", "description": "Python-код"},
    }, "required": ["code"]},
}}


def get_tools_schema() -> list:
    """Список инструментов, передаваемый модели. run_shell_command / systemd_control /
    python_repl добавляются, только если пользователь явно включил 'Автономность ИИ' —
    без этого модель даже не узнаёт о существовании таких инструментов."""
    if state.ai_autonomy:
        return TOOLS_SCHEMA + [TOOL_SCHEMA_SHELL, TOOL_SCHEMA_SYSTEMD, TOOL_SCHEMA_PYREPL]
    return TOOLS_SCHEMA

SYSTEM_PROMPT_TEMPLATE = """{folders}"""


def get_model_custom_system(model_name: str) -> str:
    """Спрашивает у Ollama, есть ли у модели свой собственный SYSTEM,
    зашитый через Modelfile (ollama create) — например персона, контекст
    о железе пользователя и т.п. Если есть, мы должны его СОХРАНИТЬ,
    а не молча перекрывать своим системным промптом."""
    try:
        resp = requests.post(f"{state.ollama_host}/api/show", json={"name": model_name}, timeout=5)
        if resp.status_code == 200:
            return (resp.json().get("system") or "").strip()
    except Exception:
        pass
    return ""


KNOWLEDGE_SOURCE_LABELS = {
    "internet": (
        "Интернет (wikipedia_search, google_search, web_search, github_search, "
        "image_search, video_search, fetch_url)"
    ),
    "archive": "Архив успешных решений (recall_solutions и автоматический дайджест похожих задач)",
    "memory": "Собственная долгосрочная память (remember/recall — заметки, которые ты сам сохранял раньше)",
}


def knowledge_priority_instruction() -> str:
    order = state.knowledge_priority or ["internet", "archive", "memory"]
    lines = [
        "Пользователь задал порядок приоритета источников знаний для технических/фактических "
        "вопросов. Сначала проверяй источник с более высоким приоритетом (1) и переходи к "
        "следующему по списку, только если там не нашлось релевантного ответа. Не пропускай "
        "более приоритетный источник просто потому, что кажется, будто следующий быстрее:"
    ]
    for i, key in enumerate(order, start=1):
        label = KNOWLEDGE_SOURCE_LABELS.get(key, key)
        lines.append(f"{i}. {label}")
    if order and order[0] == "archive":
        lines.append(
            "Архив стоит на 1-м месте — это значит, что перед тем как отвечать на "
            "технический/повторяющийся вопрос своими силами, ты ОБЯЗАН явно вызвать "
            "recall_solutions с сутью текущего запроса, даже если тебе кажется, что там "
            "ничего нет. Не полагайся только на автоматическое подмешивание похожих "
            "решений в начало этого сообщения — оно может не сработать, если ты "
            "сформулировал запрос заметно по-другому, чем при сохранении. Начинай "
            "придумывать решение с нуля только ПОСЛЕ того, как recall_solutions явно "
            "ответил, что подходящего решения нет."
        )
    return "\n".join(lines)


def build_system_message(user_text: str = "") -> dict:
    folders = "\n".join(f"- {f}" for f in state.allowed_folders) if state.allowed_folders else "(нет)"
    folder_info = f"Подключённые папки:\n{folders}"

    custom_system = get_model_custom_system(state.model)
    parts = [custom_system] if custom_system else []
    parts.append(folder_info)
    parts.append(knowledge_priority_instruction())

    # Автоматический дайджест архива решений — подмешивается СЕРВЕРОМ на
    # каждом сообщении (см. solutions_auto_digest), поэтому не зависит от
    # того, вспомнит ли модель сама проверить архив.
    digest = solutions_auto_digest(user_text)
    if digest:
        parts.append(digest)
    parts.append(
        "У тебя также есть инструменты save_solution (сохранить успешно решённую "
        "техническую задачу в архив: команда/скрипт/настройка, которые реально сработали) "
        "и recall_solutions (вручную поискать в архиве по запросу). Похожие прошлые решения "
        "сервер и так подмешивает сюда автоматически при каждом сообщении — но после того, "
        "как решишь НОВУЮ техническую задачу (запуск софта, скрипт, настройка системы) и "
        "убедишься, что решение реально сработало, обязательно вызови save_solution, чтобы "
        "не решать её заново в следующий раз. НЕ вызывай save_solution только на основании "
        "кода возврата 0 у команды — это значит лишь 'команда не упала с ошибкой', а не "
        "'результат действительно появился и работает'. Дождись подтверждения от "
        "пользователя или сам независимо перепроверь результат другим способом, прежде чем "
        "сохранять решение как успешное."
    )
    parts.append(
        "Полезный факт про создание ярлыков (.desktop-файлов) на Linux Mint/Cinnamon "
        "(Nemo) и других GTK-окружениях: одного 'chmod +x' и копирования файла на "
        "рабочий стол НЕДОСТАТОЧНО — файловый менеджер по умолчанию считает такие "
        "файлы 'недоверенными' и не показывает иконку/не даёт запустить их, пока не "
        "пометишь файл доверенным командой "
        "'gio set ~/Desktop/файл.desktop metadata::trusted true' (или пользователь "
        "не кликнет правой кнопкой → 'Allow Launching / Разрешить запуск'). Учитывай "
        "это при создании ярлыков и не считай задачу решённой, пока не упомянул этот шаг."
    )
    parts.append(
        "У тебя есть инструменты для работы с интернетом: wikipedia_search "
        "(приоритетный источник фоновых знаний — определения, факты, история), "
        "google_search (поиск актуальной информации и ссылок в Google), github_search "
        "(поиск репозиториев и кода на GitHub), fetch_url "
        "(прочитать содержимое конкретной страницы по ссылке) и download_file "
        "(скачать файл по прямой ссылке на диск). Для общих фактов и определений "
        "сначала пробуй wikipedia_search, для актуальных новостей и поиска "
        "конкретных страниц/ссылок используй google_search, для проектов и кода — "
        "github_search, а не отвечай по памяти и не придумывай ссылки."
    )
    parts.append(
        "Когда пользователь просит показать фото/картинки ('покажи фото', 'как выглядит', "
        "'пришли картинки') — используй image_search, а затем ОБЯЗАТЕЛЬНО скопируй "
        "markdown-строки вида ![...](...) из результата инструмента в свой финальный ответ "
        "БЕЗ ИЗМЕНЕНИЙ — только тогда картинки реально отобразятся в чате пользователя. "
        "Если image_search ответил 'не найдено', вызови его ЕЩЁ РАЗ с более простым и/или "
        "английским запросом (1-2 слова) — НЕ переключайся на google_search/web_search "
        "вместо этого: результаты обычного веб-поиска — это ссылки на страницы сайтов "
        "(галереи, магазины), а не на сами файлы изображений, и картинки из них в чате "
        "не покажутся. Когда просит видео/обзор/туториал — используй video_search и вставь "
        "возвращённые ссылки на YouTube в ответ как обычный текст без изменений — чат сам "
        "покажет встроенный видеоплеер."
    )
    parts.append(
        "КРИТИЧЕСКИ ВАЖНО: НИКОГДА не пиши markdown-ссылку на картинку ![...](...) с URL, "
        "который ты не скопировал дословно из результата image_search (или fetch_url/"
        "web_search). В частности, НИКОГДА не сочиняй ссылки вида "
        "'https://images.unsplash.com/photo-XXXXXXXXXX-XXXXXXXXXXXX' или похожие на других "
        "стоках по памяти — такие ID случайны, ты не можешь угадать реальный, и ссылка "
        "гарантированно окажется битой (404). Если ты хочешь показать картинку, но ещё не "
        "вызывал image_search в этом ответе — сначала вызови его по-настоящему, не пропускай "
        "этот шаг и не подменяй его придуманным URL."
    )
    parts.append(
        "Дополнительные инструменты: read_document (текст из PDF/DOCX/XLSX), ocr_image "
        "(распознать текст на изображении), archive_extract/archive_create (архивы), "
        "diff_files (сравнить два файла), remember/recall (долгосрочная память между "
        "чатами — используй remember, когда пользователь сообщает важный факт о себе "
        "или проекте, который пригодится в будущем), process_list (запущенные процессы), "
        "git_tool (git-команды), notify (push-уведомление пользователю), send_email и "
        "telegram_send (требуют настроенных SMTP/Telegram в боковой панели), lint_format "
        "(проверка/форматирование кода), get_weather, get_exchange_rate, rss_read и "
        "set_reminder (напоминание через N минут)."
    )

    if state.ai_autonomy:
        parts.append(
            "Режим «Автономность ИИ» включён: тебе доступен инструмент run_shell_command "
            "для выполнения команд в системном терминале компьютера пользователя, при "
            "необходимости с правами root (use_sudo=true). Это выходит за пределы "
            "подключённых папок и затрагивает всю систему. Используй его обдуманно: "
            "объясняй пользователю, что и зачем собираешься выполнить, избегай "
            "разрушительных команд (удаление системных файлов, форматирование, "
            "'rm -rf /', изменение критичных системных настроек) без явной необходимости "
            "и явного запроса пользователя. Запуская GUI-приложения или долгие процессы, "
            "ВСЕГДА уводи их в фон с редиректом вывода (nohup ... > /dev/null 2>&1 & disown), "
            "иначе можно словить '[таймаут]' даже при успешном запуске. Если ответ содержит "
            "'[таймаут]' — сначала проверь 'ps aux | grep <имя>', и если процесс уже "
            "работает, считай задачу выполненной; НЕ запускай программу заново другим "
            "способом только из-за пометки таймаута. "
            "ПЕРЕД запуском/повторным запуском ЛЮБОГО GUI-приложения ОБЯЗАТЕЛЬНО сначала "
            "вызови check_process_running с несколькими вариантами ключевых слов сразу "
            "(название программы, вендор, 'electron' и т.п.) — обычный pgrep по одному "
            "точному имени часто не находит Electron/AppImage-процессы (у них настоящее имя "
            "процесса не совпадает с названием программы), из-за чего легко случайно "
            "наплодить несколько запущенных копий одного приложения. У тебя ЕСТЬ доступ к "
            "терминалу через run_shell_command — НИКОГДА не проси пользователя самого "
            "открыть терминал и ввести команду, которую можешь выполнить сам; исключение — "
            "только если команда требует ручного взаимодействия с GUI (клик, ввод пароля "
            "в диалоге), которое ты физически не можешь сделать за него."
        )

    if state.think_enabled:
        # Модели-рассуждатели (Qwen3.5 и т.п.) по умолчанию "думают" на
        # английском и часто начинают блок размышлений с пересказа вопроса
        # пользователя ("Пользователь спрашивает о...") — просим этого не
        # делать и сразу переходить к сути анализа, на русском языке.
        parts.append(
            "Когда рассуждаешь в блоке размышлений (thinking): думай СТРОГО на "
            "русском языке, никогда не переключайся на английский. Не пересказывай "
            "и не повторяй вопрос пользователя в начале рассуждения (не пиши фразы "
            "вроде \"Пользователь спрашивает о...\" или \"The user is asking...\") — "
            "сразу переходи к сути анализа."
        )

    return {"role": "system", "content": "\n\n".join(parts)}


# ----------------------------- Общение с Ollama -------------------------------

def flatten_for_strict_template(messages: list[dict]) -> list[dict]:
    """Некоторые шаблоны моделей (Gemma, Mistral и т.п.) требуют строгого
    чередования ролей user/assistant и не понимают роль 'tool' или системные
    сообщения посередине диалога. Эта функция сплющивает историю в предельно
    простую и совместимую форму: system сливается в первое user-сообщение,
    роль 'tool' и вспомогательные сообщения присоединяются текстом к ближайшему
    user-сообщению, соседние сообщения одной роли объединяются."""
    flat = []
    pending_user_parts = []
    system_text = None

    def flush_user():
        nonlocal pending_user_parts
        if pending_user_parts:
            flat.append({"role": "user", "content": "\n\n".join(pending_user_parts)})
            pending_user_parts = []

    for m in messages:
        role = m.get("role")
        content = m.get("content") or ""
        if role == "system":
            system_text = content
        elif role == "user":
            pending_user_parts.append(content)
        elif role == "tool":
            pending_user_parts.append(f"[результат инструмента]\n{content}")
        elif role == "assistant":
            flush_user()
            if content:
                flat.append({"role": "assistant", "content": content})
        else:
            pending_user_parts.append(content)
    flush_user()

    if system_text:
        if flat and flat[0]["role"] == "user":
            flat[0]["content"] = system_text + "\n\n" + flat[0]["content"]
        else:
            flat.insert(0, {"role": "user", "content": system_text})

    cleaned = []
    for m in flat:
        if not m["content"].strip():
            continue
        if cleaned and cleaned[-1]["role"] == m["role"]:
            cleaned[-1]["content"] += "\n\n" + m["content"]
        else:
            cleaned.append(dict(m))
    return cleaned


def stream_ollama(messages, cancel_flag=None):
    base_payload = {"model": state.model, "messages": messages, "stream": True,
                     "think": state.think_enabled, "options": state.reasoning_options()}

    payload = dict(base_payload, tools=get_tools_schema())
    resp = requests.post(f"{state.ollama_host}/api/chat", json=payload, stream=True, timeout=None)

    if resp.status_code >= 400 and "tool" in resp.text.lower() and "alternate" not in resp.text.lower():
        # Модель не поддерживает function calling (например, чисто vision-модели
        # вроде Qwen2.5-VL). Тихо повторяем запрос обычным чатом, без инструментов.
        resp.close()
        resp = requests.post(f"{state.ollama_host}/api/chat", json=base_payload, stream=True, timeout=None)

    if resp.status_code >= 400 and ("alternate" in resp.text.lower() or "raise_exception" in resp.text.lower()):
        # Шаблон модели требует строгого чередования user/assistant и не понимает
        # роль 'tool' в истории — тихо упрощаем историю и повторяем без инструментов.
        # Заодно сохраняем очищенную историю в сам диалог, чтобы ошибка не
        # повторялась на каждом следующем сообщении в этом же чате.
        resp.close()
        flat_messages = flatten_for_strict_template(messages)
        flat_payload = {"model": state.model, "messages": flat_messages, "stream": True,
                         "think": state.think_enabled, "options": state.reasoning_options()}
        resp = requests.post(f"{state.ollama_host}/api/chat", json=flat_payload, stream=True, timeout=None)
        if resp.status_code == 200:
            # для сохранения в историю чата выпрямляем без системного сообщения —
            # оно и так добавляется заново на каждый следующий ход отдельно
            history_only = flatten_for_strict_template(messages[1:])
            yield {"__replace_history__": history_only}

    if resp.status_code != 200:
        raise RuntimeError(f"Ollama вернул {resp.status_code}: {resp.text[:400]}")

    try:
        for line in resp.iter_lines():
            if cancel_flag is not None and cancel_flag.get("cancelled"):
                yield {"__cancelled__": True}
                return
            if not line:
                continue
            try:
                yield json.loads(line)
            except json.JSONDecodeError:
                continue
    finally:
        resp.close()


def ask_vision_model(image_b64: str, question: str) -> str:
    payload = {"model": state.vision_model,
               "messages": [{"role": "user", "content": question, "images": [image_b64]}],
               "stream": False}
    resp = requests.post(f"{state.ollama_host}/api/chat", json=payload, timeout=None)
    if resp.status_code != 200:
        raise RuntimeError(f"Vision-модель вернула {resp.status_code}: {resp.text[:300]}")
    return resp.json().get("message", {}).get("content", "").strip()


# ----------------------------- SSE-эндпоинт чата ------------------------------

def sse(event: dict) -> str:
    return f"data: {json.dumps(event, ensure_ascii=False)}\n\n"


@app.route("/api/chat", methods=["POST"])
def api_chat():
    body = request.get_json(force=True)
    user_text = body.get("message", "")
    image_b64 = body.get("image")
    files = body.get("files", [])  # [{name, content, type}]

    # Фиксируем, в какой именно чат идёт этот ответ, ПРЯМО СЕЙЧАС — если пользователь
    # переключится на другой чат, пока ИИ ещё печатает, ответ всё равно сохранится
    # туда, куда был реально адресован, а не 'уплывёт' в чат, открытый в данный момент.
    target_chat_id = state.current_chat_id

    cancel_flag = {"cancelled": False}
    state.active_cancel = cancel_flag

    def generate():
        def emit(ev):
            pass  # заменяется ниже очередью

        events = []
        emit_lock = threading.Lock()

        def emit(ev):
            with emit_lock:
                events.append(ev)
            # 'done' — общая точка выхода для ВСЕХ сценариев (отмена, ошибка,
            # обычное завершение) — снимаем здесь флаг генерации, чтобы не
            # дублировать эту логику в каждом return внутри worker().
            if ev.get("type") == "done":
                chat = state.chats.get(target_chat_id)
                if chat is not None:
                    chat["generating"] = False
                # Финальное состояние сохраняем на диск сразу, а не ждём
                # фонового потока — на "горячем пути" (частые touch() во
                # время генерации) это неважно, а вот на самом завершении
                # результат должен попасть на диск без задержки.
                state._save_chats_now()

        # Собрать изображения и текст из вложений
        all_images = []
        file_texts = []
        if image_b64:
            all_images.append(image_b64)
        for f in files:
            fname = f.get("name", "файл")
            fcontent = f.get("content", "")
            ftype = f.get("type", "")
            ext = Path(fname).suffix.lower()
            is_image = ext in IMAGE_EXTENSIONS or ftype.startswith("image/")
            if is_image and fcontent:
                # Изображение — отправляем как base64 image
                all_images.append(fcontent)
            elif fcontent:
                # Текстовый файл — вставляем содержимое в сообщение
                file_texts.append(f"--- Файл: {fname} ---\n{fcontent}\n--- Конец файла: {fname} ---")
            else:
                file_texts.append(f"[Вложение: {fname} ({ftype or 'неизвестный тип'}) — содержимое не может быть прочитано]")

        # Собрать итоговый текст сообщения
        parts = []
        if file_texts:
            parts.append("\n\n".join(file_texts))
        if user_text:
            parts.append(user_text)
        final_text = "\n\n".join(parts) if parts else ""

        user_msg = {"role": "user", "content": final_text, "time": time.time()}
        if all_images:
            user_msg["images"] = all_images
        # Сохраняем информацию о вложениях для отображения в чате
        if files:
            user_msg["attachments"] = [{"name": f.get("name", ""), "type": f.get("type", "")} for f in files]
        cur = state.messages_for(target_chat_id)
        cur.append(user_msg)
        state.touch(target_chat_id, maybe_title=user_text if user_text else None)
        # Флаг "идёт генерация" — используется как страховка на фронтенде:
        # если основной механизм live-рендера (моментальное переключение
        # DOM-узла) по какой-то причине не подхватит новые токены сразу после
        # возврата в чат, опрос по этому флагу всё равно подтянет актуальный
        # текст в течение пары секунд.
        chat_obj = state.chats.get(target_chat_id)
        if chat_obj is not None:
            chat_obj["generating"] = True
            state.save_chats()

        def worker():
            messages = [build_system_message(user_text)] + state.messages_for(target_chat_id)
            retried = False
            # Периодически слать "thinking" событие, чтобы фронтенд знал что ИИ жив
            thinking_timer = [None]
            def send_thinking():
                while True:
                    time.sleep(5)
                    emit({"type": "thinking", "elapsed": int(time.time() - start_time)})
            start_time = time.time()
            thinking_timer[0] = threading.Thread(target=send_thinking, daemon=True)
            thinking_timer[0].start()

            for _ in range(MAX_TOOL_ITERATIONS):
                full_content = ""
                full_thinking = ""
                tool_calls = None
                token_count = 0
                try:
                    for chunk in stream_ollama(messages, cancel_flag=cancel_flag):
                        if "__cancelled__" in chunk:
                            if full_content:
                                state.messages_for(target_chat_id).append(
                                    {"role": "assistant", "content": full_content + "\n\n[остановлено пользователем]"})
                            emit({"type": "cancelled"})
                            emit({"type": "done"})
                            return
                        if "__replace_history__" in chunk:
                            # сохраняем уже "выпрямленную" историю, чтобы та же самая
                            # ошибка шаблона не повторялась на каждом следующем сообщении
                            cur_chat = state.chats.get(target_chat_id)
                            if cur_chat is not None:
                                cur_chat["messages"] = chunk["__replace_history__"]
                                state.save_chats()
                            continue
                        msg = chunk.get("message", {})
                        delta = msg.get("content", "")
                        thinking = msg.get("thinking", "")
                        if thinking:
                            full_thinking += thinking
                            emit({"type": "thinking_token", "content": thinking})
                            # Периодически сохранять рассуждения (каждые 200 символов)
                            if len(full_thinking) % 200 < 20 and full_thinking.strip():
                                cur = state.messages_for(target_chat_id)
                                if cur and cur[-1].get("role") == "assistant" and cur[-1].get("_partial"):
                                    cur.pop()
                                display = full_thinking.strip()
                                if full_content.strip():
                                    display += "\n\n---\n\n" + full_content.strip()
                                cur.append({"role": "assistant", "content": display, "_partial": True, "time": time.time()})
                                state.touch(target_chat_id)
                        if delta:
                            full_content += delta
                            token_count += 1
                            emit({"type": "token", "content": delta, "n": token_count})
                            # Периодически сохранять частичный ответ в чат (каждые 20 токенов)
                            if token_count % 20 == 0 and (full_content.strip() or full_thinking.strip()):
                                cur = state.messages_for(target_chat_id)
                                if cur and cur[-1].get("role") == "assistant" and cur[-1].get("_partial"):
                                    cur.pop()
                                display = full_thinking.strip()
                                if full_content.strip():
                                    display += "\n\n---\n\n" + full_content.strip() if display else full_content.strip()
                                cur.append({"role": "assistant", "content": display, "_partial": True, "time": time.time()})
                                state.touch(target_chat_id)
                        if msg.get("tool_calls"):
                            tool_calls = msg["tool_calls"]
                        if chunk.get("done"):
                            eval_count = chunk.get("eval_count")
                            eval_duration = chunk.get("eval_duration")  # наносекунды
                            if eval_count and eval_duration:
                                emit({
                                    "type": "token_stats",
                                    "eval_count": eval_count,
                                    "tokens_per_sec": round(eval_count / (eval_duration / 1_000_000_000), 1),
                                    "prompt_eval_count": chunk.get("prompt_eval_count"),
                                    "total_duration_s": round(chunk.get("total_duration", 0) / 1_000_000_000, 2),
                                })
                            break
                except Exception as e:
                    emit({"type": "error", "message": friendly_ollama_error(e)})
                    emit({"type": "done"})
                    return

                if not tool_calls:
                    if not full_content.strip():
                        # Если были рассуждения но нет контента — показать рассуждения
                        if full_thinking.strip():
                            full_content = full_thinking.strip()
                        elif not retried:
                            retried = True
                            emit({"type": "notice", "message": "Пустой ответ, повторяю запрос…"})
                            continue
                    # Если были только thinking токены и нет контента — не шлём done,
                    # модель возможно ещё решит вызвать инструменты
                    if full_thinking.strip() and not full_content.strip() and not retried:
                        retried = True
                        emit({"type": "notice", "message": "Модель рассуждала, повторяю для получения ответа…"})
                        continue
                    # === ПАРСИНГ ТЕКСТА: если модель описала создание файла ===
                    # Qwen3 не генерирует tool_calls — парсим текст и выполняем сами
                    import re as _re
                    text_to_parse = full_content.strip()

                    # Паттерн: "Файл сохранён в /path/file.ext"
                    file_path_match = _re.search(
                        r'(?:сохранён|сохранено|создан|записан|Создам|Сохрани[мт])\s+(?:в|по пути|по адресу)?\s*[`"]?(/[^\s`"\n,]+?\.\w+)[`"]?',
                        text_to_parse, _re.IGNORECASE
                    )
                    if file_path_match:
                        fpath = file_path_match.group(1).strip(".")
                        # Найти код в блоках ```
                        code_match = _re.search(r'```(?:html|javascript|css|python|bash)?\s*\n(.*?)```', text_to_parse, _re.DOTALL)
                        if code_match:
                            code = code_match.group(1).strip()
                        else:
                            code = ""
                        if code and len(code) > 20:
                            emit({"type": "notice", "message": f"Создаю файл: {fpath}"})
                            try:
                                target = safe_path(fpath)
                                target.parent.mkdir(parents=True, exist_ok=True)
                                target.write_text(code, encoding="utf-8")
                                emit({"type": "tool_call", "name": "write_file", "args": {"path": fpath}})
                                emit({"type": "tool_result", "name": "write_file", "result": f"Файл '{fpath}' создан ({len(code)} символов)."})
                            except Exception as e:
                                emit({"type": "tool_result", "name": "write_file", "result": f"Ошибка: {e}"})

                    # Финальное сохранение
                    cur = state.messages_for(target_chat_id)
                    if cur and cur[-1].get("role") == "assistant" and cur[-1].get("_partial"):
                        cur.pop()
                    display = full_thinking.strip() + "\n\n---\n\n" + full_content.strip() if full_thinking.strip() and full_content.strip() != full_thinking.strip() else full_content.strip()
                    cur.append({"role": "assistant", "content": display, "time": time.time()})
                    state.touch(target_chat_id)
                    emit({"type": "done"})
                    return

                # Удалить частичный ответ перед сохранением с tool_calls
                cur = state.messages_for(target_chat_id)
                if cur and cur[-1].get("role") == "assistant" and cur[-1].get("_partial"):
                    cur.pop()
                assistant_msg = {"role": "assistant", "content": full_content, "tool_calls": tool_calls}
                messages.append(assistant_msg)
                cur.append(assistant_msg)

                for call in tool_calls:
                    fn = call.get("function", {})
                    name = fn.get("name")
                    args = fn.get("arguments") or {}
                    if isinstance(args, str):
                        try:
                            args = json.loads(args)
                        except json.JSONDecodeError:
                            args = {}
                    emit({"type": "tool_call", "name": name, "args": args})

                    try:
                        if name == "view_image":
                            img_b64, result = tool_view_image(**args)
                            if img_b64 and state.vision_model and state.vision_model != state.model:
                                emit({"type": "vision_thinking", "model": state.vision_model})
                                question = ("Подробно опиши, что изображено на этой картинке, включая "
                                            "весь видимый текст, интерфейсы, диаграммы и детали.")
                                description = ask_vision_model(img_b64, question)
                                emit({"type": "vision_result", "model": state.vision_model, "content": description})
                                tool_result = f"[Описание от vision-модели '{state.vision_model}']\n{description}"
                                tool_msg = {"role": "tool", "content": tool_result}
                                messages.append(tool_msg)
                                state.messages_for(target_chat_id).append(tool_msg)
                                continue
                            tool_msg = {"role": "tool", "content": str(result)}
                            messages.append(tool_msg)
                            state.messages_for(target_chat_id).append(tool_msg)
                            if img_b64:
                                image_msg = {"role": "user", "content": f"Изображение '{args.get('path','')}':",
                                             "images": [img_b64]}
                                messages.append(image_msg)
                                state.messages_for(target_chat_id).append(image_msg)
                            continue

                        if name == "write_file":
                            result = tool_write_file(args.get("path", ""), args.get("content", ""), emit)
                        elif name == "edit_file":
                            result = tool_edit_file(args.get("path", ""), args.get("old_text", ""),
                                                     args.get("new_text", ""), emit)
                        elif name == "list_dir":
                            result = tool_list_dir(args.get("path", ""))
                        elif name == "read_file":
                            result = tool_read_file(args.get("path", ""))
                        elif name == "search_files":
                            result = tool_search_files(args.get("query", ""), args.get("path", ""))
                        elif name == "run_shell_command":
                            result = tool_run_shell(args.get("command", ""), emit,
                                                     use_sudo=bool(args.get("use_sudo")))
                        elif name == "web_search":
                            result = tool_web_search(args.get("query", ""), args.get("max_results", 5))
                        elif name == "fetch_url":
                            result = tool_fetch_url(args.get("url", ""))
                        elif name == "wikipedia_search":
                            result = tool_wikipedia_search(args.get("query", ""), args.get("lang", "ru"),
                                                            args.get("max_results", 3))
                        elif name == "google_search":
                            result = tool_google_search(args.get("query", ""), args.get("max_results", 5))
                        elif name == "github_search":
                            result = tool_github_search(args.get("query", ""), args.get("max_results", 5))
                        elif name == "image_search":
                            result = tool_image_search(args.get("query", ""), args.get("max_results", 6))
                        elif name == "video_search":
                            result = tool_video_search(args.get("query", ""), args.get("max_results", 4))
                        elif name == "download_file":
                            result = tool_download_file(args.get("url", ""), args.get("path", ""), emit)
                        elif name == "read_document":
                            result = tool_read_document(args.get("path", ""))
                        elif name == "ocr_image":
                            result = tool_ocr_image(args.get("path", ""), args.get("lang", "rus+eng"))
                        elif name == "archive_extract":
                            result = tool_archive_extract(args.get("path", ""), args.get("dest", ""), emit)
                        elif name == "archive_create":
                            result = tool_archive_create(args.get("source", ""), args.get("dest", ""), emit)
                        elif name == "diff_files":
                            result = tool_diff_files(args.get("path_a", ""), args.get("path_b", ""))
                        elif name == "remember":
                            result = tool_remember(args.get("key", ""), args.get("text", ""))
                        elif name == "recall":
                            result = tool_recall(args.get("key", ""))
                        elif name == "save_solution":
                            result = tool_save_solution(args.get("task", ""), args.get("solution", ""))
                        elif name == "recall_solutions":
                            result = tool_recall_solutions(args.get("query", ""), args.get("max_results", 5))
                        elif name == "process_list":
                            result = tool_process_list(args.get("filter_text", ""))
                        elif name == "check_process_running":
                            result = tool_check_process_running(args.get("hints", ""))
                        elif name == "git_tool":
                            result = tool_git(args.get("action", ""), args.get("repo_path", ""),
                                               args.get("args", ""), emit)
                        elif name == "systemd_control":
                            result = tool_systemd(args.get("action", ""), args.get("service", ""), emit,
                                                   use_sudo=args.get("use_sudo", True))
                        elif name == "notify":
                            result = tool_notify(args.get("message", ""), args.get("title", ""))
                        elif name == "send_email":
                            result = tool_send_email(args.get("to", ""), args.get("subject", ""),
                                                      args.get("body", ""), emit)
                        elif name == "telegram_send":
                            result = tool_telegram_send(args.get("message", ""), emit)
                        elif name == "python_repl":
                            result = tool_python_repl(args.get("code", ""), target_chat_id, emit)
                        elif name == "lint_format":
                            result = tool_lint_format(args.get("path", ""), args.get("tool", "ruff"),
                                                       bool(args.get("fix", False)), emit)
                        elif name == "get_weather":
                            result = tool_get_weather(args.get("location", ""))
                        elif name == "get_exchange_rate":
                            result = tool_get_exchange_rate(args.get("amount", 1), args.get("from_currency", ""),
                                                             args.get("to_currency", ""))
                        elif name == "rss_read":
                            result = tool_rss_read(args.get("url", ""), args.get("max_items", 10))
                        elif name == "set_reminder":
                            result = tool_set_reminder(args.get("minutes", 0), args.get("message", ""), target_chat_id)
                        else:
                            result = f"Неизвестный инструмент: {name}"
                    except SandboxError as e:
                        result = f"Ошибка доступа: {e}"
                    except TypeError as e:
                        result = f"Ошибка аргументов: {e}"
                    except Exception as e:
                        result = f"Ошибка выполнения: {e}"

                    tool_msg = {"role": "tool", "content": str(result)}
                    messages.append(tool_msg)
                    state.messages_for(target_chat_id).append(tool_msg)
                    emit({"type": "tool_result", "name": name, "result": str(result)})

            emit({"type": "done"})

        t = threading.Thread(target=worker, daemon=True)
        t.start()

        idx = 0
        last_yield_at = time.time()
        while True:
            with emit_lock:
                pending = events[idx:]
                idx = len(events)
            for ev in pending:
                yield sse(ev)
                last_yield_at = time.time()
                if ev.get("type") == "done":
                    return
            # Keepalive: шлём комментарий каждую секунду, чтобы соединение не закрылось
            now = time.time()
            if now - last_yield_at > 1:
                yield ":keepalive\n\n"
                last_yield_at = now
            if not t.is_alive() and idx >= len(events):
                yield sse({"type": "done"})
                return
            time.sleep(0.03)

    return Response(generate(), mimetype="text/event-stream",
                     headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/confirm", methods=["POST"])
def api_confirm():
    body = request.get_json(force=True)
    cid = body.get("id")
    approved = bool(body.get("approved"))
    entry = pending_confirmations.get(cid)
    if not entry:
        return jsonify({"ok": False, "error": "unknown id"}), 404
    entry["decision"] = approved
    entry["event"].set()
    return jsonify({"ok": True})


@app.route("/api/chat/stop", methods=["POST"])
def api_chat_stop():
    """Экстренно останавливает текущую генерацию ответа — закрывает соединение
    с Ollama и прекращает дальнейшую обработку на сервере."""
    state.active_cancel["cancelled"] = True
    return jsonify({"ok": True})


# ----------------------------- Прочие эндпоинты -------------------------------

@app.route("/api/state", methods=["GET"])
def api_state():
    return jsonify({
        "model": state.model,
        "vision_model": state.vision_model,
        "allowed_folders": [str(f) for f in state.allowed_folders],
        "auto_approve": state.auto_approve,
        "ollama_host": state.ollama_host,
        "reasoning_level": state.reasoning_level,
        "num_ctx": state.num_ctx,
        "use_cpu": state.use_cpu,
        "use_ram": state.use_ram,
        "use_vram": state.use_vram,
        "think_enabled": state.think_enabled,
        "sudo_password": state.sudo_password,
        "ai_autonomy": state.ai_autonomy,
        "google_api_key": state.google_api_key,
        "google_cx": state.google_cx,
        "smtp_host": state.smtp_host,
        "smtp_port": state.smtp_port,
        "smtp_user": state.smtp_user,
        "smtp_password": state.smtp_password,
        "smtp_from": state.smtp_from,
        "telegram_bot_token": state.telegram_bot_token,
        "telegram_chat_id": state.telegram_chat_id,
        "ntfy_topic": state.ntfy_topic,
        "solutions_archive_path": state.solutions_archive_path,
        "knowledge_priority": state.knowledge_priority,
        "reasoning_levels": [
            {"id": k, "label": v["label"]} for k, v in REASONING_LEVELS.items()
        ],
    })


@app.route("/api/settings", methods=["POST"])
def api_settings():
    body = request.get_json(force=True)
    if "model" in body and body["model"]:
        state.model = body["model"]
    if "vision_model" in body:
        state.vision_model = body["vision_model"] or None
    if "auto_approve" in body:
        state.auto_approve = bool(body["auto_approve"])
    if "reasoning_level" in body and body["reasoning_level"] in REASONING_LEVELS:
        state.reasoning_level = body["reasoning_level"]
    if "num_ctx" in body:
        try:
            nc = int(body["num_ctx"])
            if nc <= 0:
                return jsonify({"ok": False, "error": "num_ctx должен быть положительным числом"}), 400
            state.num_ctx = nc
        except (TypeError, ValueError):
            return jsonify({"ok": False, "error": "num_ctx должен быть числом"}), 400
    if "use_cpu" in body:
        state.use_cpu = bool(body["use_cpu"])
    if "use_ram" in body:
        state.use_ram = bool(body["use_ram"])
    if "use_vram" in body:
        state.use_vram = bool(body["use_vram"])
    if not (state.use_cpu or state.use_ram or state.use_vram):
        state.use_cpu = state.use_ram = state.use_vram = True
    if "think_enabled" in body:
        state.think_enabled = bool(body["think_enabled"])
    if "sudo_password" in body:
        state.sudo_password = str(body["sudo_password"] or "")
    if "ai_autonomy" in body:
        state.ai_autonomy = bool(body["ai_autonomy"])
    if "google_api_key" in body:
        state.google_api_key = str(body["google_api_key"] or "")
    if "google_cx" in body:
        state.google_cx = str(body["google_cx"] or "")
    if "smtp_host" in body:
        state.smtp_host = str(body["smtp_host"] or "")
    if "smtp_port" in body:
        try:
            state.smtp_port = int(body["smtp_port"] or 587)
        except (TypeError, ValueError):
            pass
    if "smtp_user" in body:
        state.smtp_user = str(body["smtp_user"] or "")
    if "smtp_password" in body:
        state.smtp_password = str(body["smtp_password"] or "")
    if "smtp_from" in body:
        state.smtp_from = str(body["smtp_from"] or "")
    if "telegram_bot_token" in body:
        state.telegram_bot_token = str(body["telegram_bot_token"] or "")
    if "telegram_chat_id" in body:
        state.telegram_chat_id = str(body["telegram_chat_id"] or "")
    if "ntfy_topic" in body:
        state.ntfy_topic = str(body["ntfy_topic"] or "")
    if "solutions_archive_path" in body:
        state.solutions_archive_path = str(body["solutions_archive_path"] or "")
    if "knowledge_priority" in body:
        kp = body["knowledge_priority"]
        valid = {"internet", "archive", "memory"}
        if isinstance(kp, list) and set(kp) == valid and len(kp) == 3:
            state.knowledge_priority = kp
        else:
            return jsonify({"ok": False, "error": "knowledge_priority должен содержать ровно internet/archive/memory без повторов"}), 400
    state.save()
    return jsonify({
        "ok": True,
        "use_cpu": state.use_cpu, "use_ram": state.use_ram, "use_vram": state.use_vram,
        "think_enabled": state.think_enabled,
        "ai_autonomy": state.ai_autonomy,
    })


@app.route("/api/server-settings", methods=["POST"])
def api_server_settings():
    """Смена адреса Ollama (сразу) и/или порта самого сайта (требует перезапуска процесса)."""
    body = request.get_json(force=True)
    restarting = False
    new_port = None

    if "ollama_host" in body and body["ollama_host"]:
        state.ollama_host = body["ollama_host"].strip().rstrip("/")
        state.save()

    if "web_port" in body and body["web_port"]:
        try:
            new_port = int(body["web_port"])
        except (TypeError, ValueError):
            return jsonify({"ok": False, "error": "Порт должен быть числом"}), 400
        if new_port != int(os.environ.get("WEB_PORT", WEB_PORT)):
            restarting = True

    if restarting:
        def do_restart():
            time.sleep(0.6)  # даём успеть уйти HTTP-ответу
            env = os.environ.copy()
            env["WEB_PORT"] = str(new_port)
            os.execve(sys.executable, [sys.executable] + sys.argv, env)
        threading.Thread(target=do_restart, daemon=True).start()

    return jsonify({"ok": True, "restarting": restarting, "new_port": new_port})


@app.route("/api/browse", methods=["GET"])
def api_browse():
    """Обзор файловой системы сервера — для выбора .gguf-файла модели через интерфейс."""
    raw = request.args.get("path", "") or str(Path.home())
    p = Path(raw).expanduser()
    if not p.exists():
        p = Path.home()
    if p.is_file():
        p = p.parent

    try:
        entries = []
        for item in sorted(p.iterdir(), key=lambda x: (x.is_file(), x.name.lower())):
            try:
                if item.name.startswith("."):
                    continue
                if item.is_dir():
                    entries.append({"name": item.name, "path": str(item), "type": "dir"})
                elif item.suffix.lower() == ".gguf":
                    size_mb = round(item.stat().st_size / 1_000_000, 1)
                    entries.append({"name": item.name, "path": str(item), "type": "gguf", "size_mb": size_mb})
            except (PermissionError, OSError):
                continue
    except (PermissionError, OSError) as e:
        return jsonify({"error": f"Нет доступа: {e}"}), 403

    return jsonify({
        "current": str(p.resolve()),
        "parent": str(p.resolve().parent) if p.resolve() != p.resolve().parent else None,
        "entries": entries,
    })


@app.route("/api/delete-model", methods=["POST"])
def api_delete_model():
    body = request.get_json(force=True)
    name = (body.get("name") or "").strip()
    if not name:
        return jsonify({"ok": False, "error": "Не указано имя модели"}), 400
    if name == state.model:
        return jsonify({"ok": False, "error": "Нельзя удалить модель, которая сейчас используется как «мозг»."}), 400
    if name == state.vision_model:
        return jsonify({"ok": False, "error": "Нельзя удалить модель, которая сейчас используется как «глаза»."}), 400
    try:
        result = subprocess.run(["ollama", "rm", name], capture_output=True, text=True, timeout=120)
    except FileNotFoundError:
        return jsonify({"ok": False, "error": "Команда 'ollama' не найдена в PATH на сервере."}), 500
    except subprocess.TimeoutExpired:
        return jsonify({"ok": False, "error": "Удаление заняло слишком много времени."}), 500
    if result.returncode != 0:
        return jsonify({"ok": False, "error": result.stderr[-800:] or "Неизвестная ошибка ollama rm"}), 500
    return jsonify({"ok": True})


PERCENT_RE = re.compile(r"(\d{1,3})\s*%")


def stream_process_output(cmd: list):
    """Запускает процесс и построчно отдаёт его вывод, учитывая, что прогресс-бары
    консольных утилит обычно перерисовываются через \\r, а не обычный перенос строки."""
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, bufsize=0)
    buf = b""
    try:
        while True:
            chunk = proc.stdout.read(1)
            if not chunk:
                break
            if chunk in (b"\r", b"\n"):
                line = buf.decode("utf-8", errors="replace").strip()
                buf = b""
                if line:
                    yield line
            else:
                buf += chunk
    finally:
        if buf:
            line = buf.decode("utf-8", errors="replace").strip()
            if line:
                yield line
        proc.wait()
        yield f"__RETURNCODE__{proc.returncode}"


@app.route("/api/import-model", methods=["POST"])
def api_import_model():
    """Импортирует .gguf-файл(ы) как новую модель Ollama (аналог ollama create -f Modelfile),
    транслируя прогресс в реальном времени через SSE, чтобы было видно, что модель
    действительно загружается в Ollama, а не 'зависла'."""
    body = request.get_json(force=True)
    name = (body.get("name") or "").strip()
    gguf_path = body.get("gguf_path")
    mmproj_path = body.get("mmproj_path")

    if not name or not gguf_path:
        return jsonify({"ok": False, "error": "Не указано имя модели или файл .gguf"}), 400
    if not Path(gguf_path).exists():
        return jsonify({"ok": False, "error": "Файл .gguf не найден"}), 400

    modelfile_lines = [f"FROM {gguf_path}"]
    if mmproj_path:
        if not Path(mmproj_path).exists():
            return jsonify({"ok": False, "error": "Файл mmproj не найден"}), 400
        modelfile_lines.append(f"FROM {mmproj_path}")

    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    tmp_modelfile = CONFIG_DIR / f"Modelfile.{uuid.uuid4().hex[:8]}"
    tmp_modelfile.write_text("\n".join(modelfile_lines) + "\n", encoding="utf-8")

    def generate():
        returncode = None
        try:
            for line in stream_process_output(["ollama", "create", name, "-f", str(tmp_modelfile)]):
                if line.startswith("__RETURNCODE__"):
                    returncode = int(line[len("__RETURNCODE__"):])
                    continue
                m = PERCENT_RE.search(line)
                percent = int(m.group(1)) if m else None
                yield sse({"type": "progress", "line": line, "percent": percent})

            if returncode == 0:
                yield sse({"type": "done", "ok": True, "name": name})
            else:
                yield sse({"type": "done", "ok": False,
                           "error": f"ollama create завершился с ошибкой (код {returncode})"})
        except FileNotFoundError:
            yield sse({"type": "done", "ok": False, "error": "Команда 'ollama' не найдена в PATH на сервере."})
        except Exception as e:
            yield sse({"type": "done", "ok": False, "error": str(e)})
        finally:
            try:
                tmp_modelfile.unlink(missing_ok=True)
            except Exception:
                pass

    return Response(generate(), mimetype="text/event-stream",
                     headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/api/rebuild-model", methods=["POST"])
def api_rebuild_model():
    """Пересобирает УЖЕ ИМПОРТИРОВАННУЮ модель Ollama с текущими num_ctx/temperature,
    зашитыми прямо в новую модель через 'FROM <существующая модель>' — файл .gguf
    заново не нужен, Ollama берёт веса из уже загруженной модели."""
    body = request.get_json(force=True)
    base_model = (body.get("base_model") or "").strip()
    new_name = (body.get("new_name") or "").strip()
    num_ctx = body.get("num_ctx")
    temperature = body.get("temperature")

    if not base_model:
        return jsonify({"ok": False, "error": "Не указана исходная модель"}), 400
    if not new_name:
        return jsonify({"ok": False, "error": "Не указано имя новой модели"}), 400
    try:
        num_ctx = int(num_ctx)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "error": "num_ctx должен быть числом"}), 400
    try:
        temperature = float(temperature)
    except (TypeError, ValueError):
        temperature = 0.5

    modelfile_lines = [
        f"FROM {base_model}",
        f"PARAMETER num_ctx {num_ctx}",
        f"PARAMETER temperature {temperature}",
    ]
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    tmp_modelfile = CONFIG_DIR / f"Modelfile.{uuid.uuid4().hex[:8]}"
    tmp_modelfile.write_text("\n".join(modelfile_lines) + "\n", encoding="utf-8")

    def generate():
        returncode = None
        try:
            for line in stream_process_output(["ollama", "create", new_name, "-f", str(tmp_modelfile)]):
                if line.startswith("__RETURNCODE__"):
                    returncode = int(line[len("__RETURNCODE__"):])
                    continue
                m = PERCENT_RE.search(line)
                percent = int(m.group(1)) if m else None
                yield sse({"type": "progress", "line": line, "percent": percent})

            if returncode == 0:
                yield sse({"type": "done", "ok": True, "name": new_name})
            else:
                yield sse({"type": "done", "ok": False,
                           "error": f"ollama create завершился с ошибкой (код {returncode})"})
        except FileNotFoundError:
            yield sse({"type": "done", "ok": False, "error": "Команда 'ollama' не найдена в PATH на сервере."})
        except Exception as e:
            yield sse({"type": "done", "ok": False, "error": str(e)})
        finally:
            try:
                tmp_modelfile.unlink(missing_ok=True)
            except Exception:
                pass

    return Response(generate(), mimetype="text/event-stream",
                     headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


def _try_commands(command_lists, input_data=None):
    """Пробует по очереди несколько команд, возвращает первую успешную попытку
    либо все ошибки, если ни одна не сработала."""
    errors = []
    for cmd in command_lists:
        try:
            result = subprocess.run(cmd, input=input_data, capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return True, " ".join(cmd), result.stdout.strip()
            errors.append(f"{' '.join(cmd)} -> {result.stderr.strip() or result.stdout.strip()}")
        except FileNotFoundError:
            errors.append(f"{' '.join(cmd)} -> команда не найдена")
        except subprocess.TimeoutExpired:
            errors.append(f"{' '.join(cmd)} -> превышено время ожидания")
    return False, None, "\n".join(errors)


def _ollama_is_actually_alive() -> bool:
    """Проверяет, отвечает ли Ollama на запросы прямо сейчас — независимо
    от того, что говорит systemd. Если systemd считает службу остановленной,
    а Ollama при этом реально отвечает — значит работает какой-то ДРУГОЙ,
    не-systemd процесс (например, автозапуск рабочего стола), и journalctl
    показывает не про него."""
    try:
        r = requests.get(f"{state.ollama_host}/api/tags", timeout=3)
        return r.status_code == 200
    except Exception:
        return False


@app.route("/api/ollama/logs", methods=["GET"])
def api_ollama_logs():
    """Отдаёт последние строки лога сервера Ollama — для панели
    'Терминал событий Ollama' в интерфейсе. Ollama может быть запущена
    по-разному (systemd-сервис, systemd --user, вручную/автозапуском
    рабочего стола), поэтому пробуем несколько источников и отдаём
    наиболее СВЕЖИЙ, а не просто первый непустой — иначе старый (но
    непустой) вывод journalctl мог навсегда перекрыть живой файл-лог."""
    try:
        n = max(20, min(int(request.args.get("lines", "300")), 2000))
    except ValueError:
        n = 300

    log_file = Path.home() / ".ollama-server.log"
    # Если файл-лог существует, не пуст и реально свежий (обновлялся совсем
    # недавно — то есть кто-то прямо сейчас в него пишет) — он гораздо
    # надёжнее journalctl, у которого может просто застрять старая, но
    # непустая историческая запись про давно остановленный systemd-юнит.
    if log_file.exists():
        try:
            stat = log_file.stat()
            fresh = (time.time() - stat.st_mtime) < 120  # писали в последние 2 минуты
            if stat.st_size > 0 and fresh:
                result = subprocess.run(["tail", "-n", str(n), str(log_file)],
                                         capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    return jsonify({"ok": True, "source": f"файл {log_file}", "log": result.stdout[-40000:],
                                    "mismatch_warning": None})
        except Exception:
            pass

    attempts = [
        ("systemd (ollama)", ["journalctl", "-u", "ollama", "-n", str(n), "--no-pager", "-o", "cat"]),
        ("systemd --user (ollama)", ["journalctl", "--user", "-u", "ollama", "-n", str(n), "--no-pager", "-o", "cat"]),
    ]
    for source, cmd in attempts:
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout.strip():
                log_text = result.stdout[-40000:]
                tail = log_text[-400:]
                looks_stopped = ("Stopped ollama.service" in tail or "Deactivated successfully" in tail) \
                    and "Listening on" not in tail
                mismatch_warning = None
                if looks_stopped and _ollama_is_actually_alive():
                    mismatch_warning = (
                        "ПРОТИВОРЕЧИЕ: по этому логу служба systemd остановлена, но Ollama прямо "
                        "сейчас реально отвечает на запросы. Значит работает какой-то другой процесс "
                        "Ollama — не тот, что управляется systemd (скорее всего, автозапуск рабочего "
                        "стола) — и его вывод отсюда не виден, лог ниже устарел и не про него.\n"
                        "Проверьте: ls ~/.config/autostart/ | grep -i ollama"
                    )
                return jsonify({"ok": True, "source": source, "log": log_text, "mismatch_warning": mismatch_warning})
        except Exception:
            continue

    # Ни systemd-журнал, ни свежий файл-лог ничего не дали — но если файл
    # СУЩЕСТВУЕТ (просто устарел или пуст), всё равно покажем его, чем ничего.
    if log_file.exists():
        try:
            result = subprocess.run(["tail", "-n", str(n), str(log_file)],
                                     capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout.strip():
                return jsonify({"ok": True, "source": f"файл {log_file} (устарел)", "log": result.stdout[-40000:],
                                "mismatch_warning": None})
        except Exception:
            pass

    alive = _ollama_is_actually_alive()
    return jsonify({
        "ok": False,
        "error": (
            ("Ollama сейчас РАБОТАЕТ и отвечает на запросы, но " if alive else "Ollama ") +
            "журнал недоступен через journalctl — вероятно, она запущена не через systemd, а "
            "напрямую (автозапуском рабочего стола / вручную из терминала), и вывод такого "
            "процесса отсюда физически не виден.\n\n"
            "Чтобы получить живой журнал в этом случае, откройте настройки автозагрузки "
            "(«Автозагрузка» → «Ollama AI Server» → изменить программу) и замените команду\n"
            "    ollama serve\n"
            "на:\n"
            "    bash -c 'ollama serve >> ~/.ollama-server.log 2>&1'\n"
            "Перезапустите Ollama после этого — журнал начнёт писаться в файл "
            f"{log_file}, и эта панель сможет его читать."
        ),
    })


def _kill_stray_ollama_serve():
    """Если Ollama запускается не только через systemd, но ЕЩЁ и через
    автозапуск рабочего стола (GNOME autostart, Exec=ollama serve) — второй,
    пользовательский процесс 'ollama serve' живёт независимо от systemd и
    держит видеопамять. Убиваем его несколькими способами: по имени
    процесса, по PID, слушающему порт 11434 (кто бы это ни был), И
    ОТДЕЛЬНО — дочерний процесс 'llama-server', который реально грузит
    модель в видеопамять. Просто убить родителя 'ollama serve' НЕ убивает
    llama-server — он остаётся сиротой и продолжает занимать VRAM, а при
    следующем запуске Ollama плодит ещё один такой процесс поверх старого
    (видеопамять и ОЗУ забиваются несколькими одновременно загруженными
    копиями модели)."""
    try:
        subprocess.run(["pkill", "-9", "-u", getpass.getuser(), "-f", "ollama serve"],
                        capture_output=True, text=True, timeout=10)
    except Exception:
        pass
    try:
        # Дочерний процесс инференса — держит видеопамять независимо от
        # родителя 'ollama serve'. Ищем по имени бинарника llama-server,
        # запущенного из директории ollama (чтобы не задеть что-то чужое).
        subprocess.run(["pkill", "-9", "-u", getpass.getuser(), "-f", "llama-server"],
                        capture_output=True, text=True, timeout=10)
    except Exception:
        pass
    try:
        # fuser самый надёжный способ узнать, кто реально слушает порт,
        # независимо от того, как называется процесс и от чьего он имени.
        subprocess.run(["fuser", "-k", "-9", "11434/tcp"],
                        capture_output=True, text=True, timeout=10)
    except Exception:
        pass


def _ollama_port_busy() -> bool:
    """Проверить, занят ли порт 11434 прямо сейчас."""
    try:
        result = subprocess.run(["fuser", "11434/tcp"], capture_output=True, text=True, timeout=5)
        return bool(result.stdout.strip())
    except Exception:
        return False


def _start_ollama_directly():
    """Если юнита systemd 'ollama' вообще не существует (Ollama запускается
    только автозапуском рабочего стола или вручную) — 'systemctl start'
    никогда не сработает ('Unit ollama.service not found'), тут нечего
    чинить правами доступа. В этом случае запускаем 'ollama serve' сами,
    отдельным процессом, с выводом в лог-файл — так и кнопка "Запустить"
    заработает, и заодно появится живой источник для панели терминала."""
    ollama_bin = shutil.which("ollama") or "ollama"
    log_file = Path.home() / ".ollama-server.log"
    try:
        f = open(log_file, "a")
        subprocess.Popen(
            [ollama_bin, "serve"],
            stdout=f, stderr=subprocess.STDOUT,
            start_new_session=True,
            cwd=str(Path.home()),
        )
        return True, str(log_file)
    except Exception as e:
        return False, str(e)


@app.route("/api/ollama/start", methods=["POST"])
def api_ollama_start():
    # Сначала на всякий случай убираем любые осиротевшие процессы (см.
    # _kill_stray_ollama_serve) и ждём, пока порт 11434 реально освободится
    # — иначе новый запуск может упасть с "address already in use".
    # 'start' проще и надёжнее 'restart': ему не нужно сначала останавливать
    # ещё что-то живое — просто поднять сервис с нуля.
    _kill_stray_ollama_serve()
    for _ in range(10):
        if not _ollama_port_busy():
            break
        time.sleep(0.3)

    ok, used_cmd, detail = _try_commands([
        ["systemctl", "start", "ollama"],
        ["systemctl", "--user", "start", "ollama"],
    ])
    if ok:
        return jsonify({"ok": True, "message": f"Ollama запущена ({used_cmd})."})

    # "Unit ollama.service not found" означает, что юнита systemd для
    # Ollama попросту НЕТ (а не что нет прав или он настроен иначе) — на
    # этой машине Ollama управляется только автозапуском рабочего стола.
    # В этом случае запускаем процесс напрямую, а не пытаемся чинить
    # несуществующий systemd-сервис.
    if "not found" in (detail or ""):
        started, info = _start_ollama_directly()
        if started:
            time.sleep(1.0)
            return jsonify({
                "ok": True,
                "message": f"Юнита systemd для Ollama нет — запустил процесс напрямую. Журнал пишется в {info}.",
            })
        return jsonify({"ok": False, "error": f"Не удалось запустить Ollama напрямую: {info}"}), 500

    return jsonify({
        "ok": False,
        "error": (
            "Не удалось запустить Ollama автоматически (нет прав или служба настроена иначе). "
            "Выполните вручную в терминале: 'sudo systemctl start ollama' "
            "или, если Ollama запущена вручную: 'ollama serve'.\n\n"
            f"Подробности:\n{detail}"
        ),
    }), 500


@app.route("/api/ollama/stop", methods=["POST"])
def api_ollama_stop():
    body = request.get_json(force=True) or {}
    sudo_password = (body.get("sudo_password") or "").strip() or state.sudo_password

    messages = []

    # --- Остановка Ollama ---
    if sudo_password:
        # Сначала пробуем с sudo-паролем
        ok, used_cmd, detail = _try_commands([
            ["sudo", "-S", "systemctl", "stop", "ollama"],
            ["sudo", "-S", "systemctl", "--user", "stop", "ollama"],
        ], input_data=sudo_password + "\n")
        if not ok:
            # Без sudo тоже пробуем (может, уже остановлена или passwordless)
            ok, used_cmd, detail = _try_commands([
                ["systemctl", "stop", "ollama"],
                ["systemctl", "--user", "stop", "ollama"],
            ])
    else:
        ok, used_cmd, detail = _try_commands([
            ["systemctl", "stop", "ollama"],
            ["systemctl", "--user", "stop", "ollama"],
        ])
    _kill_stray_ollama_serve()
    time.sleep(0.3)
    actually_stopped = not _ollama_port_busy()

    if ok or actually_stopped:
        if ok:
            messages.append(f"Ollama остановлена ({used_cmd}).")
        else:
            messages.append("Ollama остановлена (процесс завершён напрямую).")
    else:
        return jsonify({
            "ok": False,
            "error": (
                "Не удалось остановить Ollama. Выполните вручную:\n"
                "'sudo systemctl stop ollama' или 'killall ollama'.\n\n"
                f"Подробности:\n{detail}"
            ),
        }), 500

    # --- Очистка кэша страниц ---
    cache_ok = False
    if sudo_password:
        try:
            proc = subprocess.run(
                ["sudo", "-S", "bash", "-c", "sync; echo 3 > /proc/sys/vm/drop_caches"],
                input=sudo_password + "\n",
                capture_output=True, text=True, timeout=15,
            )
            cache_ok = proc.returncode == 0
        except Exception:
            pass
    if not cache_ok:
        # Пробуем без пароля (passwordless sudo)
        try:
            proc = subprocess.run(
                ["sudo", "bash", "-c", "sync; echo 3 > /proc/sys/vm/drop_caches"],
                capture_output=True, text=True, timeout=15,
            )
            cache_ok = proc.returncode == 0
        except Exception:
            pass

    if cache_ok:
        messages.append("Кэш системы очищен.")
    else:
        messages.append("Кэш не удалось очистить (проверьте пароль sudo).")

    # --- Сохраняем пароль, если он сработал ---
    if sudo_password and sudo_password != state.sudo_password:
        state.sudo_password = sudo_password
        state.save()

    return jsonify({"ok": True, "message": " ".join(messages), "cache_cleared": cache_ok})


@app.route("/api/clear-cache", methods=["POST"])
def api_clear_cache():
    """Очищает кэш страниц Linux (pagecache, dentries, inodes).
    Использует сохранённый пароль sudo из настроек."""
    sudo_password = state.sudo_password
    cmd = ["sudo", "-S", "bash", "-c", "sync; echo 3 > /proc/sys/vm/drop_caches"]
    try:
        proc = subprocess.run(
            cmd, input=(sudo_password + "\n") if sudo_password else "",
            capture_output=True, text=True, timeout=15,
        )
        if proc.returncode == 0:
            return jsonify({"ok": True, "message": "Кэш системы очищен."})
        # Если пароль не подошёл — пробуем без пароля (passwordless sudo)
        if "authentication" in proc.stderr.lower() or "incorrect" in proc.stderr.lower():
            proc2 = subprocess.run(
                ["sudo", "bash", "-c", "sync; echo 3 > /proc/sys/vm/drop_caches"],
                capture_output=True, text=True, timeout=15,
            )
            if proc2.returncode == 0:
                return jsonify({"ok": True, "message": "Кэш системы очищен."})
        return jsonify({"ok": False, "error": proc.stderr.strip() or "Не удалось очистить кэш — проверьте пароль sudo в настройках"}), 500
    except FileNotFoundError:
        return jsonify({"ok": False, "error": "Команда sudo не найдена."}), 500
    except subprocess.TimeoutExpired:
        return jsonify({"ok": False, "error": "Превышено время ожидания."}), 500


@app.route("/api/open-path", methods=["POST"])
def api_open_path():
    """Открывает папку (или родительскую папку файла) в системном файловом
    менеджере через xdg-open. Работает, т.к. сервер запущен на том же
    компьютере, где сидит пользователь (локальное приложение)."""
    body = request.get_json(force=True)
    path = body.get("path", "")
    try:
        target = safe_path(path)
    except SandboxError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    folder = target if target.is_dir() else target.parent
    if not folder.exists():
        return jsonify({"ok": False, "error": "Папка не найдена."}), 404
    try:
        subprocess.Popen(["xdg-open", str(folder)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except FileNotFoundError:
        return jsonify({"ok": False, "error": "Команда xdg-open не найдена в системе."}), 500
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify({"ok": True})


@app.route("/api/run-file", methods=["POST"])
def api_run_file():
    """Запускает файл: если он исполняемый — выполняет напрямую, иначе
    открывает программой по умолчанию через xdg-open."""
    body = request.get_json(force=True)
    path = body.get("path", "")
    try:
        target = safe_path(path)
    except SandboxError as e:
        return jsonify({"ok": False, "error": str(e)}), 400
    if not target.exists() or not target.is_file():
        return jsonify({"ok": False, "error": "Файл не найден."}), 404
    try:
        if os.access(target, os.X_OK) and target.suffix.lower() not in (".txt", ".md", ".json", ".csv"):
            subprocess.Popen([str(target)], cwd=str(target.parent),
                              stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        else:
            subprocess.Popen(["xdg-open", str(target)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except FileNotFoundError:
        return jsonify({"ok": False, "error": "Команда xdg-open не найдена в системе."}), 500
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify({"ok": True})


@app.route("/api/folders", methods=["GET", "POST", "DELETE"])
def api_folders():
    if request.method == "GET":
        return jsonify({"folders": [str(f) for f in state.allowed_folders]})
    body = request.get_json(force=True)
    if request.method == "POST":
        path = Path(body.get("path", "")).expanduser()
        if not path.exists() or not path.is_dir():
            return jsonify({"ok": False, "error": "Папка не найдена"}), 400
        resolved = path.resolve()
        if resolved not in state.allowed_folders:
            state.allowed_folders.append(resolved)
            state.save()
        return jsonify({"ok": True, "folders": [str(f) for f in state.allowed_folders]})
    if request.method == "DELETE":
        path = body.get("path")
        if path == "all":
            state.allowed_folders = []
        else:
            resolved = Path(path).expanduser().resolve()
            if resolved in state.allowed_folders:
                state.allowed_folders.remove(resolved)
        state.save()
        return jsonify({"ok": True, "folders": [str(f) for f in state.allowed_folders]})


def get_cpu_temp():
    """Пытается получить температуру процессора через датчики psutil (Linux hwmon)."""
    if psutil is None or not hasattr(psutil, "sensors_temperatures"):
        return None
    try:
        temps = psutil.sensors_temperatures()
    except Exception:
        return None
    if not temps:
        return None
    # типичные имена датчиков именно процессора (не диска, не GPU, не Wi-Fi)
    for name in ("coretemp", "k10temp", "zenpower", "cpu_thermal", "acpitz"):
        if name in temps and temps[name]:
            return round(temps[name][0].current, 1)
    for entries in temps.values():
        if entries:
            return round(entries[0].current, 1)
    return None


def get_gpu_temp():
    """Температуры ВСЕХ видеокарт. Возвращает список [{name, temp_c}]."""
    temps = []
    # NVIDIA — может показать все карты
    try:
        out = subprocess.run(
            ["nvidia-smi", "--query-gpu=index,name,temperature.gpu", "--format=csv,noheader,nounits"],
            capture_output=True, text=True, timeout=3,
        )
        if out.returncode == 0 and out.stdout.strip():
            for line in out.stdout.strip().splitlines():
                parts = [p.strip() for p in line.split(",")]
                if len(parts) >= 3:
                    temps.append({"name": f"GPU {parts[0]}: {parts[1]}", "temp_c": float(parts[2])})
            if temps:
                return temps
    except (FileNotFoundError, subprocess.TimeoutExpired, ValueError, Exception):
        pass
    # AMD rocm-smi
    try:
        out = subprocess.run(
            ["rocm-smi", "--showtemp", "--json"],
            capture_output=True, text=True, timeout=3,
        )
        if out.returncode == 0 and out.stdout.strip():
            data = json.loads(out.stdout)
            for card_name, card_data in data.items():
                for key, val in card_data.items():
                    if "edge" in key.lower() or "temperature" in key.lower():
                        temps.append({"name": card_name, "temp_c": float(val)})
                        break
            if temps:
                return temps
    except Exception:
        pass
    # amdgpu sysfs
    try:
        for card_dir in sorted(glob.glob("/sys/class/drm/card*/device")):
            for hw in glob.glob(os.path.join(card_dir, "hwmon", "hwmon*")):
                temp_path = os.path.join(hw, "temp1_input")
                if os.path.exists(temp_path):
                    with open(temp_path) as f:
                        millideg = int(f.read().strip())
                    card_name = os.path.basename(os.path.dirname(card_dir))
                    temps.append({"name": card_name, "temp_c": round(millideg / 1000, 1)})
                    break
    except Exception:
        pass
    return temps


def _safe_int(s):
    try: return int(s)
    except (ValueError, TypeError): return None

def _safe_float(s):
    try: return round(float(s), 1)
    except (ValueError, TypeError): return None

def _get_amd_gpu_names():
    """Получает имена AMD GPU через lspci."""
    names = {}
    try:
        out = subprocess.run(["lspci"], capture_output=True, text=True, timeout=3)
        if out.returncode == 0:
            for line in out.stdout.splitlines():
                if "VGA" in line or "Display" in line:
                    # Формат: "XX:XX.X VGA compatible controller: AMD Device Name"
                    # Нужно взять часть после "controller: " или "controller "
                    match = re.search(r'controller[s]?:\s*(.+)', line)
                    if match:
                        idx = len(names)
                        names[idx] = match.group(1).strip()[:50]
    except Exception:
        pass
    return names


def get_vram_stats():
    """Пытается получить занятую/общую видеопамять ВСЕХ видеокарт.
    Возвращает список [{used_mb, total_mb, source, name, ...}]."""
    gpus = []
    # NVIDIA — расширенный запрос
    try:
        out = subprocess.run(
            ["nvidia-smi", "--query-gpu=index,name,memory.used,memory.total,utilization.gpu,power.draw,clocks.current.graphics,driver_version",
             "--format=csv,noheader,nounits"],
            capture_output=True, text=True, timeout=3,
        )
        if out.returncode == 0 and out.stdout.strip():
            for line in out.stdout.strip().splitlines():
                parts = [p.strip() for p in line.split(",")]
                if len(parts) >= 4:
                    used, total = int(parts[2]), int(parts[3])
                    if total > 0:
                        gpu_entry = {"used_mb": used, "total_mb": total, "source": "nvidia", "name": f"GPU {parts[0]}: {parts[1]}"}
                        if len(parts) >= 8:
                            gpu_entry["utilization"] = _safe_int(parts[4])
                            gpu_entry["power_w"] = _safe_float(parts[5])
                            gpu_entry["clock_mhz"] = _safe_int(parts[6])
                            gpu_entry["driver"] = parts[7] if parts[7] != "[N/A]" else None
                        gpus.append(gpu_entry)
            if gpus:
                return gpus
    except (FileNotFoundError, subprocess.TimeoutExpired, ValueError, Exception):
        pass
    # AMD rocm-smi
    try:
        out = subprocess.run(
            ["rocm-smi", "--showmeminfo", "vram", "--json"],
            capture_output=True, text=True, timeout=3,
        )
        if out.returncode == 0 and out.stdout.strip():
            data = json.loads(out.stdout)
            amd_names = _get_amd_gpu_names()
            for idx, (card_name, card_data) in enumerate(data.items()):
                used = int(card_data.get("VRAM Total Used Memory (B)", 0)) // (1024 * 1024)
                total = int(card_data.get("VRAM Total Memory (B)", 0)) // (1024 * 1024)
                if total > 0:
                    name = amd_names.get(idx, card_name)
                    gpus.append({"used_mb": used, "total_mb": total, "source": "amd", "name": name})
            if gpus:
                return gpus
    except Exception:
        pass
    # amdgpu sysfs
    try:
        amd_names = _get_amd_gpu_names()
        sysfs_idx = 0
        for card_dir in sorted(glob.glob("/sys/class/drm/card*/device")):
            used_path = os.path.join(card_dir, "mem_info_vram_used")
            total_path = os.path.join(card_dir, "mem_info_vram_total")
            if os.path.exists(used_path) and os.path.exists(total_path):
                with open(used_path) as f:
                    used_bytes = int(f.read().strip())
                with open(total_path) as f:
                    total_bytes = int(f.read().strip())
                if total_bytes > 0:
                    name = amd_names.get(sysfs_idx, f"GPU {sysfs_idx}")
                    gpus.append({
                        "used_mb": used_bytes // (1024 * 1024),
                        "total_mb": total_bytes // (1024 * 1024),
                        "source": "amdgpu-sysfs",
                        "name": name,
                    })
                    sysfs_idx += 1
    except Exception:
        pass
    return gpus


@app.route("/api/system-stats", methods=["GET"])
def api_system_stats():
    stats = {"cpu": None, "ram": None, "gpus": []}

    if psutil is not None:
        try:
            cpu_percent = psutil.cpu_percent(interval=0.15)
            per_core = psutil.cpu_percent(interval=0, percpu=True)
            freq = psutil.cpu_freq()
            stats["cpu"] = {
                "percent": round(cpu_percent, 1),
                "freq_mhz": round(freq.current, 0) if freq else None,
                "temp_c": get_cpu_temp(),
                "cores": len(per_core),
                "per_core": [round(c, 1) for c in per_core],
            }
        except Exception:
            pass
        try:
            mem = psutil.virtual_memory()
            total = mem.total
            used = total - mem.available  # = приложения + кэш + буферы + slab
            cached = getattr(mem, 'cached', 0)
            buffers = getattr(mem, 'buffers', 0)
            slab = getattr(mem, 'slab', 0)
            shared = getattr(mem, 'shared', 0)
            free = getattr(mem, 'free', 0)
            # Кэш как в системном мониторе: cached + buffers + slab
            cache_total = cached + buffers + slab
            def _gb(b): return round(b / 1_000_000_000, 2)  # ГБ (1000^3), как в системном мониторе
            stats["ram"] = {
                "used_gb": _gb(used),
                "total_gb": _gb(total),
                "percent": round(used / total * 100, 1),
                "cached_gb": _gb(cache_total),
                "cached_only_gb": _gb(cached),
                "buffers_gb": _gb(buffers),
                "slab_gb": _gb(slab),
                "shared_gb": _gb(shared),
                "free_gb": _gb(free),
                "app_used_gb": _gb(max(0, used - cache_total)),
            }
        except Exception:
            pass
    else:
        stats["error"] = "psutil не установлен на сервере — выполните pip install -r requirements-web.txt"

    gpus_info = get_vram_stats()
    gpu_temps = get_gpu_temp()
    temp_map = {}
    for t in gpu_temps:
        temp_map[t["name"]] = t["temp_c"]

    for i, gpu in enumerate(gpus_info):
        temp = temp_map.get(gpu["name"])
        if temp is None and i < len(gpu_temps):
            temp = gpu_temps[i]["temp_c"]
        stats["gpus"].append({
            "used_mb": gpu["used_mb"],
            "total_mb": gpu["total_mb"],
            "percent": round(gpu["used_mb"] / gpu["total_mb"] * 100, 1) if gpu["total_mb"] else None,
            "source": gpu["source"],
            "name": gpu.get("name", f"GPU {i}"),
            "temp_c": temp,
            "utilization": gpu.get("utilization"),
            "power_w": gpu.get("power_w"),
            "clock_mhz": gpu.get("clock_mhz"),
            "driver": gpu.get("driver"),
        })

    return jsonify(stats)


@app.route("/api/models", methods=["GET"])
def api_models():
    try:
        resp = requests.get(f"{state.ollama_host}/api/tags", timeout=10)
        resp.raise_for_status()
        return jsonify(resp.json())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/clear", methods=["POST"])
def api_clear():
    cid = state.new_chat()
    return jsonify({"ok": True, "chat_id": cid})


@app.route("/api/chats", methods=["GET"])
def api_chats_list():
    items = sorted(state.chats.values(), key=lambda c: c["updated_at"], reverse=True)
    return jsonify({
        "current_chat_id": state.current_chat_id,
        "chats": [{"id": c["id"], "title": c["title"], "updated_at": c["updated_at"]} for c in items],
    })


@app.route("/api/chats/<chat_id>/select", methods=["POST"])
def api_chats_select(chat_id):
    if chat_id not in state.chats:
        return jsonify({"ok": False, "error": "Чат не найден"}), 404
    state.current_chat_id = chat_id
    state.save_chats()
    chat = state.chats[chat_id]
    visible = [m for m in chat["messages"] if m.get("role") in ("user", "assistant") and m.get("content")]
    activity = chat.get("activity", [])
    return jsonify({"ok": True, "id": chat_id, "title": chat["title"], "messages": visible, "activity": activity,
                     "generating": bool(chat.get("generating", False))})


@app.route("/api/chats/<chat_id>/activity", methods=["GET"])
def api_chats_activity_get(chat_id):
    """Получить активность чата."""
    if chat_id not in state.chats:
        return jsonify({"ok": False, "error": "Чат не найден"}), 404
    activity = state.chats[chat_id].get("activity", [])
    return jsonify({"ok": True, "activity": activity})


@app.route("/api/scheduled-chat", methods=["GET"])
def api_scheduled_chat():
    """Получить ID чата последнего запланированного задания."""
    global _last_scheduled_chat_id
    if _last_scheduled_chat_id:
        cid = _last_scheduled_chat_id
        return jsonify({"ok": True, "chat_id": cid})
    return jsonify({"ok": False})


@app.route("/api/scheduled-chat/confirm", methods=["POST"])
def api_scheduled_chat_confirm():
    """Подтвердить что чат задания открыт."""
    global _last_scheduled_chat_id
    _last_scheduled_chat_id = None
    return jsonify({"ok": True})


@app.route("/api/chats/<chat_id>/activity", methods=["POST"])
def api_chats_activity(chat_id):
    """Добавить запись активности в чат."""
    if chat_id not in state.chats:
        return jsonify({"ok": False, "error": "Чат не найден"}), 404
    body = request.get_json(force=True)
    entry = body.get("entry")
    if not entry:
        return jsonify({"ok": False, "error": "Пустая запись"}), 400
    chat = state.chats[chat_id]
    if "activity" not in chat:
        chat["activity"] = []
    chat["activity"].append(entry)
    # Ограничить историю — не больше 200 записей
    if len(chat["activity"]) > 200:
        chat["activity"] = chat["activity"][-200:]
    chat["updated_at"] = time.time()
    state.save_chats()
    return jsonify({"ok": True})


@app.route("/api/chats/<chat_id>/activity", methods=["DELETE"])
def api_chats_activity_clear(chat_id):
    """Очистить активность чата."""
    if chat_id not in state.chats:
        return jsonify({"ok": False, "error": "Чат не найден"}), 404
    state.chats[chat_id]["activity"] = []
    state.chats[chat_id]["updated_at"] = time.time()
    state.save_chats()
    return jsonify({"ok": True})


@app.route("/api/chats/<chat_id>", methods=["DELETE"])
def api_chats_delete(chat_id):
    if chat_id not in state.chats:
        return jsonify({"ok": False, "error": "Чат не найден"}), 404
    del state.chats[chat_id]
    if not state.chats:
        state.new_chat()
    elif state.current_chat_id == chat_id:
        newest = max(state.chats.values(), key=lambda c: c["updated_at"])
        state.current_chat_id = newest["id"]
    state.save_chats()
    return jsonify({"ok": True, "current_chat_id": state.current_chat_id})


# ----------------------------- API заданий ------------------------------------

@app.route("/api/tasks", methods=["GET"])
def api_tasks_list():
    return jsonify({"tasks": tasks_store.list_all()})


@app.route("/api/tasks", methods=["POST"])
def api_tasks_create():
    body = request.get_json(force=True)
    title = (body.get("title") or "").strip()
    description = (body.get("description") or "").strip()
    files = body.get("files", [])
    scheduled_time = body.get("scheduled_time") or None
    repeat_daily = bool(body.get("repeat_daily", False))
    command = (body.get("command") or "").strip() or None
    sudo_password = (body.get("sudo_password") or "").strip() or None
    run_file = (body.get("run_file") or "").strip() or None
    if not title:
        return jsonify({"ok": False, "error": "Укажите название задания"}), 400
    task = tasks_store.create(title, description, files, scheduled_time, repeat_daily, command, sudo_password, run_file)
    return jsonify({"ok": True, "task": task})


@app.route("/api/tasks/<task_id>", methods=["GET"])
def api_tasks_get(task_id):
    task = tasks_store.get(task_id)
    if not task:
        return jsonify({"ok": False, "error": "Задание не найдено"}), 404
    return jsonify({"ok": True, "task": task})


@app.route("/api/tasks/<task_id>", methods=["PUT"])
def api_tasks_update(task_id):
    task = tasks_store.get(task_id)
    if not task:
        return jsonify({"ok": False, "error": "Задание не найдено"}), 404
    if task["status"] == "in_progress":
        return jsonify({"ok": False, "error": "Нельзя редактировать задание, пока оно выполняется"}), 400

    body = request.get_json(force=True)
    title = (body.get("title") or "").strip()
    if not title:
        return jsonify({"ok": False, "error": "Укажите название задания"}), 400

    # Пароль не приходит обратно на фронтенд при редактировании (из
    # соображений безопасности) — если поле оставили пустым, сохраняем
    # прежний пароль, а не затираем его.
    new_password = (body.get("sudo_password") or "").strip()
    sudo_password = new_password or task.get("sudo_password")

    updated = tasks_store.update(
        task_id,
        title=title,
        description=(body.get("description") or "").strip(),
        files=body.get("files", task.get("files", [])),
        scheduled_time=body.get("scheduled_time") or None,
        repeat_daily=bool(body.get("repeat_daily", False)),
        command=(body.get("command") or "").strip() or None,
        sudo_password=sudo_password,
        run_file=(body.get("run_file") or "").strip() or None,
    )
    return jsonify({"ok": True, "task": updated})


@app.route("/api/tasks/<task_id>", methods=["DELETE"])
def api_tasks_delete(task_id):
    ok = tasks_store.delete(task_id)
    if not ok:
        return jsonify({"ok": False, "error": "Задание не найдено"}), 404
    return jsonify({"ok": True})


@app.route("/api/tasks/<task_id>/execute", methods=["POST"])
def api_tasks_execute(task_id):
    """Запускает ИИ-выполнение задания: разбивает на шаги и выполняет поэтапно."""
    task = tasks_store.get(task_id)
    if not task:
        return jsonify({"ok": False, "error": "Задание не найдено"}), 404
    if task["status"] == "in_progress":
        return jsonify({"ok": False, "error": "Задание уже выполняется"}), 400

    # Сбросить статус
    tasks_store.update(task_id, status="in_progress", steps=[], result=None)

    target_chat_id = state.current_chat_id
    cancel_flag = {"cancelled": False}
    state.active_cancel = cancel_flag

    def generate():
        events = []
        emit_lock = threading.Lock()

        def emit(ev):
            with emit_lock:
                events.append(ev)
            # Помимо очереди для SSE, сразу сохраняем карточку файловых
            # изменений в постоянный лог активности чата — панель
            # "Активность в файлах" читает именно его, а не поток SSE.
            if ev.get("type") == "file_change":
                push_activity(target_chat_id, activity_card_file_change(
                    ev.get("kind", ""), ev.get("path", ""), ev.get("preview", ""), bool(ev.get("auto_approved"))))

        # Собрать текст задания с вложениями
        parts = []
        if task["files"]:
            for f in task["files"]:
                fname = f.get("name", "файл")
                fcontent = f.get("content", "")
                ftype = f.get("type", "")
                ext = Path(fname).suffix.lower()
                is_image = ext in IMAGE_EXTENSIONS or ftype.startswith("image/")
                if not is_image and fcontent:
                    parts.append(f"--- Файл: {fname} ---\n{fcontent}\n--- Конец файла: {fname} ---")
                elif is_image:
                    parts.append(f"[Изображение: {fname}]")
                else:
                    parts.append(f"[Вложение: {fname}]")

        task_text = f"ЗАДАНИЕ: {task['title']}\n\n{task['description']}"
        if parts:
            task_text += "\n\nПрикреплённые файлы:\n" + "\n".join(parts)

        # Системный промпт для выполнения заданий
        task_system = """Ты — ИИ-ассистент, выполняющий задания поэтапно.
Пользователь даёт тебе задание с описанием и возможно файлами.
Твоя задача:
1. Проанализируй задание и определи конкретные шаги для его выполнения.
2. Для каждого шага используй инструменты (read_file, write_file, edit_file, list_dir, search_files) если нужно.
3. Выполняй шаги последовательно.
4. По завершении каждого шага сообщи о результате.
5. В самом конце дай краткий итог выполнения.

Формат ответа: сначала перечисли шаги, потом выполняй каждый. Используй маркированный список для шагов."""

        folders = "\n".join(f"- {f}" for f in state.allowed_folders) if state.allowed_folders else "(нет)"
        task_system += f"\n\nПодключённые папки:\n{folders}"
        if state.think_enabled:
            task_system += (
                "\n\nКогда рассуждаешь в блоке размышлений (thinking): думай СТРОГО "
                "на русском языке, никогда не переключайся на английский. Не "
                "пересказывай и не повторяй задание в начале рассуждения — сразу "
                "переходи к сути анализа."
            )

        user_msg = {"role": "user", "content": task_text, "time": time.time()}
        # Добавить изображения если есть
        all_images = []
        for f in task["files"]:
            fname = f.get("name", "")
            fcontent = f.get("content", "")
            ftype = f.get("type", "")
            ext = Path(fname).suffix.lower()
            is_image = ext in IMAGE_EXTENSIONS or ftype.startswith("image/")
            if is_image and fcontent:
                all_images.append(fcontent)
        if all_images:
            user_msg["images"] = all_images

        messages = [{"role": "system", "content": task_system}, user_msg]

        # Сразу записать сообщение с заданием в ПОСТОЯННУЮ историю чата и
        # пометить чат как "идёт генерация" — раньше это писалось в чат
        # только в самом конце (после десятков секунд/минут работы), поэтому
        # если пользователь уходил на другой чат и возвращался, сервер отдавал
        # пустую/старую историю, а весь видимый прогресс существовал только
        # в DOM браузера и терялся при следующем переключении.
        cur = state.messages_for(target_chat_id)
        cur.append(user_msg)
        state.chats[target_chat_id]["generating"] = True
        state.touch(target_chat_id, maybe_title=task["title"])

        def worker():
            full_content = ""
            token_count = 0
            cur_local = state.messages_for(target_chat_id)
            try:
                for chunk in stream_ollama(messages, cancel_flag=cancel_flag):
                    if "__cancelled__" in chunk:
                        if cur_local and cur_local[-1].get("role") == "assistant" and cur_local[-1].get("_partial"):
                            cur_local.pop()
                        tasks_store.update(task_id, status="failed", result="Отменено пользователем")
                        emit({"type": "cancelled"})
                        emit({"type": "done"})
                        return
                    if "__replace_history__" in chunk:
                        continue
                    msg = chunk.get("message", {})
                    delta = msg.get("content", "")
                    if delta:
                        full_content += delta
                        token_count += 1
                        emit({"type": "token", "content": delta})
                        # Периодически сохранять частичный ТЕКСТ ответа модели
                        # (не служебные статусы инструментов — те теперь видны
                        # только карточками в "Активности в файлах"), чтобы
                        # прогресс не терялся при переключении чатов.
                        if token_count % 15 == 0:
                            if cur_local and cur_local[-1].get("role") == "assistant" and cur_local[-1].get("_partial"):
                                cur_local.pop()
                            cur_local.append({"role": "assistant", "content": full_content.strip(), "_partial": True, "time": time.time()})
                            state.touch(target_chat_id)
                    if msg.get("tool_calls"):
                        # Убрать частичный текст перед вызовом инструментов —
                        # дальше либо придёт новый текст, либо снова инструмент.
                        if cur_local and cur_local[-1].get("role") == "assistant" and cur_local[-1].get("_partial"):
                            cur_local.pop()
                        # Выполнить инструменты
                        for call in msg["tool_calls"]:
                            fn = call.get("function", {})
                            name = fn.get("name")
                            args = fn.get("arguments") or {}
                            if isinstance(args, str):
                                try:
                                    args = json.loads(args)
                                except json.JSONDecodeError:
                                    args = {}
                            emit({"type": "tool_call", "name": name, "args": args})
                            push_activity(target_chat_id, activity_card_tool_call(name, args))

                            try:
                                if name == "list_dir":
                                    result = tool_list_dir(args.get("path", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "read_file":
                                    result = tool_read_file(args.get("path", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "search_files":
                                    result = tool_search_files(args.get("query", ""), args.get("path", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "write_file":
                                    # force_auto=True: у этого запуска нет активного
                                    # SSE-слушателя, который мог бы нажать "разрешить" —
                                    # ждать подтверждения было бы некому.
                                    result = tool_write_file(args.get("path", ""), args.get("content", ""), emit, force_auto=True)
                                elif name == "edit_file":
                                    result = tool_edit_file(args.get("path", ""), args.get("old_text", ""),
                                                           args.get("new_text", ""), emit, force_auto=True)
                                elif name == "run_shell_command":
                                    result = tool_run_shell(args.get("command", ""), emit,
                                                             use_sudo=bool(args.get("use_sudo")), force_auto=True)
                                elif name == "web_search":
                                    result = tool_web_search(args.get("query", ""), args.get("max_results", 5))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "fetch_url":
                                    result = tool_fetch_url(args.get("url", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "wikipedia_search":
                                    result = tool_wikipedia_search(args.get("query", ""), args.get("lang", "ru"),
                                                                    args.get("max_results", 3))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "google_search":
                                    result = tool_google_search(args.get("query", ""), args.get("max_results", 5))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "github_search":
                                    result = tool_github_search(args.get("query", ""), args.get("max_results", 5))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "image_search":
                                    result = tool_image_search(args.get("query", ""), args.get("max_results", 6))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "video_search":
                                    result = tool_video_search(args.get("query", ""), args.get("max_results", 4))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "download_file":
                                    result = tool_download_file(args.get("url", ""), args.get("path", ""), emit, force_auto=True)
                                elif name == "read_document":
                                    result = tool_read_document(args.get("path", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "ocr_image":
                                    result = tool_ocr_image(args.get("path", ""), args.get("lang", "rus+eng"))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "archive_extract":
                                    result = tool_archive_extract(args.get("path", ""), args.get("dest", ""), emit, force_auto=True)
                                elif name == "archive_create":
                                    result = tool_archive_create(args.get("source", ""), args.get("dest", ""), emit, force_auto=True)
                                elif name == "diff_files":
                                    result = tool_diff_files(args.get("path_a", ""), args.get("path_b", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "remember":
                                    result = tool_remember(args.get("key", ""), args.get("text", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "recall":
                                    result = tool_recall(args.get("key", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "save_solution":
                                    result = tool_save_solution(args.get("task", ""), args.get("solution", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "recall_solutions":
                                    result = tool_recall_solutions(args.get("query", ""), args.get("max_results", 5))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "process_list":
                                    result = tool_process_list(args.get("filter_text", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "check_process_running":
                                    result = tool_check_process_running(args.get("hints", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "git_tool":
                                    result = tool_git(args.get("action", ""), args.get("repo_path", ""),
                                                       args.get("args", ""), emit, force_auto=True)
                                elif name == "systemd_control":
                                    result = tool_systemd(args.get("action", ""), args.get("service", ""), emit,
                                                           use_sudo=args.get("use_sudo", True), force_auto=True)
                                elif name == "notify":
                                    result = tool_notify(args.get("message", ""), args.get("title", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "send_email":
                                    result = tool_send_email(args.get("to", ""), args.get("subject", ""),
                                                              args.get("body", ""), emit, force_auto=True)
                                elif name == "telegram_send":
                                    result = tool_telegram_send(args.get("message", ""), emit, force_auto=True)
                                elif name == "python_repl":
                                    result = tool_python_repl(args.get("code", ""), target_chat_id, emit, force_auto=True)
                                elif name == "lint_format":
                                    result = tool_lint_format(args.get("path", ""), args.get("tool", "ruff"),
                                                               bool(args.get("fix", False)), emit, force_auto=True)
                                elif name == "get_weather":
                                    result = tool_get_weather(args.get("location", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "get_exchange_rate":
                                    result = tool_get_exchange_rate(args.get("amount", 1), args.get("from_currency", ""),
                                                                     args.get("to_currency", ""))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "rss_read":
                                    result = tool_rss_read(args.get("url", ""), args.get("max_items", 10))
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "set_reminder":
                                    result = tool_set_reminder(args.get("minutes", 0), args.get("message", ""), target_chat_id)
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                elif name == "view_image":
                                    img_b64, result = tool_view_image(**args)
                                    if img_b64 and state.vision_model and state.vision_model != state.model:
                                        question = ("Подробно опиши, что изображено на этой картинке, включая "
                                                    "весь видимый текст, интерфейсы, диаграммы и детали.")
                                        description = ask_vision_model(img_b64, question)
                                        result = f"[Описание от vision-модели '{state.vision_model}']\n{description}"
                                    elif img_b64:
                                        messages.append({"role": "user", "content": f"Изображение '{args.get('path','')}':",
                                                          "images": [img_b64]})
                                    push_activity(target_chat_id, activity_card_result(str(result), ok=True))
                                else:
                                    result = f"Неизвестный инструмент: {name}"
                            except SandboxError as e:
                                result = f"Ошибка доступа: {e}"
                                push_activity(target_chat_id, activity_card_result(str(e), ok=False))
                            except Exception as e:
                                result = f"Ошибка выполнения: {e}"
                                push_activity(target_chat_id, activity_card_result(str(e), ok=False))
                            tool_msg = {"role": "tool", "content": str(result)}
                            messages.append(tool_msg)
                            emit({"type": "tool_result", "name": name, "result": str(result)})

                        # После инструментов — продолжить генерацию
                        continue

                    if chunk.get("done"):
                        break

            except Exception as e:
                if cur_local and cur_local[-1].get("role") == "assistant" and cur_local[-1].get("_partial"):
                    cur_local.pop()
                tasks_store.update(task_id, status="failed", result=str(e))
                emit({"type": "error", "message": friendly_ollama_error(e)})
                emit({"type": "done"})
                return
            finally:
                if target_chat_id in state.chats:
                    state.chats[target_chat_id]["generating"] = False
                state._save_chats_now()

            # Сохранить результат
            tasks_store.update(task_id, status="completed", result=full_content)

            # Сохранить в чат (убрав временный частичный текст)
            if cur_local and cur_local[-1].get("role") == "assistant" and cur_local[-1].get("_partial"):
                cur_local.pop()
            cur_local.append({"role": "assistant", "content": full_content, "time": time.time()})
            state.touch(target_chat_id, maybe_title=task["title"])

            emit({"type": "task_done", "task_id": task_id, "result": full_content})
            emit({"type": "done"})

        t = threading.Thread(target=worker, daemon=True)
        t.start()

        idx = 0
        last_yield_at = time.time()
        while True:
            with emit_lock:
                pending = events[idx:]
                idx = len(events)
            for ev in pending:
                yield sse(ev)
                last_yield_at = time.time()
                if ev.get("type") == "done":
                    return
            now = time.time()
            if now - last_yield_at > 1:
                yield ":keepalive\n\n"
                last_yield_at = now
            if not t.is_alive() and idx >= len(events):
                yield sse({"type": "done"})
                return
            time.sleep(0.03)

    return Response(generate(), mimetype="text/event-stream",
                     headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ----------------------------- Пользовательские команды терминала -------------

def _load_terminal_commands() -> list:
    if TERMINAL_COMMANDS_FILE.exists():
        try:
            data = json.loads(TERMINAL_COMMANDS_FILE.read_text(encoding="utf-8"))
            return data.get("commands", []) if isinstance(data, dict) else data
        except Exception:
            pass
    return []


def _save_terminal_commands(commands: list):
    try:
        CONFIG_DIR.mkdir(parents=True, exist_ok=True)
        TERMINAL_COMMANDS_FILE.write_text(
            json.dumps({"commands": commands}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        pass


@app.route("/api/terminal-commands", methods=["GET"])
def api_terminal_commands_list():
    return jsonify({"ok": True, "commands": _load_terminal_commands()})


@app.route("/api/terminal-commands", methods=["POST"])
def api_terminal_commands_create():
    body = request.get_json(force=True)
    command = (body.get("command") or "").strip()
    label = (body.get("label") or "").strip()
    if not command:
        return jsonify({"ok": False, "error": "Команда не может быть пустой"}), 400
    if not label:
        label = command[:30]
    commands = _load_terminal_commands()
    cid = uuid.uuid4().hex[:8]
    entry = {"id": cid, "command": command, "label": label, "created_at": time.time()}
    commands.append(entry)
    _save_terminal_commands(commands)
    return jsonify({"ok": True, "command": entry})


@app.route("/api/terminal-commands/<cmd_id>", methods=["DELETE"])
def api_terminal_commands_delete(cmd_id):
    commands = _load_terminal_commands()
    commands = [c for c in commands if c.get("id") != cmd_id]
    _save_terminal_commands(commands)
    return jsonify({"ok": True})


@app.route("/api/terminal-commands/<cmd_id>", methods=["PUT"])
def api_terminal_commands_update(cmd_id):
    body = request.get_json(force=True)
    command = (body.get("command") or "").strip()
    label = (body.get("label") or "").strip()
    if not command:
        return jsonify({"ok": False, "error": "Команда не может быть пустой"}), 400
    commands = _load_terminal_commands()
    for cmd in commands:
        if cmd.get("id") == cmd_id:
            cmd["command"] = command
            cmd["label"] = label or command
            cmd["updated_at"] = time.time()
            _save_terminal_commands(commands)
            return jsonify({"ok": True, "command": cmd})
    return jsonify({"ok": False, "error": "Команда не найдена"}), 404


# ----------------------------- Интерактивный терминал (много сессий) -----------

class InteractiveTerminal:
    """Persistent shell session with non-blocking I/O via a reader thread."""

    def __init__(self):
        self.process: subprocess.Popen | None = None
        self.cwd = str(Path.home())
        self._lock = threading.Lock()
        self._output_buffer: list[str] = []
        self._reader_thread: threading.Thread | None = None

    def start(self):
        with self._lock:
            if self.process and self.process.poll() is None:
                return
            if state.allowed_folders:
                self.cwd = str(state.allowed_folders[0])
            self.process = subprocess.Popen(
                ["/bin/bash", "--norc", "--noprofile", "-i"],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=0,
                cwd=self.cwd,
            )
            self._output_buffer = []
            self._reader_thread = threading.Thread(target=self._read_loop, daemon=True)
            self._reader_thread.start()

    def _read_loop(self):
        proc = self.process
        if not proc or not proc.stdout:
            return
        while True:
            try:
                ch = proc.stdout.read(1)
                if not ch:
                    break
                with self._lock:
                    self._output_buffer.append(ch)
            except Exception:
                break

    def exec_command(self, command: str) -> str:
        if not self.process or self.process.poll() is not None:
            self.start()
        with self._lock:
            self._output_buffer.clear()
        try:
            self.process.stdin.write(command + "\n")
            self.process.stdin.flush()
        except Exception as e:
            return f"Ошибка записи: {e}"
        import time as _time
        deadline = _time.time() + 30
        collected = []
        while _time.time() < deadline:
            with self._lock:
                collected.extend(self._output_buffer)
                self._output_buffer.clear()
            if collected:
                _time.sleep(0.05)
                with self._lock:
                    if not self._output_buffer:
                        break
                    collected.extend(self._output_buffer)
                    self._output_buffer.clear()
            else:
                _time.sleep(0.05)
        with self._lock:
            collected.extend(self._output_buffer)
            self._output_buffer.clear()
        return "".join(collected)

    def stop(self):
        if self.process and self.process.poll() is None:
            try:
                self.process.stdin.write("exit\n")
                self.process.stdin.flush()
            except Exception:
                pass
            try:
                self.process.terminate()
            except Exception:
                pass
        self.process = None
        with self._lock:
            self._output_buffer.clear()


# История терминалов (session_id -> html)
terminal_history: dict[str, str] = {}


def _load_terminal_history():
    global terminal_history
    if TERMINAL_HISTORY_FILE.exists():
        try:
            terminal_history = json.loads(TERMINAL_HISTORY_FILE.read_text(encoding="utf-8"))
        except Exception:
            terminal_history = {}


def _save_terminal_history():
    try:
        CONFIG_DIR.mkdir(parents=True, exist_ok=True)
        TERMINAL_HISTORY_FILE.write_text(
            json.dumps(terminal_history, ensure_ascii=False), encoding="utf-8"
        )
    except Exception:
        pass


_load_terminal_history()


# Каждая кнопка терминала — своя сессия (session_id = command_id)
terminal_sessions: dict[str, InteractiveTerminal] = {}


def _get_session(session_id: str) -> InteractiveTerminal:
    if session_id not in terminal_sessions:
        terminal_sessions[session_id] = InteractiveTerminal()
    return terminal_sessions[session_id]


@app.route("/api/terminal/start", methods=["POST"])
def api_terminal_start():
    body = request.get_json(force=True)
    session_id = (body.get("session_id") or "").strip()
    command = (body.get("command") or "").strip()
    if not session_id:
        return jsonify({"ok": False, "error": "Нет session_id"}), 400
    session = _get_session(session_id)
    session.start()
    output = ""
    if command:
        output = session.exec_command(command)
    return jsonify({"ok": True, "output": output[-10000:]})


@app.route("/api/terminal/exec", methods=["POST"])
def api_terminal_exec():
    body = request.get_json(force=True)
    session_id = (body.get("session_id") or "").strip()
    command = (body.get("command") or "").strip()
    if not session_id:
        return jsonify({"ok": False, "error": "Нет session_id"}), 400
    if not command:
        return jsonify({"ok": False, "error": "Пустая команда"}), 400
    session = _get_session(session_id)
    output = session.exec_command(command)
    return jsonify({"ok": True, "output": output[-10000:]})


@app.route("/api/terminal/stop", methods=["POST"])
def api_terminal_stop():
    body = request.get_json(force=True)
    session_id = (body.get("session_id") or "").strip()
    if session_id in terminal_sessions:
        terminal_sessions[session_id].stop()
        del terminal_sessions[session_id]
    return jsonify({"ok": True})


@app.route("/api/terminal/history", methods=["GET"])
def api_terminal_history_get():
    return jsonify({"ok": True, "history": terminal_history})


@app.route("/api/terminal/history", methods=["POST"])
def api_terminal_history_save():
    body = request.get_json(force=True)
    session_id = (body.get("session_id") or "").strip()
    html = body.get("html", "")
    if not session_id:
        return jsonify({"ok": False, "error": "Нет session_id"}), 400
    terminal_history[session_id] = html
    _save_terminal_history()
    return jsonify({"ok": True})


@app.route("/api/terminal/history/<session_id>", methods=["DELETE"])
def api_terminal_history_delete(session_id):
    terminal_history.pop(session_id, None)
    _save_terminal_history()
    return jsonify({"ok": True})


@app.route("/")
def index():
    resp = send_from_directory(app.static_folder, "index.html")
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/sw.js")
def service_worker():
    # Браузер обновляет установленный service worker очень редко (обычно раз
    # в сутки, только при обычной навигации) — если он закэширует СТАРУЮ
    # версию этого файла надолго, обновления сайта могут годами "не доезжать"
    # до пользователя, хотя на диске уже другой код. Отдаём его всегда свежим.
    resp = send_from_directory(app.static_folder, "sw.js")
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    resp.headers["Service-Worker-Allowed"] = "/"
    return resp


# ----------------------------- Планировщик заданий ----------------------------

def run_scheduled_tasks():
    """Фоновый поток: проверяет задания и выполняет по расписанию."""
    import datetime
    while True:
        now = datetime.datetime.now()
        for task_id, task in list(tasks_store.tasks.items()):
            if task["status"] == "in_progress":
                continue
            sched = task.get("scheduled_time")
            if not sched:
                continue
            try:
                sched_dt = datetime.datetime.fromisoformat(sched)
            except (ValueError, TypeError):
                continue
            # Проверить пора ли выполнять
            if now >= sched_dt:
                # Проверить не выполнялось ли уже сегодня
                last_run = task.get("last_run")
                today_str = now.strftime("%Y-%m-%d")
                if last_run and last_run.startswith(today_str):
                    continue
                # Если repeat_daily — обновить scheduled_time на следующий день
                if task.get("repeat_daily"):
                    next_day = sched_dt + datetime.timedelta(days=1)
                    tasks_store.update(task_id, scheduled_time=next_day.isoformat())
                else:
                    tasks_store.update(task_id, scheduled_time=None)
                # Выполнить задание в новом чате
                tasks_store.update(task_id, status="in_progress", last_run=now.isoformat())
                _execute_scheduled_task(task_id, task)
        time.sleep(20)


def _execute_scheduled_task(task_id, task):
    """Выполнить запланированное задание."""
    command = task.get("command")
    sudo_password = task.get("sudo_password")
    run_file = task.get("run_file")

    if run_file:
        # Выполнить файл
        _run_file_task(task_id, task, run_file, sudo_password)
    elif command:
        # Выполнить команду в терминале
        _run_terminal_command(task_id, task, command, sudo_password)
    else:
        # Выполнить через ИИ в новом чате
        _run_ai_task(task_id, task)


def _run_file_task(task_id, task, run_file, sudo_password):
    """Выполнить файл (.py, .sh, .html и т.д.)."""
    try:
        file_path = Path(run_file)
        if not file_path.exists():
            tasks_store.update(task_id, status="failed", result=f"Файл не найден: {run_file}")
            return

        ext = file_path.suffix.lower()
        # Определить как запускать по расширению
        if ext == ".py":
            cmd = f"python3 {run_file}"
        elif ext == ".sh":
            cmd = f"bash {run_file}"
        elif ext == ".html":
            cmd = f"xdg-open {run_file}"
        elif ext == ".js":
            cmd = f"node {run_file}"
        else:
            cmd = f"{run_file}"  # попробовать запустить напрямую

        if sudo_password:
            proc = subprocess.Popen(
                ["sudo", "-S", "bash", "-c", cmd],
                stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, timeout=300
            )
            stdout, _ = proc.communicate(input=sudo_password + "\n", timeout=300)
        else:
            proc = subprocess.Popen(
                ["bash", "-c", cmd],
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, timeout=300
            )
            stdout, _ = proc.communicate(timeout=300)

        result = stdout.strip() if stdout else ""
        exit_code = proc.returncode
        status = "completed" if exit_code == 0 else "failed"
        result_text = f"Файл: {run_file}\nКоманда: {cmd}\nКод возврата: {exit_code}\n\n{result}"
        tasks_store.update(task_id, status=status, result=result_text[:2000])
    except subprocess.TimeoutExpired:
        tasks_store.update(task_id, status="failed", result=f"Превышено время ожидания запуска {run_file}")
    except Exception as e:
        tasks_store.update(task_id, status="failed", result=f"Ошибка запуска {run_file}: {e}")


def _run_terminal_command(task_id, task, command, sudo_password):
    """Выполнить shell-команду и отчитаться в чате для ИИ."""
    try:
        if sudo_password:
            proc = subprocess.Popen(
                ["sudo", "-S", "bash", "-c", command],
                stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, timeout=300
            )
            stdout, _ = proc.communicate(input=sudo_password + "\n", timeout=300)
        else:
            proc = subprocess.Popen(
                ["bash", "-c", command],
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, timeout=300
            )
            stdout, _ = proc.communicate(timeout=300)

        result = stdout.strip() if stdout else ""
        exit_code = proc.returncode
        result_text = f"Код возврата: {exit_code}\n\n{result}" if result else f"Код возврата: {exit_code}"
        status = "completed" if exit_code == 0 else "failed"

        # Отправить результат в чат чтобы ИИ видело и могло отреагировать
        chat_id = state.current_chat_id
        if chat_id:
            cur = state.messages_for(chat_id)
            cur.append({"role": "assistant", "content": f"📋 **Результат команды:** `{command}`\n\n{result_text[:2000]}", "time": time.time()})
            state.touch(chat_id, maybe_title=f"Команда: {command[:30]}")
        state.save_chats()
        tasks_store.update(task_id, status=status, result=result_text[:2000])
    except subprocess.TimeoutExpired:
        tasks_store.update(task_id, status="failed", result="Превышено время ожидания (300с)")
    except Exception as e:
        tasks_store.update(task_id, status="failed", result=f"Ошибка: {e}")


def _run_file_task(task_id, task, run_file, sudo_password):
    """Выполнить файл и отчитаться в чате для ИИ."""
    try:
        file_path = Path(run_file)
        if not file_path.exists():
            tasks_store.update(task_id, status="failed", result=f"Файл не найден: {run_file}")
            return

        ext = file_path.suffix.lower()
        if ext == ".py":
            cmd = f"python3 {run_file}"
        elif ext == ".sh":
            cmd = f"bash {run_file}"
        elif ext == ".html":
            cmd = f"xdg-open {run_file}"
        elif ext == ".js":
            cmd = f"node {run_file}"
        else:
            cmd = f"{run_file}"

        if sudo_password:
            proc = subprocess.Popen(
                ["sudo", "-S", "bash", "-c", cmd],
                stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, timeout=300
            )
            stdout, _ = proc.communicate(input=sudo_password + "\n", timeout=300)
        else:
            proc = subprocess.Popen(
                ["bash", "-c", cmd],
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, timeout=300
            )
            stdout, _ = proc.communicate(timeout=300)

        result = stdout.strip() if stdout else ""
        exit_code = proc.returncode
        result_text = f"Файл: {run_file}\nКоманда: {cmd}\nКод возврата: {exit_code}\n\n{result}"
        status = "completed" if exit_code == 0 else "failed"

        # Отправить результат в чат чтобы ИИ видело
        chat_id = state.current_chat_id
        if chat_id:
            cur = state.messages_for(chat_id)
            cur.append({"role": "assistant", "content": f"📋 **Результат запуска файла:** `{run_file}`\n\n{result_text[:2000]}", "time": time.time()})
            state.touch(chat_id, maybe_title=f"Файл: {Path(run_file).name}")
        state.save_chats()
        tasks_store.update(task_id, status=status, result=result_text[:2000])
    except subprocess.TimeoutExpired:
        tasks_store.update(task_id, status="failed", result=f"Превышено время ожидания запуска {run_file}")
    except Exception as e:
        tasks_store.update(task_id, status="failed", result=f"Ошибка запуска {run_file}: {e}")


def _run_ai_task(task_id, task):
    """Выполнить задание через ИИ в новом чате."""
    global _last_scheduled_chat_id
    chat_id = state.new_chat(title=f"Задание: {task['title']}")
    state.current_chat_id = chat_id
    # Флаг "идёт генерация" — фронтенд опрашивает его, чтобы знать, когда
    # обновлять сообщения этого чата без перезагрузки страницы.
    state.chats[chat_id]["generating"] = True
    state.save_chats()
    _last_scheduled_chat_id = chat_id

    parts = []
    if task.get("files"):
        for f in task["files"]:
            fname = f.get("name", "файл")
            fcontent = f.get("content", "")
            ftype = f.get("type", "")
            ext = Path(fname).suffix.lower()
            is_image = ext in IMAGE_EXTENSIONS or ftype.startswith("image/")
            if not is_image and fcontent:
                parts.append(f"--- Файл: {fname} ---\n{fcontent}\n---")
    task_text = f"ЗАДАНИЕ: {task['title']}\n\n{task['description']}"
    if parts:
        task_text += "\n\nФайлы:\n" + "\n".join(parts)

    cur = state.messages_for(chat_id)
    cur.append({"role": "user", "content": task_text, "time": time.time()})
    state.touch(chat_id, maybe_title=task["title"])

    try:
        # Цикл: отправлять → получать tool_calls → повторять пока нет текстового ответа
        messages = [build_system_message(task_text)] + cur
        full_content = ""
        for _ in range(MAX_TOOL_ITERATIONS):
            iteration_content = ""
            tool_calls = None
            token_count = 0
            try:
                for chunk in stream_ollama(messages):
                    if "__cancelled__" in chunk:
                        break
                    if "__replace_history__" in chunk:
                        continue
                    msg = chunk.get("message", {})
                    delta = msg.get("content", "")
                    if delta:
                        iteration_content += delta
                        token_count += 1
                        # Периодически сохранять частичный текст ответа (НЕ
                        # служебные статусы инструментов — те теперь видны
                        # только карточками в "Активности в файлах", чтобы не
                        # дублировать одно и то же в двух местах и не путать
                        # порядок сообщений в чате).
                        if token_count % 15 == 0:
                            display = (full_content + iteration_content).strip()
                            if cur and cur[-1].get("role") == "assistant" and cur[-1].get("_partial"):
                                cur.pop()
                            cur.append({"role": "assistant", "content": display, "_partial": True, "time": time.time()})
                            state.touch(chat_id)
                    if msg.get("tool_calls"):
                        tool_calls = msg["tool_calls"]
                    if chunk.get("done"):
                        break
            except Exception as e:
                full_content += f"\n\nОшибка: {e}"
                break

            full_content += iteration_content

            # Убрать временное частичное сообщение — дальше либо добавим
            # финальный текст, либо перейдём к вызову инструментов.
            if cur and cur[-1].get("role") == "assistant" and cur[-1].get("_partial"):
                cur.pop()

            if not tool_calls:
                break

            # Выполнить tool_calls. В чат добавляем только РЕАЛЬНЫЙ текст
            # модели (если она что-то написала перед вызовом инструмента) —
            # сам факт вызова и его результат отображаются карточками в
            # панели "Активность в файлах", а не отдельными сообщениями в
            # чате (иначе получается дублирование и путаница с порядком).
            assistant_msg = {"role": "assistant", "content": iteration_content, "tool_calls": tool_calls, "time": time.time()}
            messages.append(assistant_msg)
            if iteration_content.strip():
                cur.append(assistant_msg)

            for call in tool_calls:
                fn = call.get("function", {})
                name = fn.get("name")
                args = fn.get("arguments") or {}
                if isinstance(args, str):
                    try: args = json.loads(args)
                    except: args = {}

                # Карточка в панель "Активность в файлах" — та же самая, что
                # видна при интерактивном чате или ручном запуске задания.
                push_activity(chat_id, activity_card_tool_call(name, args))
                bg_emit = make_persisting_emit(chat_id)

                try:
                    if name == "view_image":
                        img_b64, result = tool_view_image(**args)
                        if img_b64 and state.vision_model and state.vision_model != state.model:
                            question = ("Подробно опиши, что изображено на этой картинке, включая "
                                        "весь видимый текст, интерфейсы, диаграммы и детали.")
                            description = ask_vision_model(img_b64, question)
                            tool_result = f"[Описание от vision-модели '{state.vision_model}']\n{description}"
                            push_activity(chat_id, activity_card_result(tool_result, ok=True))
                            tool_msg = {"role": "tool", "content": tool_result, "time": time.time()}
                            messages.append(tool_msg)
                            cur.append(tool_msg)
                            state.touch(chat_id)
                            continue
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                        tool_msg = {"role": "tool", "content": str(result), "time": time.time()}
                        messages.append(tool_msg)
                        cur.append(tool_msg)
                        if img_b64:
                            image_msg = {"role": "user", "content": f"Изображение '{args.get('path','')}':",
                                         "images": [img_b64], "time": time.time()}
                            messages.append(image_msg)
                            cur.append(image_msg)
                        state.touch(chat_id)
                        continue
                    elif name == "write_file":
                        # force_auto=True: задание выполняется в фоне, спросить
                        # разрешение у живого пользователя тут не у кого — иначе
                        # операция зависла бы на 10 минут в ожидании подтверждения.
                        result = tool_write_file(args.get("path", ""), args.get("content", ""), bg_emit, force_auto=True)
                    elif name == "edit_file":
                        result = tool_edit_file(args.get("path", ""), args.get("old_text", ""), args.get("new_text", ""), bg_emit, force_auto=True)
                    elif name == "run_shell_command":
                        result = tool_run_shell(args.get("command", ""), bg_emit,
                                                 use_sudo=bool(args.get("use_sudo")), force_auto=True)
                    elif name == "web_search":
                        result = tool_web_search(args.get("query", ""), args.get("max_results", 5))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "fetch_url":
                        result = tool_fetch_url(args.get("url", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "wikipedia_search":
                        result = tool_wikipedia_search(args.get("query", ""), args.get("lang", "ru"),
                                                        args.get("max_results", 3))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "google_search":
                        result = tool_google_search(args.get("query", ""), args.get("max_results", 5))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "github_search":
                        result = tool_github_search(args.get("query", ""), args.get("max_results", 5))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "image_search":
                        result = tool_image_search(args.get("query", ""), args.get("max_results", 6))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "video_search":
                        result = tool_video_search(args.get("query", ""), args.get("max_results", 4))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "download_file":
                        result = tool_download_file(args.get("url", ""), args.get("path", ""), bg_emit, force_auto=True)
                    elif name == "read_document":
                        result = tool_read_document(args.get("path", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "ocr_image":
                        result = tool_ocr_image(args.get("path", ""), args.get("lang", "rus+eng"))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "archive_extract":
                        result = tool_archive_extract(args.get("path", ""), args.get("dest", ""), bg_emit, force_auto=True)
                    elif name == "archive_create":
                        result = tool_archive_create(args.get("source", ""), args.get("dest", ""), bg_emit, force_auto=True)
                    elif name == "diff_files":
                        result = tool_diff_files(args.get("path_a", ""), args.get("path_b", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "remember":
                        result = tool_remember(args.get("key", ""), args.get("text", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "recall":
                        result = tool_recall(args.get("key", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "save_solution":
                        result = tool_save_solution(args.get("task", ""), args.get("solution", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "recall_solutions":
                        result = tool_recall_solutions(args.get("query", ""), args.get("max_results", 5))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "process_list":
                        result = tool_process_list(args.get("filter_text", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "check_process_running":
                        result = tool_check_process_running(args.get("hints", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "git_tool":
                        result = tool_git(args.get("action", ""), args.get("repo_path", ""),
                                           args.get("args", ""), bg_emit, force_auto=True)
                    elif name == "systemd_control":
                        result = tool_systemd(args.get("action", ""), args.get("service", ""), bg_emit,
                                               use_sudo=args.get("use_sudo", True), force_auto=True)
                    elif name == "notify":
                        result = tool_notify(args.get("message", ""), args.get("title", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "send_email":
                        result = tool_send_email(args.get("to", ""), args.get("subject", ""),
                                                  args.get("body", ""), bg_emit, force_auto=True)
                    elif name == "telegram_send":
                        result = tool_telegram_send(args.get("message", ""), bg_emit, force_auto=True)
                    elif name == "python_repl":
                        result = tool_python_repl(args.get("code", ""), chat_id, bg_emit, force_auto=True)
                    elif name == "lint_format":
                        result = tool_lint_format(args.get("path", ""), args.get("tool", "ruff"),
                                                   bool(args.get("fix", False)), bg_emit, force_auto=True)
                    elif name == "get_weather":
                        result = tool_get_weather(args.get("location", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "get_exchange_rate":
                        result = tool_get_exchange_rate(args.get("amount", 1), args.get("from_currency", ""),
                                                         args.get("to_currency", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "rss_read":
                        result = tool_rss_read(args.get("url", ""), args.get("max_items", 10))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "set_reminder":
                        result = tool_set_reminder(args.get("minutes", 0), args.get("message", ""), chat_id)
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "list_dir":
                        result = tool_list_dir(args.get("path", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "read_file":
                        result = tool_read_file(args.get("path", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    elif name == "search_files":
                        result = tool_search_files(args.get("query", ""), args.get("path", ""))
                        push_activity(chat_id, activity_card_result(str(result), ok=True))
                    else:
                        result = f"Неизвестный: {name}"
                except Exception as e:
                    result = f"Ошибка: {e}"
                    push_activity(chat_id, activity_card_result(str(e), ok=False))
                tool_msg = {"role": "tool", "content": str(result), "time": time.time()}
                messages.append(tool_msg)
                cur.append(tool_msg)
                state.touch(chat_id)

        # Сохранить финальный ответ
        if full_content.strip():
            cur.append({"role": "assistant", "content": full_content.strip(), "time": time.time()})
        state.save_chats()
        tasks_store.update(task_id, status="completed", result=full_content[:2000] if full_content else "Нет ответа")
    finally:
        # Гарантированно снимаем флаг генерации, даже если что-то упало —
        # иначе фронтенд будет бесконечно ждать завершения.
        if chat_id in state.chats:
            state.chats[chat_id]["generating"] = False
        state._save_chats_now()


# Запустить планировщик в фоне
_scheduler_thread = threading.Thread(target=run_scheduled_tasks, daemon=True)
_scheduler_thread.start()


def _chats_saver_loop():
    """Фоновый поток: реально записывает историю чатов на диск не чаще
    раза в секунду, если есть несохранённые изменения. save_chats() в
    "горячих" путях (каждые ~15 токенов, каждый вызов инструмента) теперь
    только выставляет флаг — сама запись (тяжёлая JSON-сериализация всех
    чатов) сюда вынесена, чтобы не тормозить генерацию и не подвешивать
    остальные запросы к серверу (кнопки перезапуска/остановки и т.п.)."""
    while True:
        time.sleep(1.0)
        if state._chats_dirty:
            state._chats_dirty = False
            state._save_chats_now()


_saver_thread = threading.Thread(target=_chats_saver_loop, daemon=True)
_saver_thread.start()


if __name__ == "__main__":
    print(f"Ollama Web запущен: http://localhost:{WEB_PORT}  (сервер Ollama: {state.ollama_host})")
    app.run(host="127.0.0.1", port=WEB_PORT, threaded=True, debug=False)
